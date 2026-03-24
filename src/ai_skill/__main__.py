"""CLI entry point for ai-skill.

Commands:
    workspace begin            Create a new research project workspace (interactive wizard).
    research --begin           Start Research Charter (CP1).
    research --review          Apply corrections from last CP1 preview and regenerate.
    research --signoff         Approve last CP1 preview and generate final file.
    literature --begin         Start Literature Review (CP2).
    literature --review        Apply corrections from last CP2 preview and regenerate.
    literature --signoff       Approve last CP2 preview and generate final file.
    design --begin             Start Research Design (CP3).
    design --review            Apply corrections from last CP3 preview and regenerate.
    design --signoff           Approve last CP3 preview and generate final file.
    workspace list             List all existing workspaces.
    workspace files [name]     List files available inside a workspace.
    status [workspace]         Show the current status of an existing workspace.
    graph                      Print the Mermaid diagram of the pipeline graph.
    skills list                List all registered skills.
    skills show <name>         Show details for a specific skill.
"""

from __future__ import annotations

import logging
import os
import sys
from pathlib import Path
from typing import Annotated

import typer
from rich.console import Console
from rich.logging import RichHandler
from rich.panel import Panel
from rich.table import Table

# LangGraph raises GraphInterrupt when interrupt_before fires (some versions)
try:
    from langgraph.errors import GraphInterrupt as _GRAPH_INTERRUPT  # type: ignore[import-untyped]
except ImportError:
    _GRAPH_INTERRUPT = type("_NeverRaised", (BaseException,), {})  # type: ignore[assignment,misc]

app = typer.Typer(
    name="ai-skill",
    help="AI Academic Research Pipeline Agent",
    add_completion=False,
    no_args_is_help=True,
)
skills_app = typer.Typer(help="Manage and inspect skills.", no_args_is_help=True)
workspace_app = typer.Typer(help="Gerenciar workspaces de pesquisa.", no_args_is_help=True)
app.add_typer(skills_app, name="skills")
app.add_typer(workspace_app, name="workspace")

console = Console()
logger = logging.getLogger(__name__)


def _configure_logging(level: str = "WARNING") -> None:
    """Configure root logging with RichHandler.

    Args:
        level: Log level string (DEBUG, INFO, WARNING, ERROR).
    """
    logging.basicConfig(
        level=getattr(logging, level.upper(), logging.WARNING),
        format="%(message)s",
        datefmt="[%X]",
        handlers=[RichHandler(rich_tracebacks=True, console=Console(stderr=True))],
    )


def _default_workspace(topic: str) -> Path:
    """Return a workspace path derived from the topic string.

    Args:
        topic: The research topic.

    Returns:
        Path inside AI_SKILL_WORKSPACE_DIR (default: ./research-workspace).
    """
    base = Path(os.environ.get("AI_SKILL_WORKSPACE_DIR", "./research-workspace"))
    slug = topic.lower().replace(" ", "-")[:40]
    return base / slug


# ---------------------------------------------------------------------------
# Workspace commands
# ---------------------------------------------------------------------------


@workspace_app.command("begin")
def workspace_begin() -> None:
    """Iniciar um novo workspace de pesquisa (wizard interativo)."""
    from ai_skill.core.project_workspace import ProjectWorkspace, default_projects_root, slugify

    console.print(Panel("[bold]Novo Workspace de Pesquisa[/bold]", expand=False))
    console.print(
        "\nEscolha um [bold]nome curto[/bold] para identificar esta pesquisa.\n"
        "[dim]Use letras, números e hífens. Máximo 50 caracteres.\n"
        "Caracteres especiais e espaços são convertidos automaticamente.[/dim]\n"
    )

    # --- Prompt for name ---
    while True:
        raw_name = typer.prompt("Nome da pesquisa").strip()
        if not raw_name:
            console.print("[red]Nome não pode ser vazio.[/red]")
            continue
        try:
            slug = slugify(raw_name)
        except ValueError as exc:
            console.print(f"[red]{exc}[/red]")
            continue
        if len(slug) > 50:
            console.print(
                f"[red]Nome muito longo ({len(slug)} chars após sanitização). "
                "Use no máximo 50 caracteres.[/red]"
            )
            continue
        console.print(f"  [dim]→ pasta: [/dim][cyan]{slug}[/cyan]")
        break

    # --- Prompt for topic ---
    topic = typer.prompt(
        "\nTópico ou questão de pesquisa (pode ser mais descritivo)",
        default=raw_name,
    ).strip()

    # --- Create workspace ---
    root = default_projects_root()
    ws = ProjectWorkspace(raw_name, root=root)

    if ws.exists():
        console.print(f"\n[yellow]Workspace já existe:[/yellow] {ws.path}")
        if not typer.confirm("Deseja abrir o workspace existente?", default=True):
            raise typer.Exit(0)
    else:
        try:
            ws.create(topic=topic)
        except FileExistsError:
            console.print(f"[red]Workspace já existe:[/red] {ws.path}")
            raise typer.Exit(1)

        console.print(f"\n[green]✓[/green] Workspace criado em: [bold]{ws.path}[/bold]")
        console.print(f"  [green]✓[/green] README.md  — Metodologia da Pesquisa")
        console.print(f"  [green]✓[/green] attachments/  — coloque aqui arquivos para o agente")
        console.print(f"  [green]✓[/green] workspace.yaml  — metadados do projeto")

    console.print(
        f"\n[dim]Para iniciar a pesquisa neste workspace:[/dim]\n"
        f'  ai-skill research --begin --workspace "{ws.path}"'
    )

    if typer.confirm("\nIniciar a pesquisa agora?", default=True):
        begin_research(topic=topic, workspace=ws.path, model="")


@workspace_app.command("list")
def workspace_list() -> None:
    """Listar todos os workspaces de pesquisa disponíveis."""
    from ai_skill.core.project_workspace import default_projects_root, list_workspaces

    root = default_projects_root()
    workspaces = list_workspaces(root)

    if not workspaces:
        console.print(f"[dim]Nenhum workspace encontrado em:[/dim] {root}")
        console.print("[dim]Crie um com:[/dim]  ai-skill workspace begin")
        return

    table = Table(title=f"Workspaces de Pesquisa  [dim]({root})[/dim]")
    table.add_column("#", style="dim", width=3)
    table.add_column("Nome", style="cyan")
    table.add_column("Tópico")
    table.add_column("Checkpoint", justify="center")
    table.add_column("Status")
    table.add_column("Criado em", style="dim")

    for i, ws in enumerate(workspaces, start=1):
        meta = ws.load_metadata()
        table.add_row(
            str(i),
            meta.get("name", ws.slug),
            (meta.get("topic", "?"))[:45] + ("…" if len(meta.get("topic", "")) > 45 else ""),
            str(meta.get("current_checkpoint", 0)),
            meta.get("status", "?"),
            (meta.get("created", "?"))[:10],
        )

    console.print(table)


@workspace_app.command("files")
def workspace_files(
    name: Annotated[
        str | None,
        typer.Argument(help="Nome ou slug do workspace. Omita para seleção interativa."),
    ] = None,
) -> None:
    """Listar arquivos disponíveis em um workspace (checkpoints + attachments)."""
    ws = _resolve_workspace(name)
    if ws is None:
        raise typer.Exit(1)

    files = ws.list_all_files()
    if not files:
        console.print(f"[dim]Nenhum arquivo disponível em:[/dim] {ws.path}")
        console.print(
            f"  Outputs de checkpoint aparecem aqui automaticamente.\n"
            f"  Coloque arquivos de entrada em: [cyan]{ws.attachments_path}[/cyan]"
        )
        return

    console.print(f"\n[bold]Arquivos em:[/bold] {ws.path}\n")
    checkpoints = ws.list_checkpoints()
    attachments = ws.list_attachments()

    if checkpoints:
        console.print("[bold]Checkpoints:[/bold]")
        for i, f in enumerate(checkpoints, start=1):
            console.print(f"  [cyan]{i:>2}.[/cyan] {f.name}")

    if attachments:
        offset = len(checkpoints)
        console.print("\n[bold]Attachments:[/bold]")
        for i, f in enumerate(attachments, start=offset + 1):
            rel = f.relative_to(ws.path)
            console.print(f"  [cyan]{i:>2}.[/cyan] {rel}")

    console.print()


# ---------------------------------------------------------------------------
# Commands
# ---------------------------------------------------------------------------


@app.command("begin-workspace")
def begin_workspace_alias() -> None:
    """Criar um novo workspace de pesquisa (alias de 'workspace begin')."""
    workspace_begin()


@app.command("research")
def research_cmd(
    begin: Annotated[
        bool,
        typer.Option("--begin", is_flag=True, help="Iniciar Research Charter (CP1)."),
    ] = False,
    review: Annotated[
        bool,
        typer.Option("--review", is_flag=True, help="Aplicar correções ao último preview do CP1."),
    ] = False,
    signoff: Annotated[
        bool,
        typer.Option("--signoff", is_flag=True, help="Aprovar preview e gerar arquivo final sem marcações."),
    ] = False,
    workspace: Annotated[
        Path | None,
        typer.Option("--workspace", "-w", help="Diretório do workspace."),
    ] = None,
    model: Annotated[
        str,
        typer.Option("--model", "-m", help="Claude model ID."),
    ] = "",
) -> None:
    """Gerenciar o Research Charter (Checkpoint 1).

    Use uma das flags para selecionar a ação:\n
      --begin     Iniciar uma nova pesquisa\n
      --review    Revisar o último preview (aplicar correções do Word)\n
      --signoff   Aprovar o preview atual e gerar arquivo final
    """
    flags = [begin, review, signoff]
    if sum(flags) > 1:
        console.print("[red]Use apenas uma flag por vez: --begin, --review ou --signoff.[/red]")
        raise typer.Exit(1)
    if not any(flags):
        console.print(
            "\n[bold]ai-skill research[/bold]  —  Research Charter (CP1)\n\n"
            "  [cyan]--begin[/cyan]     Iniciar uma nova pesquisa\n"
            "  [cyan]--review[/cyan]    Revisar o último preview (aplicar correções do Word)\n"
            "  [cyan]--signoff[/cyan]   Aprovar o preview atual e gerar arquivo final\n"
        )
        raise typer.Exit(0)

    if begin:
        begin_research(topic="", workspace=workspace, model=model)
    elif review:
        resume(workspace=workspace, correct=True)
    elif signoff:
        ws_path = _resolve_workspace_path(workspace, state_subdir=True)
        if ws_path is None:
            raise typer.Exit(1)
        _signoff_checkpoint(ws_path, 1)


@app.command("literature")
def literature_cmd(
    begin: Annotated[
        bool,
        typer.Option("--begin", is_flag=True, help="Iniciar Revisão Bibliográfica (CP2)."),
    ] = False,
    review: Annotated[
        bool,
        typer.Option("--review", is_flag=True, help="Aplicar correções ao último preview do CP2."),
    ] = False,
    signoff: Annotated[
        bool,
        typer.Option("--signoff", is_flag=True, help="Aprovar preview e gerar arquivo final sem marcações."),
    ] = False,
    workspace: Annotated[
        Path | None,
        typer.Option("--workspace", "-w", help="Diretório do workspace."),
    ] = None,
) -> None:
    """Gerenciar a Revisão Bibliográfica (Checkpoint 2).

    Use uma das flags para selecionar a ação:\n
      --begin     Iniciar a pesquisa bibliográfica\n
      --review    Revisar o último preview (aplicar correções do Word)\n
      --signoff   Aprovar o preview atual e gerar arquivo final
    """
    flags = [begin, review, signoff]
    if sum(flags) > 1:
        console.print("[red]Use apenas uma flag por vez: --begin, --review ou --signoff.[/red]")
        raise typer.Exit(1)
    if not any(flags):
        console.print(
            "\n[bold]ai-skill literature[/bold]  —  Revisão Bibliográfica (CP2)\n\n"
            "  [cyan]--begin[/cyan]     Iniciar a pesquisa bibliográfica\n"
            "  [cyan]--review[/cyan]    Revisar o último preview (aplicar correções do Word)\n"
            "  [cyan]--signoff[/cyan]   Aprovar o preview atual e gerar arquivo final\n"
        )
        raise typer.Exit(0)

    if begin:
        begin_literature(workspace=workspace)
    elif review:
        resume(workspace=workspace, correct=True)
    elif signoff:
        ws_path = _resolve_workspace_path(workspace, state_subdir=True)
        if ws_path is None:
            raise typer.Exit(1)
        _signoff_checkpoint(ws_path, 2)


@app.command("begin-research")
def begin_research(
    topic: Annotated[
        str,
        typer.Argument(help="Research topic or question. If omitted, you will be prompted."),
    ] = "",
    workspace: Annotated[
        Path | None,
        typer.Option("--workspace", "-w", help="Custom workspace directory."),
    ] = None,
    model: Annotated[
        str,
        typer.Option("--model", "-m", help="Claude model ID to use."),
    ] = "",
) -> None:
    """Iniciar o Research Charter (Checkpoint 1) para um novo projeto de pesquisa."""
    _configure_logging(os.environ.get("AI_SKILL_LOG_LEVEL", "WARNING"))

    if model:
        os.environ["AI_SKILL_MODEL"] = model

    # --- Interactive workspace + topic selection when arguments are omitted ---
    if workspace is None or not topic.strip():
        console.print(Panel("[bold]Novo Research Charter — CP1[/bold]", expand=False))

    if workspace is None:
        workspace = _pick_workspace_for_research()
        if workspace is None:
            raise typer.Exit(0)

    # Pre-fill topic from workspace.yaml when not provided on the command line
    if not topic.strip():
        _pre = _topic_from_workspace(workspace)
        prompt_default = _pre if _pre else ""
        while True:
            entered = typer.prompt(
                "\nTópico ou questão de pesquisa",
                default=prompt_default if prompt_default else ...,  # type: ignore[arg-type]
            ).strip()
            if entered:
                topic = entered
                break
            console.print("[red]O tópico não pode ser vazio.[/red]")
        console.print(f"\n  [dim]→[/dim] [bold]{topic}[/bold]")
        if not typer.confirm("Confirmar?", default=True):
            console.print("[dim]Operação cancelada.[/dim]")
            raise typer.Exit(0)

    ws_path = workspace

    # If the given path is a ProjectWorkspace, keep research artefacts inside
    # it but store the LangGraph internal state under <project>/.state/
    if (ws_path / "workspace.yaml").exists():
        state_path = ws_path / ".state"
        _warn_overwrite(ws_path, checkpoints=[1])
    else:
        state_path = ws_path

    console.print(f"[bold]Starting research:[/bold] {topic}")
    console.print(f"[dim]Workspace:[/dim] {ws_path}")

    # Late imports keep CLI startup fast
    from ai_skill.core.graph import build_cp1_graph
    from ai_skill.core.state import initial_state
    from ai_skill.core.workspace import ResearchWorkspace

    ws = ResearchWorkspace(state_path)
    ws.initialise(topic)

    graph = build_cp1_graph()
    state = initial_state(state_path, topic=topic)
    current_state: dict = dict(state)

    try:
        for event in graph.stream(state, stream_mode="updates"):
            for node_name, node_output in event.items():
                if node_name == "__interrupt__":
                    ws.save_state(current_state)
                    _handle_checkpoint(node_output, current_state, state_path)
                    return
                _print_node_progress(node_name, node_output, len(current_state.get("findings", [])))
                if isinstance(node_output, dict):
                    current_state.update(node_output)
                    ws.save_state(current_state)  # persist after every node
    except _GRAPH_INTERRUPT as _exc:  # type: ignore[misc]
        # LangGraph raises GraphInterrupt instead of yielding an event in some versions
        ws.save_state(current_state)
        _handle_checkpoint(_exc, current_state, state_path)
        return
    except KeyboardInterrupt:
        ws.save_state(current_state)
        console.print("\n[yellow]Pesquisa pausada. Retome com:[/yellow]")
        console.print(f"  ai-skill research --review --workspace \"{state_path}\"")
        sys.exit(0)


@app.command()
def status(
    workspace: Annotated[
        Path | None,
        typer.Option("--workspace", "-w", help="Workspace directory to inspect."),
    ] = None,
) -> None:
    """Show the current status of an existing research workspace."""
    _configure_logging()

    from ai_skill.core.workspace import ResearchWorkspace

    ws_path = _resolve_workspace_path(workspace, state_subdir=True)
    if ws_path is None:
        raise typer.Exit(1)

    ws = ResearchWorkspace(ws_path)
    state = ws.load_state()
    if state is None:
        console.print(f"[red]No research state found in:[/red] {workspace}")
        raise typer.Exit(1)

    objective = state.get("objective", {})
    console.print(f"[bold]Topic:[/bold] {objective.get('topic', '(unknown)')}")
    console.print(f"[bold]Stage:[/bold] {state.get('stage', '?')}")
    console.print(f"[bold]Status:[/bold] {state.get('status', '?')}")
    console.print(f"[bold]Attempt:[/bold] {state.get('attempt', 0)}")
    findings_count = len(state.get("findings", []))
    console.print(f"[bold]Findings:[/bold] {findings_count}")


@app.command()
def resume(
    workspace: Annotated[
        Path | None,
        typer.Option("--workspace", "-w", help="Workspace directory to resume."),
    ] = None,
    correct: Annotated[
        bool,
        typer.Option(
            "--correct", "-c",
            help="Read the last preview docx, extract corrections and generate a new preview.",
            is_flag=True,
        ),
    ] = False,
) -> None:
    """Resume a paused research pipeline.

    Without flags: approve the current checkpoint and advance.\n
    With --correct: read the last preview .docx, apply the researcher's
    edits and generate a new preview for review.
    """
    _configure_logging(os.environ.get("AI_SKILL_LOG_LEVEL", "WARNING"))

    from ai_skill.core.graph import build_cp1_graph, build_cp2_graph, build_cp3_graph
    from ai_skill.core.workspace import ResearchWorkspace

    ws_path = _resolve_workspace_path(workspace, state_subdir=True)
    if ws_path is None:
        raise typer.Exit(1)
    workspace = ws_path

    ws = ResearchWorkspace(workspace)
    state = ws.load_state()
    project_path = workspace.parent
    is_project_workspace = (project_path / "workspace.yaml").exists()

    if state is None:
        if correct and is_project_workspace:
            console.print(
                "[yellow]State file ausente — reconstruindo a partir do workspace...[/yellow]"
            )
            try:
                state = _reconstruct_state_for_correct(project_path, workspace)
            except Exception as _rec_exc:
                logger.debug("State reconstruction failed: %s", _rec_exc)
                state = None
        if state is None:
            console.print(f"[red]No state found in:[/red] {workspace}")
            console.print(
                f"[dim]Re-run:[/dim]  ai-skill research --begin --workspace \"{project_path}\""
            )
            raise typer.Exit(1)

    charter_approved: bool = bool(state.get("charter_approved"))
    active_cp: int = int(state.get("active_checkpoint", 1))

    # --- Approve path (no --correct): directly record approval, no graph re-run ---
    if not correct:
        if not charter_approved:
            # CP1 approval
            state["charter_approved"] = True
            ws.save_state(state)
            _finalize_checkpoint(workspace, 1)
            _handle_cp2_start(state, workspace)
        elif active_cp == 2 and not state.get("literature_approved"):
            # CP2 approval
            state["literature_approved"] = True
            ws.save_state(state)
            _finalize_checkpoint(workspace, 2)
            _handle_cp3_start(state, workspace)
        elif active_cp == 3 and not state.get("design_approved"):
            # CP3 approval
            state["design_approved"] = True
            ws.save_state(state)
            _finalize_checkpoint(workspace, 3)
            console.print(
                "\n[bold green]✓ Checkpoint 3 aprovado![/bold green] "
                "Research Design finalizado."
            )
        else:
            console.print(
                "[yellow]Nenhum checkpoint pendente encontrado.[/yellow]\n"
                "[dim]Use --correct para aplicar correções a um preview existente.[/dim]"
            )
        return

    # --- Correct path (--correct): extract feedback, re-run appropriate graph ---
    if not is_project_workspace:
        console.print(
            "[red]--correct requer um ProjectWorkspace. "
            "Use --workspace apontando para o diretório .state/ dentro do workspace.[/red]"
        )
        raise typer.Exit(1)

    from ai_skill.core.project_workspace import ProjectWorkspace
    pw = ProjectWorkspace.from_path(project_path)

    # Guard: if [final].docx already exists, confirm the user wants to revise it
    final_path = pw.checkpoint_final_path(active_cp)
    if final_path.exists():
        console.print(
            f"\n[bold yellow]⚠ Já existe um arquivo aprovado:[/bold yellow]\n"
            f"  {final_path.name}"
        )
        if not typer.confirm(
            "Deseja continuar e gerar um novo preview a partir do arquivo final?",
            default=True,
        ):
            console.print("[dim]Operação cancelada.[/dim]")
            raise typer.Exit(0)

    last_preview = pw.get_last_preview(active_cp)
    if last_preview is None:
        if final_path.exists():
            last_preview = final_path
        else:
            cp_names = {1: "Checkpoint 1 (Research Charter)", 2: "Checkpoint 2 (Literature Review)", 3: "Checkpoint 3 (Research Design)"}
            console.print(
                f"[red]Nenhum preview encontrado para {cp_names.get(active_cp, f'Checkpoint {active_cp}')}. "
                "Execute a pesquisa primeiro.[/red]"
            )
            raise typer.Exit(1)

    console.print(f"[dim]Analisando preview:[/dim] {last_preview.name}")
    corrections = _extract_docx_corrections(last_preview)

    if corrections is None:
        console.print(
            "\n[green]✓ Sem correções pendentes[/green] — o documento não contém "
            "marcas de revisão, comentários nem realces."
        )
        if typer.confirm("\nDeseja gerar o arquivo final?", default=True):
            # Use _signoff_checkpoint so CP2 also downloads PDFs to attachments/,
            # updates literature_approved in state, and triggers CP3 start.
            _signoff_checkpoint(workspace, active_cp)
        raise typer.Exit(0)

    state["user_feedback"] = _format_corrections_for_llm(corrections)

    # Select graph based on active checkpoint
    if active_cp == 1:
        graph = build_cp1_graph()
    elif active_cp == 3:
        graph = build_cp3_graph()
    else:
        graph = build_cp2_graph()
    console.print(f"[bold]Aplicando correções — CP{active_cp}:[/bold] {state.get('objective', {}).get('topic', '?')}")
    current_state: dict = dict(state)

    try:
        for event in graph.stream(state, stream_mode="updates"):
            for node_name, node_output in event.items():
                if node_name == "__interrupt__":
                    ws.save_state(current_state)
                    _handle_checkpoint(node_output, current_state, workspace)
                    return
                _print_node_progress(node_name, node_output, len(current_state.get("findings", [])))
                if isinstance(node_output, dict):
                    current_state.update(node_output)
                    ws.save_state(current_state)
    except _GRAPH_INTERRUPT as _exc:  # type: ignore[misc]
        ws.save_state(current_state)
        _handle_checkpoint(_exc, current_state, workspace)
        return
    except KeyboardInterrupt:
        ws.save_state(current_state)
        console.print("\n[yellow]Pesquisa pausada. Retome com:[/yellow]")
        console.print(f"  ai-skill research --review --workspace \"{workspace}\"")
        sys.exit(0)


@app.command("begin-literature")
def begin_literature(
    workspace: Annotated[
        Path | None,
        typer.Option(
            "--workspace", "-w",
            help="ProjectWorkspace ou diretório .state/ do projeto.",
        ),
    ] = None,
) -> None:
    """Iniciar a Revisão Bibliográfica (Checkpoint 2).

    Requer que o Checkpoint 1 (Research Charter) já esteja aprovado,
    ou seja, que o arquivo [final].docx do CP1 exista no workspace.
    """
    _configure_logging(os.environ.get("AI_SKILL_LOG_LEVEL", "WARNING"))

    # --- Resolve project path (interactive picker when --workspace omitted) ---
    resolved = _resolve_workspace_path(workspace, state_subdir=False)
    if resolved is None:
        raise typer.Exit(1)

    # Accept both project root and .state/ as input
    if (resolved / "workspace.yaml").exists():
        project_path = resolved
        state_path = resolved / ".state"
    elif (resolved.parent / "workspace.yaml").exists():
        project_path = resolved.parent
        state_path = resolved
    else:
        console.print(
            "[red]Workspace não reconhecido.[/red]\n"
            "[dim]Aponte --workspace para o ProjectWorkspace "
            "ou para o diretório .state/ dentro dele.[/dim]"
        )
        raise typer.Exit(1)

    # --- Validate CP1 [final].docx exists ---
    from ai_skill.core.project_workspace import ProjectWorkspace
    pw = ProjectWorkspace.from_path(project_path)
    final_cp1 = pw.checkpoint_final_path(1)

    if not final_cp1.exists():
        console.print(
            f"\n[red]✗ Checkpoint 1 não aprovado.[/red]\n"
            f"  O arquivo [bold]{final_cp1.name}[/bold] não foi encontrado em:\n"
            f"  {project_path}\n"
            f"\n[dim]Complete o Checkpoint 1 antes de iniciar a Revisão Bibliográfica:[/dim]\n"
            f"  ai-skill research --begin --workspace \"{project_path}\"\n"
            f"  ai-skill research --signoff --workspace \"{state_path}\""
        )
        raise typer.Exit(1)

    console.print(f"[green]✓[/green] Checkpoint 1 aprovado: [dim]{final_cp1.name}[/dim]")

    # --- Warn if CP2 work already exists ---
    _warn_overwrite(project_path, checkpoints=[2])

    # --- Load or reconstruct state ---
    from ai_skill.core.graph import build_cp2_graph
    from ai_skill.core.workspace import ResearchWorkspace

    ws = ResearchWorkspace(state_path)
    state = ws.load_state()

    if state is None:
        state = _reconstruct_state_for_literatura(project_path, state_path)
        if state is None:
            console.print(
                "[red]Não foi possível carregar o estado do projeto.[/red]\n"
                "[dim]Tente re-rodar:[/dim]\n"
                f"  ai-skill research --begin --workspace \"{project_path}\""
            )
            raise typer.Exit(1)

    # Guarantee charter is approved so the graph passes through the gate
    state["charter_approved"] = True
    state["user_feedback"] = None
    state["attempt"] = 0

    # Inject CP1 final.docx text so compile_literature can build on it
    state["charter_document_text"] = _read_docx_text(final_cp1)

    # Set the correct stage so CP2 node routing works (CP1 left it as RESEARCH_CHARTER)
    from ai_skill.core.pipeline_stages import PipelineStage as _PS
    state["stage"] = _PS.LITERATURE_REVIEW

    # Build the CP2 handoff — the ONLY view of the research objective that CP2
    # nodes receive.  success_metrics and all other CP1-only fields are excluded
    # so the planner and evaluator cannot design steps toward global deliverables
    # (framework development, article submission, etc.) during bibliographic search.
    _full_obj = state.get("objective") or {}
    state["cp2_context"] = {
        "topic":             _full_obj.get("topic") or "",
        "goals":             list(_full_obj.get("goals") or []),
        "scope_constraints": list(_full_obj.get("scope_constraints") or []),
    }

    topic = (state.get("objective") or {}).get("topic", "")
    console.print(
        f"\n[bold]Iniciando Revisão Bibliográfica[/bold]"
        + (f": {topic}" if topic else "")
    )

    # --- Stream the CP2 graph ---
    graph_instance = build_cp2_graph()
    current_state: dict = dict(state)

    try:
        for event in graph_instance.stream(state, stream_mode="updates"):
            for node_name, node_output in event.items():
                if node_name == "__interrupt__":
                    ws.save_state(current_state)
                    _handle_checkpoint(node_output, current_state, state_path)
                    return
                _print_node_progress(node_name, node_output, len(current_state.get("findings", [])))
                if isinstance(node_output, dict):
                    current_state.update(node_output)
                    ws.save_state(current_state)
                    if (
                        node_name == "review_literature"
                        and node_output.get("literature_approved")
                    ):
                        _finalize_checkpoint(state_path, 2)
    except _GRAPH_INTERRUPT as _exc:  # type: ignore[misc]
        ws.save_state(current_state)
        _handle_checkpoint(_exc, current_state, state_path)
        return
    except KeyboardInterrupt:
        ws.save_state(current_state)
        console.print("\n[yellow]Pesquisa pausada. Retome com:[/yellow]")
        console.print(f"  ai-skill literature --review --workspace \"{state_path}\"")
        sys.exit(0)


@app.command()
def graph() -> None:
    """Print the Mermaid diagram of the research pipeline graph."""
    from ai_skill.core.graph import get_graph_mermaid

    mermaid = get_graph_mermaid()
    console.print(mermaid)


# ---------------------------------------------------------------------------
# Skills sub-commands
# ---------------------------------------------------------------------------


@skills_app.command("list")
def skills_list() -> None:
    """List all registered skills."""
    from ai_skill.skills.registry import SkillRegistry

    registry = SkillRegistry()
    registry.auto_discover()

    table = Table(title="Registered Skills")
    table.add_column("Name", style="cyan")
    table.add_column("Version")
    table.add_column("Tags")
    table.add_column("Description")

    for meta in registry.all():
        table.add_row(
            meta["name"],
            meta["version"],
            ", ".join(meta.get("tags", [])),
            meta["description"][:60] + ("…" if len(meta["description"]) > 60 else ""),
        )

    console.print(table)


@skills_app.command("show")
def skills_show(
    name: Annotated[str, typer.Argument(help="Skill name to inspect.")],
) -> None:
    """Show details for a specific skill."""
    from ai_skill.skills.registry import SkillRegistry

    registry = SkillRegistry()
    registry.auto_discover()

    skill = registry.get(name)
    if skill is None:
        console.print(f"[red]Skill not found:[/red] {name}")
        raise typer.Exit(1)

    meta = skill.SKILL_META
    console.print(f"[bold]{meta.name}[/bold] v{meta.version}")
    console.print(f"[dim]{meta.description}[/dim]")
    console.print(f"Author: {meta.author}  License: {meta.license}")
    console.print(f"Tags: {', '.join(meta.tags)}")
    if meta.dependencies:
        console.print(f"Depends on: {', '.join(meta.dependencies)}")


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _resolve_workspace_path(given: Path | None, *, state_subdir: bool) -> Path | None:
    """Return a workspace path, falling back to an interactive picker when *given* is None.

    Args:
        given: The path supplied by the user via ``--workspace``, or None.
        state_subdir: When True, return ``project_path / ".state"`` (for commands
            that hand the path to :class:`ResearchWorkspace`).  When False,
            return the project root (for commands that detect the layout themselves).

    Returns:
        The resolved path, or None if the user cancelled or no workspaces exist.
    """
    if given is not None:
        return given

    from ai_skill.core.project_workspace import default_projects_root, list_workspaces

    root = default_projects_root()
    workspaces = list_workspaces(root)

    if not workspaces:
        console.print(f"[red]Nenhum workspace encontrado em:[/red] {root}")
        console.print("[dim]Crie um com:[/dim]  ai-skill workspace begin")
        return None

    console.print("\n[bold]Workspaces disponíveis:[/bold]\n")
    for i, ws in enumerate(workspaces, start=1):
        meta = ws.load_metadata()
        name = meta.get("name", ws.slug)
        topic = meta.get("topic", "?")
        cp = meta.get("current_checkpoint", 0)
        st = meta.get("status", "?")
        topic_short = topic[:52] + "…" if len(topic) > 52 else topic
        console.print(
            f"  [cyan]{i:>2}.[/cyan] [bold]{name}[/bold]  "
            f"[dim]CP{cp} · {st}[/dim]\n"
            f"       [dim]{topic_short}[/dim]\n"
        )

    raw = typer.prompt("Número do workspace").strip()
    try:
        idx = int(raw) - 1
        if not (0 <= idx < len(workspaces)):
            console.print("[red]Número inválido.[/red]")
            return None
    except ValueError:
        console.print("[red]Entrada inválida.[/red]")
        return None

    selected = workspaces[idx]
    console.print(f"\n  [dim]→[/dim] [bold]{selected.path}[/bold]")
    if not typer.confirm("Confirmar?", default=True):
        console.print("[dim]Operação cancelada.[/dim]")
        return None

    return selected.path / ".state" if state_subdir else selected.path


def _pick_workspace_for_research() -> Path | None:
    """List existing ProjectWorkspaces (plus a 'create new' option) for begin-research.

    Returns the selected project root path, or None if the user cancels.
    """
    from ai_skill.core.project_workspace import ProjectWorkspace, default_projects_root, list_workspaces, slugify

    root = default_projects_root()
    workspaces = list_workspaces(root)

    console.print("\n[bold]Selecione um workspace:[/bold]\n")
    console.print("   [cyan] 0.[/cyan] [green]+ Criar novo workspace[/green]\n")
    for i, ws in enumerate(workspaces, start=1):
        meta = ws.load_metadata()
        name = meta.get("name", ws.slug)
        topic = meta.get("topic", "?")
        cp = meta.get("current_checkpoint", 0)
        st = meta.get("status", "?")
        topic_short = topic[:52] + "…" if len(topic) > 52 else topic
        console.print(
            f"  [cyan]{i:>2}.[/cyan] [bold]{name}[/bold]  "
            f"[dim]CP{cp} · {st}[/dim]\n"
            f"       [dim]{topic_short}[/dim]\n"
        )

    raw = typer.prompt("Número").strip()
    try:
        idx = int(raw)
    except ValueError:
        console.print("[red]Entrada inválida.[/red]")
        return None

    if idx == 0:
        # Create new workspace
        console.print()
        while True:
            raw_name = typer.prompt("Nome da pesquisa").strip()
            if not raw_name:
                console.print("[red]Nome não pode ser vazio.[/red]")
                continue
            try:
                slugify(raw_name)
            except ValueError as exc:
                console.print(f"[red]{exc}[/red]")
                continue
            break
        ws_new = ProjectWorkspace(raw_name, root=root)
        if not ws_new.exists():
            ws_new.create(topic="")
            console.print(f"  [green]✓[/green] Workspace criado: [bold]{ws_new.path}[/bold]")
        console.print(f"\n  [dim]→[/dim] [bold]{ws_new.path}[/bold]")
        if not typer.confirm("Confirmar?", default=True):
            console.print("[dim]Operação cancelada.[/dim]")
            return None
        return ws_new.path

    if not (1 <= idx <= len(workspaces)):
        console.print("[red]Número inválido.[/red]")
        return None

    selected = workspaces[idx - 1]
    console.print(f"\n  [dim]→[/dim] [bold]{selected.path}[/bold]")
    if not typer.confirm("Confirmar?", default=True):
        console.print("[dim]Operação cancelada.[/dim]")
        return None
    return selected.path


def _topic_from_workspace(ws_path: Path) -> str:
    """Return the topic stored in workspace.yaml, or empty string if unavailable.

    Args:
        ws_path: Root of the ProjectWorkspace.

    Returns:
        Topic string, or empty string.
    """
    meta_file = ws_path / "workspace.yaml"
    if not meta_file.exists():
        return ""
    try:
        import yaml as _yaml
        meta = _yaml.safe_load(meta_file.read_text(encoding="utf-8")) or {}
        return str(meta.get("topic", "")).strip()
    except Exception:
        return ""


def _resolve_workspace(name: str | None) -> "ProjectWorkspace | None":  # noqa: F821
    """Resolve a ProjectWorkspace by name or interactive selection.

    Args:
        name: Workspace name/slug, or None for interactive selection.

    Returns:
        The resolved ProjectWorkspace, or None if not found / cancelled.
    """
    from ai_skill.core.project_workspace import ProjectWorkspace, default_projects_root, list_workspaces

    root = default_projects_root()

    if name:
        ws = ProjectWorkspace(name, root=root)
        if not ws.exists():
            console.print(f"[red]Workspace não encontrado:[/red] {name}")
            console.print(f"[dim]Workspaces disponíveis em:[/dim] {root}")
            return None
        return ws

    workspaces = list_workspaces(root)
    if not workspaces:
        console.print(f"[red]Nenhum workspace encontrado em:[/red] {root}")
        console.print("[dim]Crie um com:[/dim]  ai-skill workspace begin")
        return None

    if len(workspaces) == 1:
        selected = workspaces[0]
        meta = selected.load_metadata()
        console.print(f"\n  [dim]→[/dim] [bold]{meta.get('name', selected.slug)}[/bold]  {selected.path}")
        if not typer.confirm("Confirmar?", default=True):
            console.print("[dim]Operação cancelada.[/dim]")
            return None
        return selected

    console.print("\n[bold]Workspaces disponíveis:[/bold]\n")
    for i, ws in enumerate(workspaces, start=1):
        meta = ws.load_metadata()
        name = meta.get("name", ws.slug)
        topic = meta.get("topic", "?")
        cp = meta.get("current_checkpoint", 0)
        st = meta.get("status", "?")
        topic_short = topic[:52] + "…" if len(topic) > 52 else topic
        console.print(
            f"  [cyan]{i:>2}.[/cyan] [bold]{name}[/bold]  "
            f"[dim]CP{cp} · {st}[/dim]\n"
            f"       [dim]{topic_short}[/dim]\n"
        )

    raw = typer.prompt("Número do workspace").strip()
    try:
        idx = int(raw) - 1
        if not (0 <= idx < len(workspaces)):
            console.print("[red]Número inválido.[/red]")
            return None
    except ValueError:
        console.print("[red]Entrada inválida.[/red]")
        return None

    selected = workspaces[idx]
    console.print(f"\n  [dim]→[/dim] [bold]{selected.path}[/bold]")
    if not typer.confirm("Confirmar?", default=True):
        console.print("[dim]Operação cancelada.[/dim]")
        return None
    return selected


def prompt_file_selection(workspace_path: Path) -> Path | None:
    """Present a numbered list of workspace files and return the user's choice.

    Lists checkpoint ``.docx`` files first, then files in ``attachments/``.
    Returns None if the workspace is empty or the user cancels.

    Args:
        workspace_path: Root directory of the ProjectWorkspace.

    Returns:
        The selected Path, or None.
    """
    from ai_skill.core.project_workspace import ProjectWorkspace

    ws = ProjectWorkspace.from_path(workspace_path)
    files = ws.list_all_files()
    if not files:
        console.print("[dim]Nenhum arquivo disponível no workspace.[/dim]")
        return None

    console.print("\n[bold]Arquivos disponíveis:[/bold]")
    for i, f in enumerate(files, start=1):
        rel = f.relative_to(workspace_path)
        console.print(f"  [cyan]{i:>2}.[/cyan] {rel}")

    raw = typer.prompt("\nNúmero do arquivo (Enter para cancelar)", default="")
    if not raw.strip():
        return None
    try:
        idx = int(raw) - 1
        if 0 <= idx < len(files):
            return files[idx]
        console.print("[red]Número inválido.[/red]")
        return None
    except ValueError:
        console.print("[red]Entrada inválida.[/red]")
        return None


def _warn_overwrite(project_path: Path, checkpoints: list[int]) -> None:
    """Warn the user if checkpoint work already exists and ask to confirm.

    Previews auto-increment so they are never overwritten, but if a
    [final].docx already exists the user is re-starting an approved stage.

    Args:
        project_path: Root of the ProjectWorkspace.
        checkpoints: Checkpoint numbers to inspect.
    """
    from ai_skill.core.project_workspace import ProjectWorkspace

    pw = ProjectWorkspace.from_path(project_path)
    finals, previews = [], []
    for n in checkpoints:
        final = pw.checkpoint_final_path(n)
        if final.exists():
            finals.append(final)
        else:
            last_preview = pw.get_last_preview(n)
            if last_preview is not None:
                previews.append(last_preview)

    if finals:
        console.print("\n[bold yellow]⚠ Um checkpoint já foi aprovado:[/bold yellow]")
        for f in finals:
            console.print(f"  [yellow]•[/yellow] {f.name}")
        console.print(
            "[dim]Iniciar uma nova pesquisa sobrescreverá o preview existente "
            "(o arquivo [final] não será alterado).[/dim]"
        )
        if not typer.confirm("\nDeseja continuar mesmo assim?", default=True):
            console.print("[dim]Operação cancelada.[/dim]")
            raise typer.Exit(0)
        for n in checkpoints:
            pw.reset_previews(n)
    elif previews:
        console.print("\n[bold yellow]⚠ Já existem previews neste workspace:[/bold yellow]")
        for f in previews:
            console.print(f"  [yellow]•[/yellow] {f.name}")
        console.print(
            "[dim]Iniciar uma nova pesquisa sobrescreverá os previews existentes.[/dim]"
        )
        if not typer.confirm("\nDeseja continuar mesmo assim?", default=True):
            console.print("[dim]Operação cancelada.[/dim]")
            raise typer.Exit(0)
        for n in checkpoints:
            pw.reset_previews(n)


def _read_docx_text(path: Path) -> str:
    """Extract plain text from a .docx file paragraph by paragraph.

    Args:
        path: Path to the .docx file.

    Returns:
        Plain text content of the document, or an error message.
    """
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:
        return f"(Erro ao ler o documento: {exc})"


def _finalize_checkpoint(state_path: Path, number: int) -> None:
    """Copy the latest preview to [final].docx when a checkpoint is approved.

    If a [final].docx already exists the user is asked to confirm overwrite.
    Skips silently when state_path is not inside a ProjectWorkspace.

    Args:
        state_path: Path to the ResearchWorkspace (.state/) directory.
        number: Checkpoint number to finalize.
    """
    project_path = state_path.parent
    if not (project_path / "workspace.yaml").exists():
        return
    from ai_skill.core.project_workspace import ProjectWorkspace
    pw = ProjectWorkspace.from_path(project_path)
    final_path = pw.checkpoint_final_path(number)

    if final_path.exists():
        console.print(
            f"\n[bold yellow]⚠ Já existe um arquivo final:[/bold yellow]\n"
            f"  {final_path.name}"
        )
        if not typer.confirm("Deseja sobrescrever?", default=True):
            console.print("[dim]Arquivo final mantido sem alteração.[/dim]")
            return

    final = pw.finalize_checkpoint(number)
    if final:
        console.print(f"  [green]✓[/green] Charter aprovado → [bold]{final.name}[/bold]")


def _extract_docx_corrections(path: Path) -> dict | None:
    """Extract actionable corrections from an edited .docx file.

    Scans for three types of user markup:
    - **Track changes** (``w:ins`` / ``w:del``): text the user added or removed
    - **Comments** (``word/comments.xml``): instructions written as Word comments
    - **Yellow highlights**: paragraphs/runs the user wants fully regenerated

    Args:
        path: Path to the edited .docx file.

    Returns:
        Dict with keys ``insertions``, ``deletions``, ``comments``, ``highlights``,
        or ``None`` if the document contains none of these markers.
    """
    try:
        import zipfile
        from docx import Document as _DocxDoc
        from docx.oxml.ns import qn
    except ImportError:
        return None

    try:
        doc = _DocxDoc(str(path))
    except Exception:
        return None

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    insertions: list[dict] = []
    deletions: list[dict] = []
    highlights: list[dict] = []

    for para in doc.paragraphs:
        para_text = para.text.strip()

        # Yellow highlights on individual runs
        for run in para.runs:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is not None:
                h = rpr.find(qn("w:highlight"))
                if h is not None and h.get(qn("w:val")) in ("yellow", "darkYellow"):
                    if run.text.strip():
                        highlights.append({
                            "highlighted_text": run.text.strip(),
                            "paragraph_context": para_text[:120],
                        })

        # Also flag entire paragraph if ALL text is highlighted yellow
        all_runs = para.runs
        if all_runs and all(
            r._element.find(qn("w:rPr")) is not None
            and r._element.find(qn("w:rPr")).find(qn("w:highlight")) is not None  # type: ignore[union-attr]
            and r._element.find(qn("w:rPr")).find(qn("w:highlight")).get(qn("w:val")) in ("yellow", "darkYellow")  # type: ignore[union-attr]
            for r in all_runs if r.text.strip()
        ) and para_text:
            # Deduplicate: if paragraph already covered by run-level highlights, skip
            if not any(h["paragraph_context"] == para_text[:120] for h in highlights):
                highlights.append({
                    "highlighted_text": para_text,
                    "paragraph_context": para_text[:120],
                })

        # Track changes: insertions (w:ins)
        for ins_elem in para._element.findall(".//" + qn("w:ins")):
            text = "".join(t.text or "" for t in ins_elem.findall(".//" + qn("w:t")))
            if text.strip():
                insertions.append({"text": text.strip(), "context": para_text[:120]})

        # Track changes: deletions (w:del)
        for del_elem in para._element.findall(".//" + qn("w:del")):
            text = "".join(
                t.text or "" for t in del_elem.findall(".//" + qn("w:delText"))
            )
            if text.strip():
                deletions.append({"text": text.strip(), "context": para_text[:120]})

    # Comments from word/comments.xml — also capture annotated anchor text
    # so the LLM knows which passage each comment targets.
    comments: list[dict] = []
    try:
        with zipfile.ZipFile(str(path)) as z:
            if "word/comments.xml" in z.namelist():
                from lxml import etree  # type: ignore[import-untyped]
                root = etree.fromstring(z.read("word/comments.xml"))
                for comment_elem in root.findall(f"{{{W}}}comment"):
                    author = comment_elem.get(f"{{{W}}}author", "Usuário")
                    comment_id = comment_elem.get(f"{{{W}}}id", "")
                    t_elems = comment_elem.findall(f".//{{{W}}}t")
                    text = "".join(t.text or "" for t in t_elems).strip()
                    if not text:
                        continue
                    # Find the anchor text in document.xml tied to this comment id
                    anchor = ""
                    if comment_id and "word/document.xml" in z.namelist():
                        try:
                            doc_root = etree.fromstring(z.read("word/document.xml"))
                            # Find commentRangeStart with matching id, then collect
                            # text until the paired commentRangeEnd
                            start_tag = f"{{{W}}}commentRangeStart"
                            end_tag = f"{{{W}}}commentRangeEnd"
                            collecting = False
                            anchor_parts: list[str] = []
                            for elem in doc_root.iter():
                                if elem.tag == start_tag and elem.get(f"{{{W}}}id") == comment_id:
                                    collecting = True
                                    continue
                                if elem.tag == end_tag and elem.get(f"{{{W}}}id") == comment_id:
                                    break
                                if collecting and elem.tag == f"{{{W}}}t" and elem.text:
                                    anchor_parts.append(elem.text)
                            anchor = "".join(anchor_parts).strip()[:200]
                        except Exception:
                            pass
                    entry: dict = {"author": author, "text": text}
                    if anchor:
                        entry["anchor"] = anchor
                    comments.append(entry)
    except Exception:
        pass

    if not insertions and not deletions and not comments and not highlights:
        return None

    return {
        "insertions": insertions,
        "deletions": deletions,
        "comments": comments,
        "highlights": highlights,
    }


def _format_corrections_for_llm(corrections: dict) -> str:
    """Format extracted docx corrections into a structured prompt string.

    Args:
        corrections: Output of :func:`_extract_docx_corrections`.

    Returns:
        Formatted string describing all corrections for the LLM.
    """
    parts: list[str] = []

    if corrections.get("comments"):
        parts.append("## Comentários do pesquisador (instruções a implementar):")
        for c in corrections["comments"]:
            anchor_info = f'  (sobre: "{c["anchor"]}")' if c.get("anchor") else ""
            parts.append(f'- [{c["author"]}]: {c["text"]}{anchor_info}')

    if corrections.get("insertions"):
        parts.append("\n## Trechos inseridos pelo pesquisador (incorporar):")
        for i in corrections["insertions"]:
            parts.append(f'- Inseriu: "{i["text"]}"  (contexto: {i["context"]})')

    if corrections.get("deletions"):
        parts.append("\n## Trechos removidos pelo pesquisador (excluir):")
        for d in corrections["deletions"]:
            parts.append(f'- Removeu: "{d["text"]}"  (contexto: {d["context"]})')

    if corrections.get("highlights"):
        parts.append(
            "\n## Trechos em destaque amarelo (regerar completamente — "
            "manter o tema mas reescrever o conteúdo):"
        )
        for h in corrections["highlights"]:
            parts.append(f'- "{h["highlighted_text"]}"')

    return "\n".join(parts)


def _reconstruct_state_for_literatura(project_path: Path, state_path: Path) -> dict | None:
    """Build a minimal ResearchState for the ``literatura`` command.

    Used when the state file is missing or corrupt but CP1 [final].docx exists.
    Reads the topic from workspace.yaml and restores the objective from the
    CP1 final document.

    Args:
        project_path: Root of the ProjectWorkspace.
        state_path: Path to the .state/ directory.

    Returns:
        A state dict with ``charter_approved=True``, or None on failure.
    """
    import yaml as _yaml
    from ai_skill.core.project_workspace import ProjectWorkspace
    from ai_skill.core.state import initial_state

    meta_file = project_path / "workspace.yaml"
    if not meta_file.exists():
        return None
    try:
        meta = _yaml.safe_load(meta_file.read_text(encoding="utf-8")) or {}
    except Exception:
        return None

    topic = meta.get("topic", "")
    state: dict = dict(initial_state(str(state_path), topic=topic))

    pw = ProjectWorkspace.from_path(project_path)
    final_cp1 = pw.checkpoint_final_path(1)
    if final_cp1.exists():
        objective = _parse_objective_from_docx(final_cp1)
        if objective:
            state["objective"] = objective
        state["charter_document_text"] = _read_docx_text(final_cp1)

    state["charter_approved"] = True
    return state


def _reconstruct_state_for_correct(project_path: Path, state_path: Path) -> dict | None:
    """Build a minimal ResearchState when research-state.yaml is absent.

    Reads the workspace metadata for the topic, and tries to restore the
    objective fields by parsing the last preview .docx.  This allows
    ``resume --correct`` to work even when the state file was not saved.

    Args:
        project_path: Root of the ProjectWorkspace.
        state_path: Path to the .state/ directory (ResearchWorkspace).

    Returns:
        A minimal state dict, or None if reconstruction is not possible.
    """
    import yaml as _yaml
    from ai_skill.core.project_workspace import ProjectWorkspace
    from ai_skill.core.state import initial_state

    meta_file = project_path / "workspace.yaml"
    if not meta_file.exists():
        return None
    try:
        meta = _yaml.safe_load(meta_file.read_text(encoding="utf-8")) or {}
    except Exception:
        return None

    topic = meta.get("topic", "")
    state: dict = dict(initial_state(str(state_path), topic=topic))

    # Restore objective from the last preview docx
    pw = ProjectWorkspace.from_path(project_path)
    last_preview = pw.get_last_preview(1)
    if last_preview is not None:
        objective = _parse_objective_from_docx(last_preview)
        if objective:
            state["objective"] = objective

    return state


def _parse_objective_from_docx(path: Path) -> dict | None:
    """Restore ResearchObjective fields from a charter .docx file.

    Parses the structured headings generated by :func:`_charter_to_docx`
    to reconstruct the objective dict without calling the LLM.

    Args:
        path: Path to a charter ``[preview_N].docx`` file.

    Returns:
        A partial ResearchObjective dict, or None if parsing fails.
    """
    try:
        from docx import Document as _DocxDoc
        doc = _DocxDoc(str(path))
    except Exception:
        return None

    obj: dict = {}
    current_section: str | None = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = getattr(para.style, "name", "") or ""
        if "Heading" in style or "Title" in style:
            lower = text.lower()
            if "tópico" in lower or "topic" in lower:
                current_section = "topic"
            elif "objetivo" in lower or "goal" in lower:
                current_section = "goals"
            elif "métrica" in lower or "metric" in lower or "sucesso" in lower:
                current_section = "success_metrics"
            elif "restrição" in lower or "escopo" in lower or "constraint" in lower:
                current_section = "scope_constraints"
            elif "metodolog" in lower:
                current_section = "methodology_preference"
            else:
                current_section = None
        else:
            if current_section == "topic":
                obj["topic"] = text
            elif current_section in ("goals", "success_metrics", "scope_constraints"):
                item = text.lstrip("•-– ").strip()
                if item:
                    obj.setdefault(current_section, []).append(item)
            elif current_section == "methodology_preference":
                obj["methodology_preference"] = text

    return obj if obj.get("goals") else None


def _strip_docx_markup(source: Path, dest: Path) -> bool:
    """Copy source to dest with track changes accepted and all markup removed.

    Accepts all insertions (keeps text), removes all deletions, clears
    highlight formatting and comment references.

    Args:
        source: Path to the source .docx file.
        dest: Destination path for the cleaned .docx.

    Returns:
        True if markup was stripped successfully, False if a plain copy was used.
    """
    try:
        import zipfile
        from lxml import etree as _etree  # type: ignore[import-untyped]

        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        with zipfile.ZipFile(source, "r") as z:
            contents: dict[str, bytes] = {n: z.read(n) for n in z.namelist()}

        if "word/document.xml" in contents:
            root = _etree.fromstring(contents["word/document.xml"])

            # Accept insertions: unwrap w:ins (keep child elements)
            for ins in list(root.iter(f"{{{W}}}ins")):
                parent = ins.getparent()
                if parent is None:
                    continue
                idx = list(parent).index(ins)
                for child in list(ins):
                    ins.remove(child)
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(ins)

            # Remove deletions entirely
            for del_elem in list(root.iter(f"{{{W}}}del")):
                parent = del_elem.getparent()
                if parent is not None:
                    parent.remove(del_elem)

            # Remove highlight formatting
            for hl in list(root.iter(f"{{{W}}}highlight")):
                parent = hl.getparent()
                if parent is not None:
                    parent.remove(hl)

            # Remove comment anchor elements
            for tag in (
                f"{{{W}}}commentReference",
                f"{{{W}}}commentRangeStart",
                f"{{{W}}}commentRangeEnd",
            ):
                for elem in list(root.iter(tag)):
                    parent = elem.getparent()
                    if parent is not None:
                        parent.remove(elem)

            contents["word/document.xml"] = _etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )

        # Remove comment files
        contents = {k: v for k, v in contents.items() if "comments" not in k.lower()}

        with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as z:
            for name, data in contents.items():
                z.writestr(name, data)

        return True
    except Exception:
        import shutil as _shutil
        _shutil.copy2(source, dest)
        return False


def _slugify_word(word: str) -> str:
    """Return a lowercase ASCII-safe slug for one word (used in filenames)."""
    import re
    import unicodedata
    norm = unicodedata.normalize("NFKD", word).encode("ascii", "ignore").decode()
    return re.sub(r"[^a-z0-9]", "", norm.lower())


def _make_ref_filename(ref_num: int, ref: dict) -> str:
    """Build a short filename: ``NN_firstauthor_kw1_kw2_kw3`` (no extension)."""
    authors: list = ref.get("authors") or []
    title: str = ref.get("title") or ""

    first_author = _slugify_word(
        (authors[0].split(",")[0].split()[-1] if authors else "unknown")
    ) or "unknown"

    stop = {"the", "a", "an", "of", "in", "on", "at", "for", "and", "with",
            "to", "de", "da", "do", "em", "um", "uma", "para", "com", "por"}
    keywords = [
        _slugify_word(w)
        for w in title.split()
        if len(w) > 3 and _slugify_word(w) not in stop
    ][:3]

    base = f"{ref_num:02d}_{first_author}"
    if keywords:
        base += "_" + "_".join(keywords)
    return base


def _text_to_pdf_bytes(title: str, content: str) -> bytes:
    """Create a minimal PDF from plain/markdown text using PyMuPDF.

    Returns empty bytes if PyMuPDF is unavailable or the conversion fails.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return b""
    try:
        doc = fitz.open()
        page = doc.new_page()
        rect = fitz.Rect(50, 60, 545, 780)
        text_block = f"{title}\n{'─' * 60}\n\n{content}"
        page.insert_textbox(rect, text_block, fontsize=10, fontname="helv")
        return doc.tobytes()
    except Exception as exc:
        logger.debug("_text_to_pdf_bytes failed: %s", exc)
        return b""


def _download_references_to_attachments(
    review_doc: dict,
    pw: "ProjectWorkspace",
) -> list[int]:
    """Download all referenced sources to the attachments folder.

    Strategy per reference:
    1. Try Semantic Scholar openAccessPdf (API).
    2. Try direct HTTP PDF download.
    3. Fall back to saving the abstract/summary as a plain PDF.

    Files are named ``NN_firstauthor_kw1_kw2_kw3.pdf``.

    Args:
        review_doc: The LiteratureReviewDoc dict with ``references``.
        pw: ProjectWorkspace whose ``attachments_path`` receives the files.

    Returns:
        List of reference numbers that could NOT be saved.
    """
    import os
    import urllib.error
    import urllib.request

    from ai_skill.core.nodes import (  # reuse helpers defined in nodes.py
        _fetch_via_semantic_scholar_api,
        _extract_text_from_pdf_bytes,
    )

    references: list[dict] = review_doc.get("references", [])
    attachments_dir = pw.attachments_path
    attachments_dir.mkdir(parents=True, exist_ok=True)

    api_key = os.environ.get("SEMANTIC_SCHOLAR_API_KEY", "")
    failed: list[int] = []

    _UA = (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )

    for ref in references:
        ref_num: int = ref.get("reference_number", 0)
        url: str = ref.get("url", "")
        summary: str = ref.get("summary") or ref.get("title") or ""
        base_name = _make_ref_filename(ref_num, ref)
        out_path = attachments_dir / f"{base_name}.pdf"

        if out_path.exists():
            console.print(f"  [dim][{ref_num:02d}] já existe — pulando[/dim]")
            continue

        pdf_bytes: bytes = b""

        # --- Strategy 1: Semantic Scholar openAccessPdf ---
        try:
            from ai_skill.core.nodes import _S2_PAPER_API, _S2_PAPER_FIELDS  # type: ignore[attr-defined]
            import re as _re, json as _json
            paper_id: str | None = None
            m = _re.search(r"semanticscholar\.org/paper/([A-Za-z0-9]+)", url)
            if m:
                paper_id = m.group(1)
            if paper_id is None:
                m2 = _re.search(r"(?:doi\.org/|/doi/)([^/\s]+/[^/\s]+)", url)
                if m2:
                    paper_id = f"DOI:{m2.group(1)}"
            if paper_id is None:
                m3 = _re.search(r"arxiv\.org/(?:abs|pdf)/([^\s?/]+)", url)
                if m3:
                    paper_id = f"arXiv:{m3.group(1).split('v')[0]}"

            if paper_id:
                api_url = f"{_S2_PAPER_API}/{paper_id}?fields={_S2_PAPER_FIELDS}"
                hdrs: dict[str, str] = {"Accept": "application/json"}
                if api_key:
                    hdrs["x-api-key"] = api_key
                req = urllib.request.Request(api_url, headers=hdrs)
                with urllib.request.urlopen(req, timeout=15) as resp:
                    data = _json.loads(resp.read().decode())
                oa_pdf = (data.get("openAccessPdf") or {}).get("url", "")
                if oa_pdf:
                    pdf_req = urllib.request.Request(oa_pdf, headers={"User-Agent": _UA})
                    with urllib.request.urlopen(pdf_req, timeout=30) as pdf_resp:
                        pdf_bytes = pdf_resp.read(10 * 1024 * 1024)
        except Exception as exc:
            logger.debug("[%d] S2 openAccessPdf failed: %s", ref_num, exc)

        # --- Strategy 2: Direct PDF from URL ---
        if not pdf_bytes:
            try:
                is_pdf = url.lower().split("?")[0].endswith(".pdf")
                if is_pdf:
                    pdf_req = urllib.request.Request(url, headers={"User-Agent": _UA})
                    with urllib.request.urlopen(pdf_req, timeout=30) as pdf_resp:
                        ct = pdf_resp.headers.get("Content-Type", "")
                        raw = pdf_resp.read(10 * 1024 * 1024)
                        if "pdf" in ct or is_pdf:
                            pdf_bytes = raw
            except Exception as exc:
                logger.debug("[%d] Direct PDF download failed: %s", ref_num, exc)

        # --- Strategy 3: Save abstract/summary as PDF ---
        if not pdf_bytes:
            content = (
                f"Referência [{ref_num}]\n"
                f"Título: {ref.get('title', '')}\n"
                f"Autores: {', '.join(ref.get('authors') or [])}\n"
                f"Ano: {ref.get('year', '')}\n"
                f"URL: {url}\n\n"
                f"Resumo:\n{summary}"
            )
            pdf_bytes = _text_to_pdf_bytes(
                title=f"[{ref_num}] {ref.get('title', '')}",
                content=content,
            )
            if pdf_bytes:
                logger.debug("[%d] Saved as abstract PDF.", ref_num)

        if pdf_bytes:
            out_path.write_bytes(pdf_bytes)
            console.print(f"  [green]✓[/green] [{ref_num:02d}] {out_path.name}")
        else:
            failed.append(ref_num)
            console.print(f"  [red]✗[/red] [{ref_num:02d}] não foi possível salvar")

    return failed


def _signoff_checkpoint(workspace_path: Path, checkpoint: int) -> None:
    """Approve the last preview for a checkpoint and save as final (markup stripped).

    Args:
        workspace_path: Path to the ResearchWorkspace (.state/) directory.
        checkpoint: Checkpoint number to finalize (1 or 2).
    """
    from ai_skill.core.workspace import ResearchWorkspace

    ws = ResearchWorkspace(workspace_path)
    state = ws.load_state()

    if state is None:
        console.print(f"[red]Nenhum estado encontrado em:[/red] {workspace_path}")
        raise typer.Exit(1)

    project_path = workspace_path.parent
    if not (project_path / "workspace.yaml").exists():
        console.print("[red]Workspace não reconhecido.[/red]")
        raise typer.Exit(1)

    from ai_skill.core.project_workspace import ProjectWorkspace
    pw = ProjectWorkspace.from_path(project_path)

    last_preview = pw.get_last_preview(checkpoint)
    if last_preview is None:
        console.print(
            f"[red]Nenhum preview encontrado para CP{checkpoint}.[/red]\n"
            "[dim]Execute a pesquisa primeiro.[/dim]"
        )
        raise typer.Exit(1)

    final_path = pw.checkpoint_final_path(checkpoint)
    if final_path.exists():
        console.print(
            f"\n[bold yellow]⚠ Já existe um arquivo final:[/bold yellow]\n"
            f"  {final_path.name}"
        )
        if not typer.confirm("Deseja sobrescrever?", default=True):
            console.print("[dim]Arquivo final mantido sem alteração.[/dim]")
            raise typer.Exit(0)

    console.print(f"[dim]Gerando arquivo final a partir de:[/dim] {last_preview.name}")
    stripped = _strip_docx_markup(last_preview, final_path)
    if not stripped:
        console.print("[yellow]⚠ Não foi possível remover marcações — arquivo copiado como está.[/yellow]")

    pw.update_metadata({"current_checkpoint": checkpoint, "status": f"checkpoint_{checkpoint}_final"})

    if checkpoint == 1:
        state["charter_approved"] = True
        ws.save_state(state)
        console.print(f"\n[bold green]✓ Checkpoint 1 aprovado[/bold green] → [bold]{final_path.name}[/bold]")
        _handle_cp2_start(state, workspace_path)
    elif checkpoint == 2:
        state["literature_approved"] = True
        ws.save_state(state)
        console.print(f"\n[bold green]✓ Checkpoint 2 aprovado[/bold green] → [bold]{final_path.name}[/bold]")
        # Download all referenced sources to attachments/
        review_doc = state.get("literature_review_doc") or {}
        refs = review_doc.get("references", [])
        if refs:
            console.print(
                f"\n[bold]Baixando {len(refs)} fonte(s) referenciada(s) para attachments/[/bold]"
            )
            failed_refs = _download_references_to_attachments(review_doc, pw)
            if failed_refs:
                console.print(
                    f"\n[yellow]⚠ {len(failed_refs)} fonte(s) não puderam ser salvas "
                    f"— verifique manualmente:[/yellow]"
                )
                for rn in sorted(failed_refs):
                    console.print(f"  [{rn}]")
        _handle_cp3_start(state, workspace_path)
    elif checkpoint == 3:
        state["design_approved"] = True
        ws.save_state(state)
        console.print(f"\n[bold green]✓ Checkpoint 3 aprovado[/bold green] → [bold]{final_path.name}[/bold]")
        console.print("[dim]Research Design finalizado. O próximo passo será CP4 (coleta de dados).[/dim]")
    else:
        ws.save_state(state)
        console.print(f"\n[bold green]✓ Checkpoint {checkpoint} aprovado[/bold green] → [bold]{final_path.name}[/bold]")


_NODE_LABELS: dict[str, str] = {
    # CP1 — Research Charter
    "initiate":         "Iniciando workspace",
    "align_charter":    "Redigindo Research Charter",
    "review_charter":   "Processando revisão do Charter",
    # CP2 — Literature Review
    "cp2_router":       "Iniciando Revisão Bibliográfica",
    "plan":             "Planejando próximas buscas",
    "execute":          "Executando buscas bibliográficas",
    "evaluate":         "Avaliando qualidade dos resultados",
    "request_support":  "Solicitando orientação adicional",
    "compile_literature":  "Compilando Revisão Bibliográfica",
    "verify_literature":   "Verificando fontes e referências",
    "deliver_literature":  "Gerando documento preview",
    "review_literature":   "Processando revisão da literatura",
    "recheck_sources":     "Rechecando fontes via API",
    "refine_literature":   "Refinando com base nas correções",
    # CP3 — Research Design
    "cp3_router":       "Iniciando Research Design",
    "compile_design":   "Compilando Research Design",
    "deliver_design":   "Gerando documento preview",
    "review_design":    "Processando revisão do design",
    "refine_design":    "Refinando com base nas correções",
}


def _print_node_progress(node_name: str, output: object, prev_findings_count: int = 0) -> None:
    """Print a user-friendly progress line for a completed node.

    Shows the friendly label, attempt number (for retry nodes), and
    a summary of sub-steps when the output provides that information.

    Args:
        node_name: The LangGraph node name.
        output: The node's partial state output dict.
        prev_findings_count: Number of findings in state before this node ran,
            used to compute the delta for the execute node display.
    """
    label = _NODE_LABELS.get(node_name, node_name)
    extra = ""

    if isinstance(output, dict):
        attempt = output.get("attempt")
        findings = output.get("findings")
        evaluation = output.get("evaluation")

        if attempt is not None and node_name == "plan":
            extra = f" [dim](tentativa {attempt + 1})[/dim]"
        elif node_name == "execute" and findings is not None:
            new_n = len(findings) - prev_findings_count
            total_n = len(findings)
            extra = f" [dim](+{new_n} novos · {total_n} no total)[/dim]"
        elif node_name == "evaluate" and evaluation:
            score = evaluation.get("total_score", 0.0)
            converged = evaluation.get("converged", False)
            flag = "[green]convergiu[/green]" if converged else "[yellow]ainda em progresso[/yellow]"
            extra = f" [dim](score {score:.2f} — {flag})[/dim]"
        elif node_name == "align_charter" and not output:
            label = "Research Charter (sem alterações — aguardando revisão)"
        elif node_name == "verify_literature":
            verified = (output.get("literature_review_doc") or {}).get("verified_sources", [])
            if verified:
                ok = sum(1 for v in verified if v.get("content_matches"))
                extra = f" [dim]({ok}/{len(verified)} verificadas)[/dim]"

    console.print(f"  [green]✓[/green] {label}{extra}")


def _handle_support_request(state: dict, workspace: Path) -> None:  # type: ignore[type-arg]
    """Interactively collect researcher guidance after convergence failure.

    Prints a diagnostic summary (quality score, sources consulted, per-metric
    breakdown, gaps and an agent opinion), then opens a terminal input so the
    researcher can supply revised instructions.  The guidance is saved to
    ``state["user_guidance"]`` and persisted to disk.  The attempt counter is
    also reset so the next run starts fresh.

    Args:
        state: Accumulated research state at the support interrupt point.
        workspace: Path to the ResearchWorkspace (.state/) directory.
    """
    from ai_skill.core.workspace import ResearchWorkspace

    evaluation: dict = state.get("evaluation") or {}
    score: float = evaluation.get("total_score", 0.0)
    per_metric: list = evaluation.get("per_metric", [])
    gaps: list = evaluation.get("gaps", [])
    attempt: int = int(state.get("attempt", 0))
    findings: list = state.get("findings") or []

    # Count unique source URLs across all skill outputs
    all_sources: set = set()
    for f in findings:
        for src in (f.get("sources") or []):
            all_sources.add(src)
    source_count = len(all_sources)

    # Build agent opinion based on score and metric breakdown
    if score < 0.4:
        opinion = "O agente não conseguiu encontrar fontes suficientes ou relevantes para o tema."
    elif score < 0.6:
        opinion = "As buscas retornaram resultados parciais, mas falta cobertura e profundidade temática."
    else:
        opinion = "O agente está próximo da convergência, mas critérios específicos não foram atingidos."

    if per_metric:
        worst = min(per_metric, key=lambda m: m.get("score", 1.0))
        opinion += (
            f" O critério mais deficiente foi: \"{worst.get('metric', '?')}\""
            f" ({worst.get('score', 0.0):.0%})."
        )
        if worst.get("gaps"):
            opinion += f" Lacuna específica: {worst['gaps'][0]}"

    console.print("\n[bold red]━━ ORIENTAÇÃO ADICIONAL NECESSÁRIA ━━[/bold red]")
    console.print(
        f"\n[bold]Diagnóstico — após {attempt} tentativas sem convergência:[/bold]"
    )
    console.print(f"  [cyan]Score de qualidade atual:[/cyan]  {score:.2f} / 1.00")
    console.print(f"  [cyan]Fontes únicas consultadas:[/cyan] {source_count}")

    if per_metric:
        console.print("\n[bold]Avaliação por critério:[/bold]")
        for m in per_metric:
            s = m.get("score", 0.0)
            name = m.get("metric", "?")
            color = "green" if s >= 0.75 else "yellow" if s >= 0.5 else "red"
            console.print(f"  [{color}]{s:.0%}[/{color}]  {name}")

    if gaps:
        console.print("\n[bold]Lacunas identificadas:[/bold]")
        for g in gaps:
            console.print(f"  [yellow]•[/yellow] {g}")

    console.print(f"\n[bold]Avaliação do agente:[/bold] {opinion}")
    console.print(
        "\n[dim]Forneça orientação para o agente replanejar as buscas.\n"
        "Exemplos: sugerir termos, restringir período, indicar fontes específicas.[/dim]"
    )

    guidance = typer.prompt("\nOrientação adicional").strip()

    state["user_guidance"] = guidance if guidance else None
    state["attempt"] = 0  # reset so the next run doesn't immediately re-trigger support

    from ai_skill.core.workspace import ResearchWorkspace  # noqa: F811
    ResearchWorkspace(workspace).save_state(state)

    console.print("\n[green]Orientação registrada.[/green] Retome a pesquisa com:")
    console.print(f"  ai-skill literature --begin --workspace \"{workspace.parent}\"")


def _handle_checkpoint(
    _interrupt_data: object,
    state: dict,  # type: ignore[type-arg]
    workspace: Path,
) -> None:
    """Dispatch to the correct checkpoint handler.

    Dispatch logic:
    - convergence failed (attempt >= max_retries) → support request prompt
    - active_checkpoint == 2                       → CP2 review prompt
    - active_checkpoint == 1, approved             → CP2 start prompt
    - active_checkpoint == 1, not yet              → CP1 review prompt

    Args:
        _interrupt_data: Data from the ``__interrupt__`` event (unused).
        state: Accumulated research state at the point of interruption.
        workspace: Path to the ResearchWorkspace (.state/) directory.
    """
    import os as _os
    max_retries = int(_os.environ.get("AI_SKILL_MAX_RETRIES", "5"))
    attempt = int(state.get("attempt", 0))
    evaluation: dict = state.get("evaluation") or {}

    if attempt >= max_retries and not evaluation.get("converged"):
        _handle_support_request(state, workspace)
        return

    active_cp = int(state.get("active_checkpoint", 1))
    if active_cp == 3:
        _handle_checkpoint_3(state, workspace)
    elif active_cp == 2:
        _handle_checkpoint_2(state, workspace)
    elif state.get("literature_approved"):
        _handle_cp3_start(state, workspace)
    elif state.get("charter_approved"):
        _handle_cp2_start(state, workspace)
    else:
        _handle_checkpoint_1(state, workspace)


def _handle_checkpoint_1(state: dict, workspace: Path) -> None:  # type: ignore[type-arg]
    """Display Checkpoint 1 (Research Charter) summary and resume instructions.

    Args:
        state: Accumulated research state.
        workspace: Path to the ResearchWorkspace (.state/) directory.
    """
    console.print("\n[bold yellow]━━ CHECKPOINT 1 — Research Charter ━━[/bold yellow]")

    docx_path = state.get("checkpoint_label", "")

    objective = state.get("objective") or {}
    if objective.get("goals"):
        console.print("\n[bold]Proposta do agente:[/bold]")
        console.print(f"  [cyan]Tópico:[/cyan] {objective.get('topic', '')}")
        goals = objective.get("goals", [])
        if goals:
            console.print("  [cyan]Objetivos:[/cyan]")
            for g in goals:
                console.print(f"    • {g}")
        metrics = objective.get("success_metrics", [])
        if metrics:
            console.print("  [cyan]Métricas de sucesso:[/cyan]")
            for m in metrics:
                console.print(f"    • {m}")

    if docx_path:
        console.print("\n[green]✓ Documento salvo em:[/green]")
        console.print(f"  [bold]{docx_path}[/bold]")
        console.print("\n[dim]Abra o arquivo, revise e edite conforme necessário.[/dim]")

    console.print(
        f"\n[bold]Para aprovar[/bold] (sem alterações):\n"
        f"  ai-skill research --signoff --workspace \"{workspace}\"\n"
        f"\n[bold]Para solicitar correções[/bold]:\n"
        f"  1. Abra o preview acima no Word\n"
        f"  2. Use uma das formas para indicar correções:\n"
        f"       [dim]• Comentários[/dim]       → aciona a IA para implementar a instrução\n"
        f"       [dim]• Marcas de revisão[/dim] → a IA incorpora as suas alterações\n"
        f"       [dim]• Realce amarelo[/dim]    → a IA regera o trecho completamente\n"
        f"  3. Salve o arquivo\n"
        f"  4. ai-skill research --review --workspace \"{workspace}\""
    )


def _handle_cp2_start(state: dict, workspace: Path) -> None:  # type: ignore[type-arg]
    """Display the CP1 approval confirmation and CP2 start instructions.

    Shown when the graph pauses at the ``begin_literature`` gate after the
    user approved the Research Charter.

    Args:
        state: Accumulated research state.
        workspace: Path to the ResearchWorkspace (.state/) directory.
    """
    topic = (state.get("objective") or {}).get("topic", "")
    title = f"✓ Checkpoint 1 — Research Charter aprovado"
    if topic:
        title += f": [italic]{topic}[/italic]"
    console.print(f"\n[bold green]{title}[/bold green]")

    project_path = workspace.parent
    if (project_path / "workspace.yaml").exists():
        from ai_skill.core.project_workspace import ProjectWorkspace
        pw = ProjectWorkspace.from_path(project_path)
        final = pw.checkpoint_final_path(1)
        if final.exists():
            console.print(f"  [dim]Arquivo final:[/dim] {final.name}")

    console.print(
        f"\n[bold]Para iniciar a Revisão Bibliográfica (Checkpoint 2):[/bold]\n"
        f"  ai-skill literature --begin --workspace \"{workspace.parent}\"\n"
        f"\n[dim]O agente irá pesquisar, compilar e verificar as referências "
        f"bibliográficas automaticamente.[/dim]"
    )


def _handle_checkpoint_2(state: dict, workspace: Path) -> None:  # type: ignore[type-arg]
    """Display Checkpoint 2 (Literature Review) summary and resume instructions.

    Args:
        state: Accumulated research state.
        workspace: Path to the ResearchWorkspace (.state/) directory.
    """
    console.print("\n[bold yellow]━━ CHECKPOINT 2 — Revisão Bibliográfica ━━[/bold yellow]")

    docx_path = state.get("checkpoint_label", "")
    review_doc = state.get("literature_review_doc") or {}
    sections = review_doc.get("sections", [])
    references = review_doc.get("references", [])
    verified = review_doc.get("verified_sources", [])

    if sections:
        console.print(f"\n[bold]Seções geradas ({len(sections)}):[/bold]")
        for s in sections:
            console.print(f"  [cyan]•[/cyan] {s.get('section_title', '')}")

    if references:
        console.print(f"\n[bold]Referências bibliográficas ({len(references)}):[/bold]")

    if verified:
        ok = sum(1 for v in verified if v.get("content_matches"))
        warn = sum(1 for v in verified if v.get("accessible") and not v.get("content_matches"))
        fail = sum(1 for v in verified if not v.get("accessible"))
        console.print(
            f"\n[bold]Verificação de fontes:[/bold]  "
            f"[green]✓ {ok} verificadas[/green]  "
            f"[yellow]⚠ {warn} questionáveis[/yellow]  "
            f"[red]✗ {fail} inacessíveis[/red]"
        )

    if docx_path:
        console.print(f"\n[bold green]✓ Revisão Bibliográfica concluída[/bold green]")
        console.print(f"  [dim]Preview salvo em:[/dim] [bold]{docx_path}[/bold]")
        console.print(
            "\n[dim]Abra o arquivo, avalie os textos e as referências bibliográficas.\n"
            "Use comentários, marcas de revisão ou realce amarelo para indicar correções.[/dim]"
        )
    else:
        console.print("\n[bold yellow]⚠ Revisão Bibliográfica concluída sem preview[/bold yellow]")
        console.print("[dim]Reexecute com:[/dim]  ai-skill literature --begin")

    console.print(
        f"\n[bold]Para aprovar[/bold] (sem alterações):\n"
        f"  ai-skill literature --signoff --workspace \"{workspace}\"\n"
        f"\n[bold]Para solicitar correções[/bold]:\n"
        f"  1. Abra o preview acima no Word\n"
        f"  2. Use uma das formas para indicar correções:\n"
        f"       [dim]• Comentários[/dim]       → aciona a IA para implementar a instrução\n"
        f"       [dim]• Marcas de revisão[/dim] → a IA incorpora as suas alterações\n"
        f"       [dim]• Realce amarelo[/dim]    → a IA regera o trecho completamente\n"
        f"  3. Salve o arquivo\n"
        f"  4. ai-skill literature --review --workspace \"{workspace}\""
    )


def _handle_cp3_start(state: dict, workspace: Path) -> None:  # type: ignore[type-arg]
    """Display the CP2 approval confirmation and CP3 start instructions.

    Shown after the user approves the Literature Review (CP2).

    Args:
        state: Accumulated research state.
        workspace: Path to the ResearchWorkspace (.state/) directory.
    """
    topic = (state.get("objective") or {}).get("topic", "")
    title = "✓ Checkpoint 2 — Revisão Bibliográfica aprovada"
    if topic:
        title += f": [italic]{topic}[/italic]"
    console.print(f"\n[bold green]{title}[/bold green]")

    project_path = workspace.parent
    if (project_path / "workspace.yaml").exists():
        from ai_skill.core.project_workspace import ProjectWorkspace
        pw = ProjectWorkspace.from_path(project_path)
        final = pw.checkpoint_final_path(2)
        if final.exists():
            console.print(f"  [dim]Arquivo final:[/dim] {final.name}")

    console.print(
        f"\n[bold]Para iniciar o Research Design (Checkpoint 3):[/bold]\n"
        f"  ai-skill design --begin --workspace \"{project_path}\"\n"
        f"\n[dim]O agente irá identificar o método de pesquisa, formular hipóteses,\n"
        f"operacionalizar variáveis e definir os instrumentos de coleta automaticamente.[/dim]"
    )


def _handle_checkpoint_3(state: dict, workspace: Path) -> None:  # type: ignore[type-arg]
    """Display Checkpoint 3 (Research Design) summary and resume instructions.

    Args:
        state: Accumulated research state.
        workspace: Path to the ResearchWorkspace (.state/) directory.
    """
    console.print("\n[bold yellow]━━ CHECKPOINT 3 — Research Design ━━[/bold yellow]")

    docx_path = state.get("checkpoint_label", "")
    design_doc = state.get("research_design_doc") or {}

    study_type = design_doc.get("study_type", "")
    paradigm = design_doc.get("research_paradigm", "")
    stance = design_doc.get("epistemological_stance", "")
    hypotheses = design_doc.get("hypotheses") or []
    variables = design_doc.get("variables") or []
    instruments = design_doc.get("instruments") or []
    reporting_std = design_doc.get("reporting_standard", "")
    journal_tier = design_doc.get("target_journal_tier", "")

    if study_type:
        console.print(f"\n[bold]Classificação:[/bold]")
        console.print(f"  [cyan]Tipo de estudo:[/cyan] {study_type}")
        if paradigm:
            console.print(f"  [cyan]Paradigma:[/cyan] {paradigm}")
        if stance:
            console.print(f"  [cyan]Stance epistemológica:[/cyan] {stance}")

    if hypotheses:
        console.print(f"\n[bold]Hipóteses ({len(hypotheses)}):[/bold]")
        for h in hypotheses[:3]:
            console.print(f"  [cyan]•[/cyan] {h[:120]}{'…' if len(h) > 120 else ''}")
        if len(hypotheses) > 3:
            console.print(f"  [dim]... e mais {len(hypotheses) - 3} hipóteses[/dim]")

    if variables:
        console.print(f"\n[bold]Variáveis ({len(variables)}):[/bold]")
        for v in variables[:5]:
            name = v.get("name", "") if isinstance(v, dict) else str(v)
            vtype = v.get("type", "") if isinstance(v, dict) else ""
            console.print(f"  [cyan]•[/cyan] {name}" + (f" [{vtype}]" if vtype else ""))
        if len(variables) > 5:
            console.print(f"  [dim]... e mais {len(variables) - 5} variáveis[/dim]")

    if instruments:
        console.print(f"\n[bold]Instrumentos de coleta ({len(instruments)}):[/bold]")
        for inst in instruments[:3]:
            console.print(f"  [cyan]•[/cyan] {inst[:100]}{'…' if len(inst) > 100 else ''}")

    if reporting_std or journal_tier:
        console.print(f"\n[bold]Publicação:[/bold]")
        if reporting_std:
            console.print(f"  [cyan]Padrão de reporte:[/cyan] {reporting_std}")
        if journal_tier:
            console.print(f"  [cyan]Tier SJR alvo:[/cyan] {journal_tier}")

    if docx_path:
        console.print(f"\n[bold green]✓ Research Design concluído[/bold green]")
        console.print(f"  [dim]Preview salvo em:[/dim] [bold]{docx_path}[/bold]")
        console.print(
            "\n[dim]Abra o arquivo, avalie o design metodológico.\n"
            "Use comentários, marcas de revisão ou realce amarelo para indicar correções.[/dim]"
        )
    else:
        console.print("\n[bold yellow]⚠ Research Design concluído sem preview[/bold yellow]")
        console.print("[dim]Reexecute com:[/dim]  ai-skill design --begin")

    console.print(
        f"\n[bold]Para aprovar[/bold] (sem alterações):\n"
        f"  ai-skill design --signoff --workspace \"{workspace}\"\n"
        f"\n[bold]Para solicitar correções[/bold]:\n"
        f"  1. Abra o preview acima no Word\n"
        f"  2. Use uma das formas para indicar correções:\n"
        f"       [dim]• Comentários[/dim]       → aciona a IA para implementar a instrução\n"
        f"       [dim]• Marcas de revisão[/dim] → a IA incorpora as suas alterações\n"
        f"       [dim]• Realce amarelo[/dim]    → a IA regera o trecho completamente\n"
        f"  3. Salve o arquivo\n"
        f"  4. ai-skill design --review --workspace \"{workspace}\""
    )


@app.command("design")
def design_cmd(
    begin: Annotated[
        bool,
        typer.Option("--begin", is_flag=True, help="Iniciar Research Design (CP3)."),
    ] = False,
    review: Annotated[
        bool,
        typer.Option("--review", is_flag=True, help="Aplicar correções ao último preview do CP3."),
    ] = False,
    signoff: Annotated[
        bool,
        typer.Option("--signoff", is_flag=True, help="Aprovar preview e gerar arquivo final sem marcações."),
    ] = False,
    workspace: Annotated[
        Path | None,
        typer.Option("--workspace", "-w", help="Diretório do workspace."),
    ] = None,
) -> None:
    """Gerenciar o Research Design (Checkpoint 3).

    Use uma das flags para selecionar a ação:\n
      --begin     Iniciar a pesquisa de design metodológico\n
      --review    Revisar o último preview (aplicar correções do Word)\n
      --signoff   Aprovar o preview atual e gerar arquivo final
    """
    flags = [begin, review, signoff]
    if sum(flags) > 1:
        console.print("[red]Use apenas uma flag por vez: --begin, --review ou --signoff.[/red]")
        raise typer.Exit(1)
    if not any(flags):
        console.print(
            "\n[bold]ai-skill design[/bold]  —  Research Design (CP3)\n\n"
            "  [cyan]--begin[/cyan]     Iniciar a pesquisa de design metodológico\n"
            "  [cyan]--review[/cyan]    Revisar o último preview (aplicar correções do Word)\n"
            "  [cyan]--signoff[/cyan]   Aprovar o preview atual e gerar arquivo final\n"
        )
        raise typer.Exit(0)

    if begin:
        begin_design(workspace=workspace)
    elif review:
        _design_review(workspace=workspace)
    elif signoff:
        ws_path = _resolve_workspace_path(workspace, state_subdir=True)
        if ws_path is None:
            raise typer.Exit(1)
        _signoff_checkpoint(ws_path, 3)


@app.command("begin-design")
def begin_design(
    workspace: Annotated[
        Path | None,
        typer.Option(
            "--workspace", "-w",
            help="ProjectWorkspace ou diretório .state/ do projeto.",
        ),
    ] = None,
) -> None:
    """Iniciar o Research Design (Checkpoint 3).

    Requer que o Checkpoint 2 (Literature Review) já esteja aprovado.
    """
    _configure_logging(os.environ.get("AI_SKILL_LOG_LEVEL", "WARNING"))

    # --- Resolve project path ---
    resolved = _resolve_workspace_path(workspace, state_subdir=False)
    if resolved is None:
        raise typer.Exit(1)

    if (resolved / "workspace.yaml").exists():
        project_path = resolved
        state_path = resolved / ".state"
    elif (resolved.parent / "workspace.yaml").exists():
        project_path = resolved.parent
        state_path = resolved
    else:
        console.print(
            "[red]Workspace não reconhecido.[/red]\n"
            "[dim]Aponte --workspace para o ProjectWorkspace "
            "ou para o diretório .state/ dentro dele.[/dim]"
        )
        raise typer.Exit(1)

    # --- Validate CP2 [final].docx exists ---
    from ai_skill.core.project_workspace import ProjectWorkspace
    pw = ProjectWorkspace.from_path(project_path)
    final_cp2 = pw.checkpoint_final_path(2)

    if not final_cp2.exists():
        final_cp1 = pw.checkpoint_final_path(1)
        console.print(
            f"\n[red]✗ Checkpoint 2 não aprovado.[/red]\n"
            f"  O arquivo [bold]{final_cp2.name}[/bold] não foi encontrado em:\n"
            f"  {project_path}\n"
            f"\n[dim]Complete o Checkpoint 2 antes de iniciar o Research Design:[/dim]\n"
            f"  ai-skill literature --begin --workspace \"{project_path}\"\n"
            f"  ai-skill literature --signoff --workspace \"{state_path}\""
        )
        raise typer.Exit(1)

    console.print(f"[green]✓[/green] Checkpoint 2 aprovado: [dim]{final_cp2.name}[/dim]")

    # --- Warn if CP3 work already exists ---
    _warn_overwrite(project_path, checkpoints=[3])

    # --- Load or reconstruct state ---
    from ai_skill.core.graph import build_cp3_graph
    from ai_skill.core.workspace import ResearchWorkspace

    ws = ResearchWorkspace(state_path)
    state = ws.load_state()

    if state is None:
        state = _reconstruct_state_for_design(project_path, state_path)
        if state is None:
            console.print(
                "[red]Não foi possível carregar o estado do projeto.[/red]\n"
                "[dim]Tente re-rodar:[/dim]\n"
                f"  ai-skill literature --begin --workspace \"{project_path}\""
            )
            raise typer.Exit(1)

    # Guarantee CP1 and CP2 are marked approved so graph routing passes gates
    state["charter_approved"] = True
    state["literature_approved"] = True
    state["user_feedback"] = None
    state["attempt"] = 0

    # Inject CP1 charter text for CP3 compile context
    final_cp1 = pw.checkpoint_final_path(1)
    if final_cp1.exists():
        state["charter_document_text"] = _read_docx_text(final_cp1)

    # Set correct stage for CP3 routing
    from ai_skill.core.pipeline_stages import PipelineStage as _PS
    state["stage"] = _PS.RESEARCH_DESIGN

    # Build the CP3 handoff context
    _full_obj = state.get("objective") or {}
    review_doc = state.get("literature_review_doc") or {}
    sections = review_doc.get("sections") or []
    # Condense literature summary: titles + first 300 chars of each section
    literature_summary = "\n\n".join(
        f"## {s.get('section_title', '')}\n{(s.get('content') or '')[:300]}"
        for s in sections
    )
    state["cp3_context"] = {
        "topic":                   _full_obj.get("topic") or "",
        "goals":                   list(_full_obj.get("goals") or []),
        "scope_constraints":       list(_full_obj.get("scope_constraints") or []),
        "methodology_preference":  _full_obj.get("methodology_preference") or "",
        "literature_summary":      literature_summary,
    }

    topic = _full_obj.get("topic", "")
    console.print(
        f"\n[bold]Iniciando Research Design[/bold]"
        + (f": {topic}" if topic else "")
    )

    # --- Stream the CP3 graph ---
    graph_instance = build_cp3_graph()
    current_state: dict = dict(state)

    try:
        for event in graph_instance.stream(state, stream_mode="updates"):
            for node_name, node_output in event.items():
                if node_name == "__interrupt__":
                    ws.save_state(current_state)
                    _handle_checkpoint(node_output, current_state, state_path)
                    return
                _print_node_progress(node_name, node_output, len(current_state.get("findings", [])))
                if isinstance(node_output, dict):
                    current_state.update(node_output)
                    ws.save_state(current_state)
                    if (
                        node_name == "review_design"
                        and node_output.get("design_approved")
                    ):
                        _finalize_checkpoint(state_path, 3)
    except _GRAPH_INTERRUPT as _exc:  # type: ignore[misc]
        ws.save_state(current_state)
        _handle_checkpoint(_exc, current_state, state_path)
        return
    except KeyboardInterrupt:
        ws.save_state(current_state)
        console.print("\n[yellow]Pesquisa pausada. Retome com:[/yellow]")
        console.print(f"  ai-skill design --review --workspace \"{state_path}\"")
        sys.exit(0)


def _design_review(workspace: Path | None) -> None:
    """Apply researcher corrections from the last CP3 preview and re-run."""
    from ai_skill.core.graph import build_cp3_graph
    from ai_skill.core.workspace import ResearchWorkspace

    _configure_logging(os.environ.get("AI_SKILL_LOG_LEVEL", "WARNING"))

    resolved = _resolve_workspace_path(workspace, state_subdir=True)
    if resolved is None:
        raise typer.Exit(1)

    if (resolved / "workspace.yaml").exists():
        project_path = resolved
        state_path = resolved / ".state"
    elif (resolved.parent / "workspace.yaml").exists():
        project_path = resolved.parent
        state_path = resolved
    else:
        console.print("[red]Workspace não reconhecido.[/red]")
        raise typer.Exit(1)

    from ai_skill.core.project_workspace import ProjectWorkspace
    pw = ProjectWorkspace.from_path(project_path)

    last_preview = pw.get_last_preview(3)
    if last_preview is None:
        console.print(
            "[red]Nenhum preview do CP3 encontrado.[/red]\n"
            "[dim]Execute primeiro:[/dim]  "
            f"ai-skill design --begin --workspace \"{project_path}\""
        )
        raise typer.Exit(1)

    ws = ResearchWorkspace(state_path)
    state = ws.load_state() or {}

    feedback = _extract_docx_corrections(last_preview)
    if not feedback:
        console.print(
            "[yellow]Nenhuma correção encontrada no preview.[/yellow]\n"
            "[dim]Adicione comentários, marcas de revisão ou realce amarelo no arquivo:[/dim]\n"
            f"  {last_preview}"
        )
        raise typer.Exit(0)

    formatted_feedback = _format_corrections_for_llm(feedback)
    state["user_feedback"] = formatted_feedback
    state["design_approved"] = False

    from ai_skill.core.pipeline_stages import PipelineStage as _PS
    state["stage"] = _PS.RESEARCH_DESIGN

    graph_instance = build_cp3_graph()
    current_state: dict = dict(state)

    try:
        for event in graph_instance.stream(state, stream_mode="updates"):
            for node_name, node_output in event.items():
                if node_name == "__interrupt__":
                    ws.save_state(current_state)
                    _handle_checkpoint(node_output, current_state, state_path)
                    return
                _print_node_progress(node_name, node_output, len(current_state.get("findings", [])))
                if isinstance(node_output, dict):
                    current_state.update(node_output)
                    ws.save_state(current_state)
                    if (
                        node_name == "review_design"
                        and node_output.get("design_approved")
                    ):
                        _finalize_checkpoint(state_path, 3)
    except _GRAPH_INTERRUPT as _exc:  # type: ignore[misc]
        ws.save_state(current_state)
        _handle_checkpoint(_exc, current_state, state_path)
        return
    except KeyboardInterrupt:
        ws.save_state(current_state)
        console.print("\n[yellow]Revisão pausada. Retome com:[/yellow]")
        console.print(f"  ai-skill design --review --workspace \"{state_path}\"")
        sys.exit(0)


def _reconstruct_state_for_design(project_path: Path, state_path: Path) -> dict | None:
    """Build a minimal ResearchState for the design command.

    Used when the state file is missing but CP2 [final].docx exists.

    Args:
        project_path: Root of the ProjectWorkspace.
        state_path: Path to the .state/ directory.

    Returns:
        A state dict with ``literature_approved=True``, or None on failure.
    """
    import yaml as _yaml
    from ai_skill.core.project_workspace import ProjectWorkspace
    from ai_skill.core.state import initial_state

    meta_file = project_path / "workspace.yaml"
    if not meta_file.exists():
        return None
    try:
        meta = _yaml.safe_load(meta_file.read_text(encoding="utf-8")) or {}
    except Exception:
        return None

    topic = meta.get("topic", "")
    state: dict = dict(initial_state(str(state_path), topic=topic))

    pw = ProjectWorkspace.from_path(project_path)
    final_cp1 = pw.checkpoint_final_path(1)
    final_cp2 = pw.checkpoint_final_path(2)

    if final_cp1.exists():
        objective = _parse_objective_from_docx(final_cp1)
        if objective:
            state["objective"] = objective
        state["charter_document_text"] = _read_docx_text(final_cp1)

    state["charter_approved"] = True
    state["literature_approved"] = True

    # CP2 doc text is used as context during design compilation but we don't
    # reconstruct the full literature_review_doc here — compile_design reads
    # the cp3_context["literature_summary"] which is built by begin_design().
    if final_cp2.exists():
        state["active_checkpoint"] = 2  # will be updated to 3 by compile_design

    return state


if __name__ == "__main__":
    app()
