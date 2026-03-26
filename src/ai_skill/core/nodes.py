"""LangGraph node functions for the Phase 1 research pipeline.

Each function takes a ResearchState and returns a partial ResearchState dict.
LangGraph merges partial updates into the full state automatically.

Nodes in Phase 1:
    initiate         — create workspace directory and initial state
    align_charter    — draft and refine the Research Charter with the user
    plan             — generate an ExecutionPlan using the LLM
    execute          — run the skills defined in the current plan
    evaluate         — score findings against success metrics
    deliver          — write findings.md and mark project as converged
    request_support  — display gaps to user and collect revised guidance
"""

from __future__ import annotations

import asyncio
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
import os
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

from ai_skill.core import history as _history
from ai_skill.core.llm_client import LLMClient, LLMClientError, set_current_node
from ai_skill.core.pipeline_stages import PipelineStage, ResearchStatus
from ai_skill.core.project_workspace import ProjectWorkspace
from ai_skill.core.state import (
    DataCollectionGuideDoc,
    DataCollectionSection,
    EvaluationResult,
    ExecutionPlan,
    LiteratureReviewDoc,
    LiteratureReviewSection,
    MetricScore,
    PlanStep,
    QualitySnapshot,
    ResearchDesignDoc,
    ResearchDesignSection,
    ResearchObjective,
    ResearchState,
    SkillOutput,
    SourceVerification,
)
from ai_skill.core.workspace import ResearchWorkspace
from ai_skill.prompts.alignment import (
    build_charter_draft_messages,
    build_charter_refine_messages,
)
from ai_skill.prompts.evaluation import build_evaluation_messages
from ai_skill.prompts.design import (
    build_compile_design_messages,
    build_refine_design_messages,
    build_extract_pdf_methodology_messages,
    build_ideate_design_messages,
    build_review_frameworks_messages,
    build_evaluate_objectives_messages,
)
from ai_skill.prompts.literature import (
    build_compile_messages,
    build_refine_messages,
    build_verify_messages,
)
from ai_skill.prompts.collection import (
    CP4_FRAMEWORK_QUERIES,
    build_draft_collection_guide_messages,
    build_evaluate_collection_objectives_messages,
    build_refine_collection_guide_messages,
    build_review_collection_standards_messages,
)
from ai_skill.prompts.planning import build_planning_messages
from ai_skill.core.rag import RagIndex
from ai_skill.skills.registry import SkillRegistry
from ai_skill.skills.base import SkillInput

logger = logging.getLogger(__name__)

_CONVERGENCE_THRESHOLD = float(
    os.environ.get("AI_SKILL_CONVERGENCE_THRESHOLD", "0.75")
)

_CHUNK_FINDINGS_SIZE = 13   # max findings per chunk call
_CHUNK_SECTIONS_PER_GROUP = 2  # sections per outline group (target)


# ---------------------------------------------------------------------------
# Pydantic models for structured LLM outputs
# ---------------------------------------------------------------------------

class _ExecutionPlanLLM(BaseModel):
    """LLM-parseable representation of an ExecutionPlan."""

    steps: list[dict[str, Any]] = Field(
        description="List of plan steps with skill_name, parameters, depends_on, rationale."
    )
    rationale: str = Field(description="Overall rationale for this plan.")
    estimated_cost: str = Field(
        description="Rough estimate: 'low' (<10 API calls), 'medium' (10-50), 'high' (>50)."
    )


class _EvaluationResultLLM(BaseModel):
    """LLM-parseable representation of an EvaluationResult.

    Note: total_score and converged are NOT in this model — they are computed
    deterministically in Python from per_metric scores to eliminate LLM arithmetic
    errors and premature convergence declarations.
    """

    per_metric: list[dict[str, Any]] = Field(
        description="List of {metric, score, rationale, gaps} dicts."
    )
    gaps: list[str] = Field(description="Aggregated gaps across all metrics.")


class _CharterLLM(BaseModel):
    """LLM-parseable Research Charter."""

    topic: str
    goals: list[str] = Field(min_length=1)
    success_metrics: list[str] = Field(min_length=1)
    scope_constraints: list[str] = Field(default_factory=list)
    methodology_preference: str = Field(default="")
    bibliography_style: str = Field(default="abnt")
    language: str = Field(default="pt-BR")
    stage_guidelines: dict[str, list[str]] = Field(
        default_factory=dict,
        description=(
            "Stage-specific directives keyed by PipelineStage value "
            "(literature_review, research_design, data_collection_guide, "
            "analysis_guide, results_interpretation, paper_composition, publication). "
            "Each entry is a list of 4-8 actionable items driving planning and "
            "evaluation for that stage independently of the overall success_metrics."
        ),
    )


class _LiteratureSourceLLM(BaseModel):
    """One bibliographic entry in the literature review."""

    reference_number: int = Field(description="Sequential citation index starting at 1.")
    title: str = Field(description="Full title of the work.")
    authors: str = Field(description="Author(s) in ABNT order (SOBRENOME, Nome).")
    year: str = Field(description="Publication year as a string.")
    url: str = Field(description="Direct URL to the original document or page.")
    abnt_entry: str = Field(
        description="Complete ABNT NBR 6023:2018 formatted reference entry "
                    "including the URL and the placeholder {ACCESS_DATE}."
    )
    summary: str = Field(
        description=(
            "Summary of the article's content in 300 to 500 words (pt-BR). "
            "Must cover: main argument, methodology, key findings, conclusions, "
            "and relevance to the research objectives — based solely on information "
            "in the findings. The conclusions section is mandatory."
        )
    )


class _LiteratureReviewSectionLLM(BaseModel):
    """One thematic section of the literature review."""

    section_title: str = Field(description="Heading for this section in pt-BR.")
    content: str = Field(
        description="Body text in pt-BR Markdown with inline [N] citations "
                    "referencing the bibliography."
    )


class _LiteratureReviewLLM(BaseModel):
    """Full structured literature review — enforces consistent output format."""

    sections: list[_LiteratureReviewSectionLLM] = Field(
        min_length=1,
        description="Ordered thematic sections that form the review body.",
    )
    references: list[_LiteratureSourceLLM] = Field(
        min_length=1,
        description="Numbered bibliography (must match all [N] used in sections).",
    )


class _SectionGroupLLM(BaseModel):
    """One group of thematically cohesive sections for chunked compilation."""
    titles: list[str] = Field(description="2–3 section titles (pt-BR) for this chunk.")
    themes: list[str] = Field(description="3–6 keywords summarising the themes of this group.")


class _OutlineLLM(BaseModel):
    """Outline plan for chunked literature compilation."""
    section_groups: list[_SectionGroupLLM] = Field(
        min_length=1,
        description="Ordered list of section groups; each group becomes one chunk call.",
    )


class _LiteratureChunkLLM(BaseModel):
    """One compiled chunk of the literature review (2–3 sections + new references)."""

    sections: list[_LiteratureReviewSectionLLM] = Field(
        min_length=1,
        description="Sections generated for this chunk.",
    )
    new_references: list[_LiteratureSourceLLM] = Field(
        default_factory=list,
        description=(
            "Only NEW references introduced in this chunk "
            "(not already cited in prior chunks). Numbered from ref_offset."
        ),
    )


class _ReferenceContributionItemLLM(BaseModel):
    """Marginal contribution score for one reference."""
    url: str = Field(description="URL of the reference (must match the provided list).")
    contribution: float = Field(description="Marginal contribution score 0.0–1.0.")


class _ReferenceContributionLLM(BaseModel):
    """Collection of contribution scores for newly added references."""
    contributions: list[_ReferenceContributionItemLLM] = Field(
        default_factory=list,
        description="One entry per new reference.",
    )


class _VerifyResultLLM(BaseModel):
    """Independent verification result for a single source URL."""

    content_matches: bool = Field(
        description="True when the fetched content is consistent with the cited claim."
    )
    verification_note: str = Field(
        description="One sentence explaining the verification decision."
    )


class _ResearchDesignVariableLLM(BaseModel):
    """One variable entry in the Research Design."""

    name: str = Field(description="Variable name.")
    variable_type: str = Field(
        description="'independente' | 'dependente' | 'confundidora'."
    )
    description: str = Field(description="Plain-language description of the variable.")
    operationalization: str = Field(
        description="How the variable is measured or manipulated."
    )
    measurement_scale: str = Field(
        description="nominal | ordinal | intervalar | racional."
    )


class _ResearchDesignMitigationLLM(BaseModel):
    """One validity threat + mitigation strategy."""

    threat: str = Field(description="Specific validity or reliability threat.")
    mitigation: str = Field(description="Strategy to address or reduce the threat.")
    contingency: str = Field(description="Fallback plan if mitigation is insufficient.")


class _ResearchDesignSectionLLM(BaseModel):
    """One section of the Research Design document."""

    section_title: str = Field(description="Section heading in pt-BR.")
    content: str = Field(description="Body text in pt-BR Markdown prose.")


class _ResearchDesignDocLLM(BaseModel):
    """Full structured Research Design — enforces consistent output format for CP3."""

    sections: list[_ResearchDesignSectionLLM] = Field(
        min_length=1,
        description="Ordered sections covering all 9 mandatory headings.",
    )
    study_type: str = Field(
        description=(
            "experimental | quasi-experimental | observacional | estudo-de-caso | "
            "pesquisa-ação | revisão-sistemática | misto"
        )
    )
    research_paradigm: str = Field(
        description="quantitativo | qualitativo | misto"
    )
    epistemological_stance: str = Field(
        description="pós-positivista | construtivista | pragmático | transformativo"
    )
    hypotheses: list[str] = Field(
        default_factory=list,
        description="Testable hypothesis statements in falsifiable form (H1: …, H2: …).",
    )
    research_questions: list[str] = Field(
        default_factory=list,
        description="Alternative or complementary research questions.",
    )
    variables: list[_ResearchDesignVariableLLM] = Field(
        default_factory=list,
        description="All independent, dependent, and confounding variables.",
    )
    instruments: list[str] = Field(
        default_factory=list,
        description="Data collection instruments with validation references.",
    )
    sampling_strategy: str = Field(description="Description of sampling approach and frame.")
    sample_size_justification: str = Field(
        description="Power analysis or qualitative saturation rationale."
    )
    ethical_considerations: str = Field(
        description="Consent, privacy, CEP/IRB requirements, LGPD compliance."
    )
    validity_threats: list[str] = Field(
        default_factory=list,
        description="Internal and external validity threats identified.",
    )
    mitigation_strategies: list[_ResearchDesignMitigationLLM] = Field(
        default_factory=list,
        description="One entry per validity threat with mitigation and contingency.",
    )
    data_management_plan: str = Field(
        description="FAIR-compliant data management summary."
    )
    metrics_and_kpis: list[str] = Field(
        default_factory=list,
        description="KPIs, acceptance thresholds, statistical power, minimum sample size.",
    )
    data_sources: list[str] = Field(
        default_factory=list,
        description="Primary and secondary data sources with quality criteria and FAIR statement.",
    )
    collection_protocol: str = Field(
        description="Step-by-step data collection protocol with instruments and acceptance criteria."
    )
    methodology_timeline: str = Field(
        description="PMBOK-aligned milestone schedule: CP3 (design) → CP4 (coleta) → CP5 (análise)."
    )
    reporting_standard: str = Field(
        description="CONSORT | STROBE | PRISMA | SQUIRE | outro | não aplicável"
    )
    target_journal_tier: str = Field(
        description="Target publication venue quartile (Q1/Q2/Q3/Q4)."
    )


# ---------------------------------------------------------------------------
# Node functions
# ---------------------------------------------------------------------------

def initiate(state: ResearchState) -> dict[str, Any]:
    """Create the workspace directory and set initial pipeline state.

    Args:
        state: Current ResearchState (may be partially empty on first run).

    Returns:
        Partial state update with stage, status, and workspace_path.
    """
    workspace_dir = os.environ.get("AI_SKILL_WORKSPACE_DIR", "./research-workspace")
    workspace_path = state.get("workspace_path") or workspace_dir
    Path(workspace_path).mkdir(parents=True, exist_ok=True)

    logger.info("Research workspace: %s", workspace_path)

    return {
        "stage": PipelineStage.RESEARCH_CHARTER,
        "status": ResearchStatus.PLANNING,
        "workspace_path": str(Path(workspace_path).resolve()),
        "attempt": 0,
        "findings": [],
        "quality_history": [],
        "document_version": 0,
    }


def align_charter(state: ResearchState, llm: LLMClient | None = None) -> dict[str, Any]:
    """Draft or refine the Research Charter and save it as a .docx checkpoint.

    On the first call (no goals yet), drafts the charter from the topic.
    On subsequent calls (user provided feedback via review_charter), refines it.
    Always saves ``Checkpoint 1 - Research Charter.docx`` to the ProjectWorkspace.

    Args:
        state: Current ResearchState. ``user_feedback`` carries revision
            instructions from the previous review cycle.
        llm: Optional injected LLMClient (for testing).

    Returns:
        Partial state update with the updated ``objective`` and cleared
        ``user_feedback`` (consumed by this node).
    """
    _configure_history(state, "align_charter")
    _llm = llm or LLMClient()
    user_feedback = state.get("user_feedback")
    existing_objective = state.get("objective")
    has_full_charter = bool(existing_objective and existing_objective.get("goals"))

    # Skip if a draft already exists and there's nothing new to refine.
    # This covers resume-to-approve cycles where the graph re-runs from START.
    if has_full_charter and not user_feedback:
        logger.debug("align_charter: draft exists, no feedback — skipping LLM call.")
        return {}

    if has_full_charter and user_feedback:
        # Refinement pass: user reviewed the docx and requested changes
        system, messages = build_charter_refine_messages(existing_objective, user_feedback)
    else:
        # First pass: generate from topic (user_feedback not needed here)
        topic = (existing_objective or {}).get("topic", "")
        if not topic:
            logger.warning("align_charter: no topic in state. Using placeholder.")
            topic = "Research topic not yet specified."
        system, messages = build_charter_draft_messages(topic)

    try:
        charter_obj: _CharterLLM = _llm.complete_structured(
            messages=messages,
            response_model=_CharterLLM,
            system=system,
        )
    except LLMClientError as exc:
        logger.error("Charter generation failed: %s", exc)
        return {"status": ResearchStatus.FAILED}

    # Preserve generated_at from the first draft; never update it on refinement
    existing_generated_at = (existing_objective or {}).get("generated_at", "")
    generated_at = existing_generated_at or _today_dd_mm_yyyy()

    objective = ResearchObjective(
        topic=charter_obj.topic,
        goals=charter_obj.goals,
        success_metrics=charter_obj.success_metrics,
        scope_constraints=charter_obj.scope_constraints,
        methodology_preference=charter_obj.methodology_preference,
        bibliography_style=charter_obj.bibliography_style,
        language=charter_obj.language,
        generated_at=generated_at,
        stage_guidelines=charter_obj.stage_guidelines,
    )

    workspace_path = state.get("workspace_path", "")
    checkpoint_docx_path: str = ""
    if workspace_path:
        ResearchWorkspace(Path(workspace_path)).log(
            f"Research Charter {'refined' if has_full_charter else 'drafted'} "
            f"for topic: {charter_obj.topic}"
        )
        project_ws = _get_project_workspace(workspace_path)
        if project_ws is not None:
            docx_bytes = _charter_to_docx(objective)
            if docx_bytes:
                saved = project_ws.save_checkpoint_preview(1, docx_bytes)
                checkpoint_docx_path = str(saved)
                logger.info("Checkpoint 1 preview saved: %s", saved)

    return {
        "objective": objective,
        "charter_approved": False,
        "checkpoint_label": checkpoint_docx_path,  # path shown at interrupt
        "user_feedback": None,
        "status": ResearchStatus.PLANNING,
    }


def review_charter(state: ResearchState) -> dict[str, Any]:
    """Checkpoint gate executed after the user resumes from a charter review.

    The graph pauses (``interrupt_before``) before this node.  When the user
    resumes with ``--feedback``, that text lands in ``user_feedback`` and this
    node signals that a revision cycle is needed.  When the user resumes
    without feedback, the charter is approved and the pipeline proceeds.

    Args:
        state: Current ResearchState. ``user_feedback`` carries revision
            instructions if the user requested changes.

    Returns:
        Partial state update with ``charter_approved`` set accordingly.
    """
    user_feedback = state.get("user_feedback")
    if user_feedback:
        logger.info("Charter revision requested: %s", user_feedback[:80])
        return {"charter_approved": False, "status": ResearchStatus.PLANNING}
    logger.info("Research Charter approved by user.")
    return {"charter_approved": True, "status": ResearchStatus.PLANNING}


def route_after_review_charter(state: ResearchState) -> str:
    """Conditional edge: decide whether to refine the charter or end CP1.

    Args:
        state: Current ResearchState after ``review_charter`` ran.

    Returns:
        ``"END"`` if charter approved, ``"align_charter"`` if revision needed.
    """
    return "END" if state.get("charter_approved") else "align_charter"


def cp2_router(_state: ResearchState) -> dict[str, Any]:
    """CP2 start router — no-op node; routing logic lives in route_cp2_start.

    Returns:
        Empty dict (no state changes).
    """
    return {}


def route_cp2_start(state: ResearchState) -> str:
    """Conditional edge from cp2_router: decide CP2 entry point.

    Args:
        state: Current ResearchState.

    Returns:
        ``"refine_literature"`` when a compiled review exists and user
        feedback is pending (correction cycle); ``"plan"`` otherwise (fresh
        literature research start).
    """
    if state.get("literature_review_doc") and state.get("user_feedback"):
        return "refine_literature"
    return "plan"


def plan(state: ResearchState, registry: SkillRegistry | None = None, llm: LLMClient | None = None) -> dict[str, Any]:
    """Generate an ExecutionPlan using the LLM.

    Args:
        state: Current ResearchState with objective and previous gaps.
        registry: Optional injected SkillRegistry (for testing).
        llm: Optional injected LLMClient (for testing).

    Returns:
        Partial state update with the new plan.
    """
    _configure_history(state, "plan")
    _registry = registry or _get_default_registry()
    _llm = llm or LLMClient()

    # cp2_context is the restricted handoff from CP1: only topic, goals,
    # scope_constraints.  When present we use it instead of the full objective
    # so success_metrics and other CP1-only fields are structurally absent.
    objective = state.get("cp2_context") or state.get("objective") or ResearchObjective(
        topic="", goals=[], success_metrics=[]
    )
    stage = state.get("stage", PipelineStage.LITERATURE_REVIEW)
    stage_str = str(stage.value if hasattr(stage, "value") else stage)
    attempt = state.get("attempt", 0)
    previous_eval = state.get("evaluation")
    gaps = previous_eval.get("gaps", []) if previous_eval else []
    user_guidance = state.get("user_guidance")

    # For LITERATURE_REVIEW and RESEARCH_DESIGN, always use hardcoded defaults
    # and ignore whatever the charter LLM generated.  The charter LLM consistently
    # mirrors global project outcomes into these slots instead of restricting to
    # stage-specific activities.  Hardcoding prevents the planner from planning
    # toward "submit article" or "build framework" steps during CP2/CP3.
    if stage_str == PipelineStage.LITERATURE_REVIEW.value:
        stage_guidelines: list[str] | None = _default_literature_review_guidelines()
    elif stage_str == PipelineStage.RESEARCH_DESIGN.value:
        stage_guidelines = _default_research_design_guidelines()
    else:
        raw_guidelines = (objective.get("stage_guidelines") or {}).get(stage_str)
        if raw_guidelines:
            stage_guidelines = raw_guidelines
        else:
            defaults = _default_stage_guidelines(stage_str)
            stage_guidelines = defaults if defaults else None

    # CP3 uses only the PDFs already downloaded during CP2 — no new searches.
    # extra_search_approved=True is set by request_support when the user
    # authorises fetching additional sources after a convergence failure.
    available_files: list[str] = []
    if stage_str == PipelineStage.RESEARCH_DESIGN.value:
        extra_search_approved: bool = bool(state.get("extra_search_approved"))
        if extra_search_approved:
            # User approved additional searches — expose all skills
            skill_registry_summary = _registry.all_as_dicts()
        else:
            # Restrict to skills that read already-downloaded material
            _CP3_SKILLS = {"pdf_reader", "content_summarizer", "google_drive"}
            skill_registry_summary = [
                s for s in _registry.all_as_dicts() if s.get("name") in _CP3_SKILLS
            ]
        # List PDFs available in attachments/ so the LLM generates correct file_path params
        ws_path = Path(state.get("workspace_path", ""))
        attachments_dir = (
            ws_path.parent / "attachments"
            if ws_path.name == ".state"
            else ws_path / "attachments"
        )
        if attachments_dir.is_dir():
            available_files = sorted(str(f) for f in attachments_dir.glob("*.pdf"))
    else:
        skill_registry_summary = _registry.all_as_dicts()

    system, messages = build_planning_messages(
        objective=objective,
        stage=stage_str,
        attempt=attempt,
        gaps=gaps,
        skill_registry_summary=skill_registry_summary,
        user_guidance=user_guidance,
        stage_guidelines=stage_guidelines,
        available_files=available_files or None,
    )

    try:
        plan_obj: _ExecutionPlanLLM = _llm.complete_structured(
            messages=messages,
            response_model=_ExecutionPlanLLM,
            system=system,
            max_tokens=16384,  # plans can be verbose; raise ceiling above default 8192
        )
    except LLMClientError as exc:
        logger.error("Plan generation failed: %s", exc)
        return {"status": ResearchStatus.FAILED}

    steps: list[PlanStep] = []
    for i, raw_step in enumerate(plan_obj.steps):
        steps.append(PlanStep(
            step_id=i,
            skill_name=raw_step.get("skill_name", ""),
            parameters=raw_step.get("parameters", {}),
            depends_on=raw_step.get("depends_on", []),
            rationale=raw_step.get("rationale", ""),
        ))

    execution_plan = ExecutionPlan(
        steps=steps,
        rationale=plan_obj.rationale,
        estimated_cost=plan_obj.estimated_cost,
        attempt=attempt,
    )

    # Log a plan_start chapter-header entry so llm_history.yaml has one
    # clear anchor per attempt that groups everything that follows.
    _history.log_plan_start(
        attempt=attempt,
        step_count=len(steps),
        estimated_cost=plan_obj.estimated_cost,
        plan_rationale=plan_obj.rationale,
    )

    return {
        "plan": execution_plan,
        "status": ResearchStatus.EXECUTING,
        "user_guidance": None,  # consumed; clear so it doesn't repeat on next attempt
    }


def execute(
    state: ResearchState,
    registry: SkillRegistry | None = None,
) -> dict[str, Any]:
    """Execute the current plan by running each skill step.

    Steps with no dependencies run concurrently via asyncio.gather.
    Steps with dependencies run after their prerequisites complete.

    Args:
        state: Current ResearchState with plan and objective.
        registry: Optional injected SkillRegistry (for testing).

    Returns:
        Partial state update with findings accumulated from skill outputs.
    """
    _registry = registry or _get_default_registry()
    current_plan = state.get("plan")
    if not current_plan or not current_plan.get("steps"):
        logger.warning("execute: no plan or empty steps — skipping.")
        return {"findings": state.get("findings", []), "findings_current": []}

    objective = state.get("objective")
    stage = state.get("stage", PipelineStage.LITERATURE_REVIEW)
    attempt = state.get("attempt", 0)

    steps = current_plan["steps"]
    outputs: dict[int, SkillOutput] = {}

    # Group steps by dependency level for sequential batching
    batches = _topological_batches(steps)

    for batch in batches:
        batch_inputs: list[tuple[int, SkillInput]] = []
        for step in batch:
            skill_input = SkillInput({
                "parameters": {**step["parameters"], "_skill_name": step.get("skill_name", "")},
                "objective": objective,
                "stage": str(stage.value if hasattr(stage, "value") else stage),
                "attempt": attempt,
            })
            batch_inputs.append((step["step_id"], skill_input))

        # Run batch concurrently.
        # asyncio.run() raises RuntimeError if called inside a running event
        # loop (e.g. when LangGraph is invoked from an async context).
        # Running it in a ThreadPoolExecutor thread guarantees a fresh loop.
        import concurrent.futures as _cf
        with _cf.ThreadPoolExecutor(max_workers=1) as _pool:
            batch_outputs = _pool.submit(
                asyncio.run, _run_batch_async(_registry, batch_inputs)
            ).result()
        outputs.update(batch_outputs)

        # Log each skill execution to llm_history.yaml
        _history.configure(state.get("workspace_path", ""))
        stage_str_exec = str(stage.value if hasattr(stage, "value") else stage)
        for step in batch:
            sid = step["step_id"]
            out = batch_outputs.get(sid)
            if out is not None:
                _history.log_skill_call(
                    skill=step.get("skill_name", "unknown"),
                    stage=stage_str_exec,
                    attempt=attempt,
                    step_id=step.get("step_id", -1),
                    step_rationale=step.get("rationale", ""),
                    request_sent={
                        k: v for k, v in step.get("parameters", {}).items()
                        if k not in ("content",)  # omit large text blobs
                    },
                    response_received={
                        "confidence": out.get("confidence", 0.0),
                        "sources_count": len(out.get("sources", [])),
                        "sources": out.get("sources", [])[:20],
                        "error": out.get("error"),
                        "result_keys": list(out.get("result", {}).keys()),
                    },
                )

    findings_current: list[SkillOutput] = list(outputs.values())
    existing_findings = state.get("findings", [])
    all_findings = existing_findings + findings_current

    return {
        "findings": all_findings,
        "findings_current": findings_current,
        "status": ResearchStatus.EVALUATING,
    }


def _estimate_reference_contributions(
    llm: "LLMClient",
    state: "ResearchState",
    prev_score: float,
    new_score: float,
    prev_doc: dict,
    current_doc: dict,
) -> dict[str, float]:
    """Estimate the marginal contribution of new references to a score improvement.

    Identifies references present in current_doc but absent from prev_doc,
    then asks the LLM to score each one's contribution to the quality delta.

    Args:
        llm: LLMClient instance.
        state: Current ResearchState (for charter goals and evaluation gaps).
        prev_score: Quality score of the previous best document.
        new_score: Quality score of the current (new best) document.
        prev_doc: Previous best LiteratureReviewDoc.
        current_doc: Current LiteratureReviewDoc.

    Returns:
        dict mapping reference URL → estimated contribution score (0.0–1.0).
    """
    from ai_skill.prompts.literature import build_contribution_messages

    # Identify new references not present in prev_doc
    prev_urls: set[str] = {
        (r.get("url") or "").lower().strip()
        for r in prev_doc.get("references", [])
        if r.get("url")
    }
    new_refs: list[dict] = [
        r for r in current_doc.get("references", [])
        if (r.get("url") or "").lower().strip() not in prev_urls
    ]

    if not new_refs:
        return {}

    # Collect context for the contribution prompt
    objective = state.get("cp2_context") or state.get("objective") or {}
    charter_goals: list[str] = list(objective.get("goals") or [])
    evaluation: dict = state.get("evaluation") or {}
    prev_gaps: list[str] = list(evaluation.get("gaps") or [])

    system, messages = build_contribution_messages(
        charter_goals=charter_goals,
        prev_gaps=prev_gaps,
        new_references=new_refs,
        prev_score=prev_score,
        new_score=new_score,
    )
    try:
        result: _ReferenceContributionLLM = llm.complete_structured(
            messages=messages,
            response_model=_ReferenceContributionLLM,
            system=system,
            temperature=0.0,
            max_tokens=2048,
        )
    except LLMClientError as exc:
        logger.warning("Reference contribution LLM call failed: %s", exc)
        return {}

    return {
        item.url.lower().strip(): float(item.contribution)
        for item in result.contributions
        if item.url
    }


def evaluate(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Evaluate findings against success metrics and decide whether to retry.

    Args:
        state: Current ResearchState with findings and objective.
        llm: Optional injected LLMClient (for testing).

    Returns:
        Partial state update with evaluation result and incremented attempt counter.
    """
    _configure_history(state, "evaluate")
    _llm = llm or LLMClient()
    # cp2_context is the restricted handoff from CP1 (topic, goals, scope_constraints only).
    # Using it here prevents success_metrics from ever reaching the evaluator prompt.
    objective = state.get("cp2_context") or state.get("objective") or ResearchObjective(
        topic="", goals=[], success_metrics=[]
    )
    # Use only the current attempt's findings for evaluation so that poor results
    # from earlier retries do not contaminate scoring of the current attempt.
    findings = state.get("findings_current") or state.get("findings", [])
    attempt = state.get("attempt", 0)
    quality_history = state.get("quality_history", [])
    stage = state.get("stage", PipelineStage.LITERATURE_REVIEW)
    stage_str = str(stage.value if hasattr(stage, "value") else stage)

    # For LITERATURE_REVIEW and RESEARCH_DESIGN, always use hardcoded defaults.
    # See comment in plan() for the full rationale — same root cause applies here.
    if stage_str == PipelineStage.LITERATURE_REVIEW.value:
        stage_guidelines: list[str] | None = _default_literature_review_guidelines()
    elif stage_str == PipelineStage.RESEARCH_DESIGN.value:
        stage_guidelines = _default_research_design_guidelines()
    else:
        raw_guidelines = (objective.get("stage_guidelines") or {}).get(stage_str)
        if raw_guidelines:
            stage_guidelines = raw_guidelines
        else:
            defaults = _default_stage_guidelines(stage_str)
            stage_guidelines = defaults if defaults else None

    system, messages = build_evaluation_messages(
        objective=objective,
        findings=findings,
        threshold=_CONVERGENCE_THRESHOLD,
        stage_guidelines=stage_guidelines,
    )

    try:
        eval_obj: _EvaluationResultLLM = _llm.complete_structured(
            messages=messages,
            response_model=_EvaluationResultLLM,
            system=system,
        )
    except LLMClientError as exc:
        logger.error("Evaluation failed: %s", exc)
        # Fail-safe: treat as not converged with empty gaps
        eval_obj = _EvaluationResultLLM(
            per_metric=[],
            gaps=[f"Evaluation error: {exc}"],
        )

    # Compute total_score and converged deterministically in Python.
    # Never trust the LLM to do arithmetic — it introduces non-determinism.
    per_metric: list[MetricScore] = [
        MetricScore(
            metric=m.get("metric", ""),
            score=float(m.get("score", 0.0)),
            rationale=m.get("rationale", ""),
            gaps=m.get("gaps", []),
        )
        for m in eval_obj.per_metric
    ]
    scores = [m["score"] for m in per_metric]
    total_score = sum(scores) / len(scores) if scores else 0.0
    converged = total_score >= _CONVERGENCE_THRESHOLD

    # Detect regression vs previous attempt
    previous_score = quality_history[-1]["total_score"] if quality_history else None
    regression = (
        previous_score is not None
        and (previous_score - total_score) > 0.05
    )

    evaluation = EvaluationResult(
        per_metric=per_metric,
        total_score=total_score,
        converged=converged,
        gaps=eval_obj.gaps,
        regression=regression,
    )

    # Build quality snapshot
    skills_used = list({
        f.get("skill_name", "") for f in findings if f.get("skill_name")
    })
    snapshot = QualitySnapshot(
        attempt=attempt,
        stage=stage_str,
        total_score=total_score,
        per_metric_scores={m["metric"]: float(m["score"]) for m in eval_obj.per_metric},
        skills_used=skills_used,
        cache_hit_rate=0.0,  # ResearchMemory added in Phase 2
    )

    # Update both the continuous history and the per-stage isolated history
    stage_quality_history: dict[str, list[QualitySnapshot]] = dict(
        state.get("stage_quality_history") or {}
    )
    stage_history_for_stage = list(stage_quality_history.get(stage_str, []))
    stage_history_for_stage.append(snapshot)
    stage_quality_history[stage_str] = stage_history_for_stage

    if regression:
        logger.warning(
            "Quality regression detected: %.2f → %.2f (attempt %d)",
            previous_score,
            total_score,
            attempt,
        )

    # ── CP2 anti-regression + incremental accumulation ────────────────────
    # Only active during LITERATURE_REVIEW stage to avoid interfering with CP3+.
    extra_updates: dict[str, Any] = {}
    if stage_str == PipelineStage.LITERATURE_REVIEW.value:
        best_score = float(state.get("cp2_best_score") or 0.0)
        current_doc = state.get("literature_review_doc")
        prev_best_doc: dict | None = state.get("cp2_best_doc")

        if total_score > best_score and current_doc and current_doc.get("sections"):
            # New best achieved — update best_doc and estimate reference contributions.
            extra_updates["cp2_best_doc"] = current_doc
            extra_updates["cp2_best_score"] = total_score
            logger.info(
                "CP2 new best: %.2f → %.2f (attempt %d)",
                best_score, total_score, attempt,
            )

            # Compute marginal contribution of NEW references (only when improving).
            if best_score > 0.0 and prev_best_doc:
                try:
                    ref_scores = _estimate_reference_contributions(
                        llm=_llm,
                        state=state,
                        prev_score=best_score,
                        new_score=total_score,
                        prev_doc=prev_best_doc,
                        current_doc=current_doc,
                    )
                    merged_scores = dict(state.get("cp2_reference_scores") or {})
                    merged_scores.update(ref_scores)
                    extra_updates["cp2_reference_scores"] = merged_scores
                except Exception as exc:
                    logger.warning("Reference contribution estimation failed: %s", exc)

        elif prev_best_doc and total_score < best_score - 0.01:
            # Score regressed — silently revert to the best known doc so the next
            # compile_literature starts from a better base.
            logger.warning(
                "CP2 regression %.2f → %.2f — reverting literature_review_doc to best (%.2f)",
                previous_score or 0.0, total_score, best_score,
            )
            extra_updates["literature_review_doc"] = prev_best_doc

    return {
        "evaluation": evaluation,
        "attempt": attempt + 1,
        "quality_history": quality_history + [snapshot],
        "stage_quality_history": stage_quality_history,
        "status": ResearchStatus.CONVERGED if converged else ResearchStatus.EXECUTING,
        **extra_updates,
    }


def request_support(state: ResearchState) -> dict[str, Any]:
    """Notify the user that convergence failed and collect revised guidance.

    This node sets checkpoint_pending=True so that the graph pauses via
    interrupt_before and waits for user input before replanning.

    Args:
        state: Current ResearchState with evaluation gaps.

    Returns:
        Partial state update requesting user feedback.
    """
    evaluation = state.get("evaluation") or {}
    gaps = evaluation.get("gaps", [])

    logger.warning(
        "Pipeline could not converge after 5 attempts. Requesting user support."
    )

    gaps_display = "\n".join(f"  - {g}" for g in gaps) if gaps else "  (no specific gaps identified)"
    message = (
        f"After 5 attempts, the research pipeline could not meet all success metrics.\n\n"
        f"Unresolved gaps:\n{gaps_display}\n\n"
        f"Please provide guidance to help the agent replan."
    )

    workspace_path = state.get("workspace_path", "")
    if workspace_path:
        ResearchWorkspace(Path(workspace_path)).log(
            f"User support requested. Gaps: {gaps}"
        )

    return {
        "status": ResearchStatus.CHECKPOINT_PENDING,
        "checkpoint_pending": True,
        "checkpoint_label": (
            f"SUPPORT NEEDED: {message}"
        ),
        "attempt": 0,  # Reset attempt counter for replanning
    }


# ---------------------------------------------------------------------------
# Checkpoint 2 — Literature Review nodes
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# CP2 chunked compilation helpers
# ---------------------------------------------------------------------------

import json


def _canonical_ref_url(ref: dict) -> str:
    """Return a canonical deduplication key for a reference dict."""
    url = (ref.get("url") or "").strip().rstrip("/")
    if url:
        return url.lower()
    return (ref.get("title") or "").lower().strip()


def _select_chunk_findings(
    findings: list[dict],
    themes: list[str],
    used_limit: int = _CHUNK_FINDINGS_SIZE,
) -> list[dict]:
    """Select the most theme-relevant findings for a chunk."""
    if not themes:
        return findings[:used_limit]

    def _relevance(f: dict) -> int:
        text = json.dumps(f, ensure_ascii=False).lower()
        return sum(1 for kw in themes if kw.lower() in text)

    ranked = sorted(findings, key=_relevance, reverse=True)
    return ranked[:used_limit]


def _compile_literature_chunked(
    llm: "LLMClient",
    charter_document_text: str,
    findings_summary: list[dict],
    best_doc: dict | None = None,
) -> "_LiteratureReviewLLM":
    """Compile a literature review in multiple chunks to avoid max_tokens limits.

    Strategy:
    1. Outline pass (1 call): plan N section groups with themes.
    2. Chunk pass (N calls): generate 2–3 sections per group with offset-numbered refs.
    3. Python merge: concatenate sections, deduplicate references.

    Args:
        llm: LLMClient instance.
        charter_document_text: CP1 charter text.
        findings_summary: Pre-built findings summary list.
        best_doc: Previous best LiteratureReviewDoc (used as seed context if provided).

    Returns:
        A _LiteratureReviewLLM instance ready for post-processing.
    """
    from ai_skill.prompts.literature import (
        build_outline_messages,
        build_chunk_messages,
    )

    num_chunks = max(2, (len(findings_summary) + _CHUNK_FINDINGS_SIZE - 1) // _CHUNK_FINDINGS_SIZE)

    # ── Phase 1: Outline ──────────────────────────────────────────────────
    finding_titles: list[str] = []
    for f in findings_summary:
        for p in (f.get("papers") or []):
            title = p.get("title") or ""
            if title:
                finding_titles.append(title)
        if f.get("skill_name") == "content_summarizer":
            finding_titles.append((f.get("summary") or "")[:80])

    system, messages = build_outline_messages(
        charter_document_text=charter_document_text,
        finding_titles=finding_titles,
        num_chunks=num_chunks,
    )
    try:
        outline: _OutlineLLM = llm.complete_structured(
            messages=messages,
            response_model=_OutlineLLM,
            system=system,
            max_tokens=2048,
        )
    except LLMClientError as exc:
        logger.warning("Outline phase failed (%s) — falling back to even split", exc)
        # Fallback: create equal groups without thematic assignment
        section_titles_fallback = [
            f"Seção {i + 1}" for i in range(num_chunks * _CHUNK_SECTIONS_PER_GROUP)
        ]
        groups = [
            _SectionGroupLLM(
                titles=section_titles_fallback[i : i + _CHUNK_SECTIONS_PER_GROUP],
                themes=[],
            )
            for i in range(0, len(section_titles_fallback), _CHUNK_SECTIONS_PER_GROUP)
        ]
        outline = _OutlineLLM(section_groups=groups)

    # If best_doc has refs, start global_refs with them so chunks can cite them
    global_refs: list[dict] = []  # list of reference dicts (already numbered)
    seen_url_keys: set[str] = set()

    if best_doc:
        for ref in best_doc.get("references", []):
            key = _canonical_ref_url(ref)
            if key and key not in seen_url_keys:
                seen_url_keys.add(key)
                global_refs.append(ref)

    all_sections: list[_LiteratureReviewSectionLLM] = []

    # ── Phase 2: Chunk generation ─────────────────────────────────────────
    for group in outline.section_groups:
        chunk_findings = _select_chunk_findings(findings_summary, group.themes)
        ref_offset = len(global_refs) + 1

        chunk_system, chunk_messages = build_chunk_messages(
            charter_document_text=charter_document_text,
            section_titles=group.titles,
            findings=chunk_findings,
            already_cited=global_refs,
            ref_offset=ref_offset,
        )
        try:
            chunk: _LiteratureChunkLLM = llm.complete_structured(
                messages=chunk_messages,
                response_model=_LiteratureChunkLLM,
                system=chunk_system,
                max_tokens=16384,
            )
        except LLMClientError as exc:
            logger.warning(
                "Chunk '%s' failed (%s) — skipping",
                ", ".join(group.titles),
                exc,
            )
            continue

        all_sections.extend(chunk.sections)

        # Merge new references, deduplicating by canonical URL
        for ref in chunk.new_references:
            key = _canonical_ref_url({"url": ref.url, "title": ref.title})
            if key and key not in seen_url_keys:
                seen_url_keys.add(key)
                # Convert _LiteratureSourceLLM → dict for global_refs
                global_refs.append({
                    "reference_number": ref.reference_number,
                    "title": ref.title,
                    "authors": ref.authors,
                    "year": ref.year,
                    "url": ref.url,
                    "abnt_entry": ref.abnt_entry,
                    "summary": ref.summary,
                })

    if not all_sections:
        raise LLMClientError("Chunked compilation produced no sections.")

    # ── Phase 3: Rebuild final _LiteratureReviewLLM ───────────────────────
    final_sources = [
        _LiteratureSourceLLM(
            reference_number=r.get("reference_number", i + 1),
            title=r.get("title", ""),
            authors=r.get("authors", ""),
            year=str(r.get("year", "")),
            url=r.get("url", ""),
            abnt_entry=r.get("abnt_entry", ""),
            summary=r.get("summary", ""),
        )
        for i, r in enumerate(global_refs)
    ]

    return _LiteratureReviewLLM(sections=all_sections, references=final_sources)


def compile_literature(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Compile a structured literature review from the accumulated findings.

    Calls the LLM with a fixed output schema (_LiteratureReviewLLM) to ensure
    a consistent format across multiple runs.

    Args:
        state: Current ResearchState with findings and objective.
        llm: Optional injected LLMClient (for testing).

    Returns:
        Partial state update with ``literature_review_doc`` and
        ``active_checkpoint=2``.
    """
    _configure_history(state, "compile_literature")
    _llm = llm or LLMClient()
    findings = state.get("findings", [])
    charter_document_text = state.get("charter_document_text", "")

    # Sort by confidence (desc) and take top 40 to prevent token overflow.
    # Uses all accumulated findings (not just current attempt) for richer compilation.
    # Build reference scores index from cp2_reference_scores (url → contribution score).
    # This boosts findings that previously produced high-contribution references so that
    # incremental iterations keep improving on the best result.
    ref_scores: dict[str, float] = dict(state.get("cp2_reference_scores") or {})
    best_doc_refs: set[str] = set()
    best_doc: dict | None = state.get("cp2_best_doc")
    if best_doc:
        for r in best_doc.get("references", []):
            u = (r.get("url") or "").strip().lower()
            if u:
                best_doc_refs.add(u)

    def _finding_priority(f: dict) -> float:
        """Combine confidence + max marginal contribution score for contained papers."""
        base = float(f.get("confidence", 0.0))
        bonus = 0.0
        for p in (f.get("papers") or []):
            url = (p.get("url") or p.get("doi") or "").lower()
            if url in ref_scores:
                bonus = max(bonus, ref_scores[url])
        return base + bonus * 0.5  # blend: confidence + half the contribution bonus

    sorted_findings = sorted(findings, key=_finding_priority, reverse=True)
    top_findings = sorted_findings[:40]

    findings_summary: list[dict[str, Any]] = []
    for f in top_findings:
        skill_name = f.get("skill_name", "")
        result = f.get("result") or {}
        entry: dict[str, Any] = {
            "skill": skill_name,
            "confidence": f.get("confidence", 0.0),
            "sources": f.get("sources", [])[:5],
        }
        if skill_name == "article_search":
            papers = result.get("papers", [])
            entry["papers"] = [
                {
                    "title": p.get("title", ""),
                    "abstract": (p.get("abstract", "") or "")[:300],
                    "url": p.get("url", ""),
                    "doi": p.get("doi") or "",           # prefer doi.org URL for verify
                    "arxiv_id": p.get("arxiv_id") or "", # prefer arxiv.org URL for verify
                    "year": p.get("year"),
                    "authors": p.get("authors", [])[:3],
                    "venue": p.get("venue", ""),
                }
                for p in papers[:15]
            ]
        elif skill_name == "content_summarizer":
            entry["summary"] = (result.get("summary", "") or "")[:500]
            entry["key_points"] = result.get("key_points", [])[:5]
        else:
            for k, v in list(result.items())[:5]:
                if isinstance(v, str):
                    entry[k] = v[:400]
                elif isinstance(v, list):
                    entry[k] = v[:10]
                else:
                    entry[k] = v
        findings_summary.append(entry)

    system, messages = build_compile_messages(
        charter_document_text=charter_document_text,
        findings=findings_summary,
    )

    # Attempt 1: single call with full max_tokens ceiling (64 000 is the API max for this model).
    # If the response is truncated (max_tokens hit), fall back to chunked compilation.
    review_obj: _LiteratureReviewLLM | None = None
    try:
        review_obj = _llm.complete_structured(
            messages=messages,
            response_model=_LiteratureReviewLLM,
            system=system,
            temperature=0.3,
            max_tokens=64000,
        )
    except LLMClientError as exc:
        if "max_tokens" in str(exc).lower() or "truncat" in str(exc).lower():
            logger.warning(
                "Single compilation hit max_tokens — switching to chunked mode (%d findings)",
                len(top_findings),
            )
            try:
                review_obj = _compile_literature_chunked(
                    llm=_llm,
                    charter_document_text=charter_document_text,
                    findings_summary=findings_summary,
                    best_doc=state.get("cp2_best_doc"),
                )
            except LLMClientError as chunk_exc:
                logger.error("Chunked compilation also failed: %s", chunk_exc)
                return {
                    "status": ResearchStatus.FAILED,
                    "active_checkpoint": 2,
                    "literature_review_doc": LiteratureReviewDoc(sections=[], references=[], verified_sources=[]),
                }
        else:
            logger.error("Literature review compilation failed: %s", exc)
            return {
                "status": ResearchStatus.FAILED,
                "active_checkpoint": 2,
                "literature_review_doc": LiteratureReviewDoc(sections=[], references=[], verified_sources=[]),
            }

    sections: list[LiteratureReviewSection] = [
        LiteratureReviewSection(section_title=s.section_title, content=s.content)
        for s in review_obj.sections
    ]
    # Build a title → finding lookup so we can repair LLM-generated URLs
    _finding_by_title: dict[str, dict[str, Any]] = {}
    for f in top_findings:
        if f.get("skill_name") == "article_search":
            for p in (f.get("result") or {}).get("papers", []):
                t = (p.get("title") or "").lower().strip()
                if t:
                    _finding_by_title[t] = p

    def _best_url(llm_url: str, title: str) -> str:
        """Return the most API-resolvable URL for a reference."""
        # If the LLM URL already looks like a DOI or arXiv link, keep it
        if any(kw in llm_url for kw in ("doi.org", "arxiv.org", "semanticscholar.org")):
            return llm_url
        # Try to find the paper in findings by title and use its DOI/arXiv/S2 URL
        finding = _finding_by_title.get(title.lower().strip())
        if finding:
            if finding.get("doi"):
                return f"https://doi.org/{finding['doi']}"
            if finding.get("arxiv_id"):
                return f"https://arxiv.org/abs/{finding['arxiv_id']}"
            if finding.get("url"):
                return finding["url"]
        return llm_url  # keep as-is — HTTP fallback will try it

    references: list[dict[str, Any]] = [
        {
            "reference_number": r.reference_number,
            "title": r.title,
            "authors": r.authors,
            "year": r.year,
            "url": _best_url(r.url or "", r.title or ""),
            "abnt_entry": r.abnt_entry,
            "summary": r.summary,
        }
        for r in review_obj.references
    ]

    literature_review_doc = LiteratureReviewDoc(
        sections=sections,
        references=references,
        verified_sources=[],
    )

    workspace_path = state.get("workspace_path", "")
    if workspace_path:
        ResearchWorkspace(Path(workspace_path)).log(
            f"Literature review compiled: {len(sections)} sections, "
            f"{len(references)} references."
        )

    return {
        "literature_review_doc": literature_review_doc,
        "active_checkpoint": 2,
        "status": ResearchStatus.EXECUTING,
    }


def verify_literature(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Verify each bibliographic source with an independent LLM agent.

    For every reference in ``literature_review_doc``:
    1. Attempt to fetch the URL via HTTP.
    2. Send fetched content + claim to a separate LLM call (temperature=0)
       to check factual consistency.
    3. Attach a :class:`SourceVerification` record to the document.

    This node is intentionally separate from ``compile_literature`` so that
    the verification agent cannot rationalise its own previous claims.

    Smart re-verification rules (correction cycle only — when previous
    ``verified_sources`` already exist):

    - ✓ Green tick (accessible + content_matches): skip unless the reference
      is explicitly mentioned in ``user_feedback`` (comment context).
    - ⚠ Warning (accessible, content divergence): skip unless the reference
      appears in ``user_feedback`` (user highlighted/commented that section).
    - ✗ Red X (inaccessible): retry if *any* feedback/marks exist in the
      document (``user_feedback`` non-empty); skip if there are no marks at all.
    - New reference (not in previous verified_sources): always verify.

    Args:
        state: Current ResearchState with ``literature_review_doc``.
        llm: Optional injected LLMClient (for testing).

    Returns:
        Partial state update with verified ``literature_review_doc``.
    """
    import re as _re

    _configure_history(state, "verify_literature")
    _llm = llm or LLMClient()
    review_doc: dict[str, Any] = dict(state.get("literature_review_doc") or {})
    references: list[dict[str, Any]] = list(review_doc.get("references", []))

    if not references:
        logger.warning("verify_literature: no references to verify.")
        return {"literature_review_doc": LiteratureReviewDoc(**review_doc)}

    access_date = _abnt_access_date()

    # Replace ACCESS_DATE placeholder in all references upfront
    for ref in references:
        abnt_entry = ref.get("abnt_entry", "")
        ref["abnt_entry"] = abnt_entry.replace("{ACCESS_DATE}", access_date)

    # Build lookup of previously verified sources (empty on initial pass)
    prev_verified: list[dict[str, Any]] = list(review_doc.get("verified_sources") or [])
    prev_by_num: dict[int, dict[str, Any]] = {
        v.get("reference_number", 0): dict(v) for v in prev_verified
    }
    is_correction_cycle = bool(prev_by_num)

    # Reference numbers explicitly mentioned in feedback (comment context)
    user_feedback: str = state.get("user_feedback") or ""
    feedback_ref_nums: set[int] = set()
    if user_feedback:
        for m in _re.findall(r"\[(\d+)\]", user_feedback):
            feedback_ref_nums.add(int(m))
        for m in _re.findall(
            r"(?:fonte|ref(?:er[eê]ncia)?|source|recheck|verificar|checar)\s+(\d+)",
            user_feedback,
            _re.IGNORECASE,
        ):
            feedback_ref_nums.add(int(m))

    def _verify_one(ref: dict[str, Any]) -> SourceVerification:
        """Verify a single reference — runs in a thread pool worker."""
        url = ref.get("url", "")
        ref_num = ref.get("reference_number", 0)
        title = ref.get("title", "")
        summary = ref.get("summary", "")

        # --- Smart skip logic (correction cycle only) ---
        if is_correction_cycle:
            prev = prev_by_num.get(ref_num)
            in_comment = ref_num in feedback_ref_nums

            if prev is not None and not in_comment:
                was_accessible = prev.get("accessible", False)
                was_match = prev.get("content_matches", False)

                # ✓ Green tick: skip — already verified correctly
                if was_accessible and was_match:
                    logger.debug(
                        "verify_literature: [%d] ✓ already verified — skipping.", ref_num
                    )
                    return SourceVerification(**prev)

                # ⚠ Warning: always re-verify on correction cycles — the claim may
                # have been updated by refine_literature (which clears user_feedback
                # before verify runs, so feedback_ref_nums is unreliable here).
                if was_accessible and not was_match:
                    logger.debug(
                        "verify_literature: [%d] ⚠ correction cycle — re-verifying "
                        "(claim may have changed).",
                        ref_num,
                    )
                    # Fall through to actual fetch + LLM verification

                # ✗ Red X: always retry in correction cycles — a previous failure may
                # have been caused by a transient error (rate-limit, timeout, etc.)
                # rather than genuine inaccessibility.
                if not was_accessible:
                    logger.debug(
                        "verify_literature: [%d] ✗ correction cycle — retrying.", ref_num
                    )

        # --- Perform actual fetch + LLM verification ---
        accessible, fetched_content, access_method = _fetch_url_content(url)

        content_matches = False
        note = "URL não acessível — verificação automática não foi possível."

        if accessible:
            system, messages = build_verify_messages(
                reference_number=ref_num,
                title=title,
                summary=summary,
                url=url,
                fetched_content=fetched_content,
            )
            try:
                v_result: _VerifyResultLLM = _llm.complete_structured(
                    messages=messages,
                    response_model=_VerifyResultLLM,
                    system=system,
                    temperature=0.0,
                )
                content_matches = v_result.content_matches
                note = v_result.verification_note
            except LLMClientError as exc:
                logger.warning(
                    "Verification LLM call failed for ref [%d]: %s", ref_num, exc
                )
                note = f"Verificação automática falhou: {exc}"

        return SourceVerification(
            reference_number=ref_num,
            url=url,
            title=title,
            accessible=accessible,
            content_matches=content_matches,
            verification_note=note,
            access_date=access_date,
            access_method=access_method,
        )

    verified: list[SourceVerification] = []
    max_workers = min(5, max(1, len(references)))
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(_verify_one, ref): ref for ref in references}
        for future in as_completed(futures):
            try:
                verified.append(future.result())
            except Exception as exc:
                ref = futures[future]
                logger.warning("Verification failed for ref [%s]: %s", ref.get("url", "?"), exc)

    review_doc["references"] = references
    review_doc["verified_sources"] = verified  # type: ignore[assignment]

    workspace_path = state.get("workspace_path", "")
    if workspace_path:
        ok = sum(1 for v in verified if v.get("content_matches"))
        warn = sum(1 for v in verified if v.get("accessible") and not v.get("content_matches"))
        fail = sum(1 for v in verified if not v.get("accessible"))
        ResearchWorkspace(Path(workspace_path)).log(
            f"Source verification: ✓ {ok} verified, ⚠ {warn} questionable, ✗ {fail} inaccessible."
        )

    return {"literature_review_doc": LiteratureReviewDoc(**review_doc)}


def deliver_literature(state: ResearchState) -> dict[str, Any]:
    """Convert the verified literature review to a .docx and save Checkpoint 2 preview.

    Renders each section with its content, then a bibliography section with
    coloured verification marks (✓ green / ⚠ orange / ✗ red) beside each entry.

    Args:
        state: Current ResearchState with a verified ``literature_review_doc``.

    Returns:
        Partial state update with ``checkpoint_label`` pointing to the saved preview.
    """
    review_doc: dict[str, Any] = state.get("literature_review_doc") or {}
    workspace_path = state.get("workspace_path", "")
    checkpoint_docx_path = ""

    if not review_doc.get("sections"):
        logger.warning("deliver_literature: empty review doc — skipping docx generation.")
        return {"checkpoint_label": "", "literature_approved": False, "status": ResearchStatus.FAILED, "active_checkpoint": 2}

    if workspace_path:
        docx_bytes = _literature_review_to_docx(review_doc)
        if docx_bytes:
            project_ws = _get_project_workspace(workspace_path)
            if project_ws is not None:
                saved = project_ws.save_checkpoint_preview(2, docx_bytes)
                checkpoint_docx_path = str(saved)
                logger.info("Checkpoint 2 preview saved: %s", saved)
        ResearchWorkspace(Path(workspace_path)).log("Literature Review preview saved.")

    return {
        "checkpoint_label": checkpoint_docx_path,
        "literature_approved": False,
        "status": ResearchStatus.PLANNING,
    }


def review_literature(state: ResearchState) -> dict[str, Any]:
    """Checkpoint gate executed after the user resumes from the literature review.

    The graph pauses (``interrupt_before``) before this node.  When the user
    resumes with ``--correct``, formatted feedback lands in ``user_feedback``
    and this node signals that a revision cycle is needed.  Without feedback
    the review is approved and the pipeline ends.

    Args:
        state: Current ResearchState. ``user_feedback`` carries revision
            instructions if the user requested changes.

    Returns:
        Partial state update with ``literature_approved`` set accordingly.
    """
    user_feedback = state.get("user_feedback")
    if user_feedback:
        logger.info("Literature review revision requested.")
        return {"literature_approved": False, "status": ResearchStatus.PLANNING}
    logger.info("Literature Review approved by user.")
    return {"literature_approved": True, "status": ResearchStatus.COMPLETED}


def refine_literature(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Refine the literature review based on researcher corrections from the .docx.

    Args:
        state: Current ResearchState with ``user_feedback`` and
            ``literature_review_doc``.
        llm: Optional injected LLMClient (for testing).

    Returns:
        Partial state update with refined ``literature_review_doc`` and
        cleared ``user_feedback``.
    """
    _configure_history(state, "refine_literature")
    _llm = llm or LLMClient()
    review_doc: dict[str, Any] = state.get("literature_review_doc") or {}
    user_feedback = state.get("user_feedback", "")

    system, messages = build_refine_messages(
        review_doc=review_doc,
        feedback=user_feedback or "",
    )

    try:
        refined_obj: _LiteratureReviewLLM = _llm.complete_structured(
            messages=messages,
            response_model=_LiteratureReviewLLM,
            system=system,
            temperature=0.3,
            max_tokens=32768,
        )
    except LLMClientError as exc:
        logger.error("Literature review refinement failed: %s", exc)
        return {"status": ResearchStatus.FAILED}

    sections: list[LiteratureReviewSection] = [
        LiteratureReviewSection(section_title=s.section_title, content=s.content)
        for s in refined_obj.sections
    ]
    references: list[dict[str, Any]] = [
        {
            "reference_number": r.reference_number,
            "title": r.title,
            "authors": r.authors,
            "year": r.year,
            "url": r.url,
            "abnt_entry": r.abnt_entry,
            "summary": r.summary,
        }
        for r in refined_obj.references
    ]

    # Preserve existing verified_sources so verify_literature can apply
    # smart skip logic (green ticks are not re-fetched unnecessarily).
    existing_verified = list(review_doc.get("verified_sources") or [])
    updated_doc = LiteratureReviewDoc(
        sections=sections,
        references=references,
        verified_sources=existing_verified,
    )

    workspace_path = state.get("workspace_path", "")
    if workspace_path:
        ResearchWorkspace(Path(workspace_path)).log(
            "Literature review refined based on researcher corrections."
        )

    return {
        "literature_review_doc": updated_doc,
        "user_feedback": None,
        "status": ResearchStatus.EXECUTING,
    }


def route_after_review_literature(state: ResearchState) -> str:
    """Conditional edge: decide next node after literature review gate.

    Args:
        state: Current ResearchState after ``review_literature`` ran.

    Returns:
        ``"END"`` if approved, ``"recheck_sources"`` if revision needed
        (recheck_sources is a no-op when feedback contains no recheck instructions,
        then always continues to refine_literature).
    """
    return "END" if state.get("literature_approved") else "recheck_sources"


def recheck_sources(state: ResearchState) -> dict[str, Any]:
    """Re-fetch specific sources via Semantic Scholar API when the user requests a recheck.

    Parses ``user_feedback`` for reference numbers using patterns such as
    ``[N]``, ``fonte N``, ``referência N``, or ``recheck N``.  For each
    mentioned reference, calls ``_fetch_via_semantic_scholar_api`` to obtain
    fresh metadata (abstract, authors, year) and updates the corresponding
    entry in ``verified_sources``.

    This node is always called in the revision path (even without explicit recheck
    instructions), acting as a no-op when no reference numbers are found.
    ``refine_literature`` runs next to apply any prose corrections.

    Args:
        state: Current ResearchState with ``user_feedback`` and
            ``literature_review_doc``.

    Returns:
        Partial state update with ``literature_review_doc`` updated (verified_sources
        refreshed for mentioned references), or ``{}`` if no references were targeted.
    """
    import re

    feedback = state.get("user_feedback") or ""
    review_doc: dict[str, Any] = dict(state.get("literature_review_doc") or {})
    references: list[dict[str, Any]] = list(review_doc.get("references", []))
    verified: list[dict[str, Any]] = list(review_doc.get("verified_sources", []))

    # Parse reference numbers from feedback
    ref_nums: set[int] = set()
    for m in re.findall(r"\[(\d+)\]", feedback):
        ref_nums.add(int(m))
    for m in re.findall(
        r"(?:fonte|ref(?:er[eê]ncia)?|source|recheck|verificar|checar)\s+(\d+)",
        feedback,
        re.IGNORECASE,
    ):
        ref_nums.add(int(m))

    if not ref_nums:
        logger.info("recheck_sources: no reference numbers found in feedback — skipping.")
        return {}

    logger.info("recheck_sources: rechecking references %s", sorted(ref_nums))
    refs_by_num: dict[int, dict[str, Any]] = {
        r.get("reference_number", 0): r for r in references
    }
    verified_by_num: dict[int, dict[str, Any]] = {
        v.get("reference_number", 0): dict(v) for v in verified
    }
    access_date = _abnt_access_date()

    for ref_num in ref_nums:
        ref = refs_by_num.get(ref_num)
        if ref is None:
            logger.debug("recheck_sources: reference [%d] not found in review doc.", ref_num)
            continue

        url = ref.get("url", "")
        result = _fetch_via_semantic_scholar_api(url)
        if result is not None:
            accessible, content, method = result
        else:
            accessible, content, method = _fetch_url_content(url)

        existing = verified_by_num.get(ref_num, {})
        existing.update({
            "reference_number": ref_num,
            "url": url,
            "title": ref.get("title", existing.get("title", "")),
            "accessible": accessible,
            "content_matches": existing.get("content_matches", False),  # preserve LLM result
            "verification_note": (
                f"Rechecado via {method}. "
                + (existing.get("verification_note") or "")
            )[:500],
            "access_date": access_date,
            "access_method": method,
        })
        verified_by_num[ref_num] = existing
        logger.info(
            "recheck_sources: [%d] %s — accessible=%s method=%s",
            ref_num, url[:60], accessible, method,
        )

    review_doc["verified_sources"] = list(verified_by_num.values())
    workspace_path = state.get("workspace_path", "")
    if workspace_path:
        ResearchWorkspace(Path(workspace_path)).log(
            f"Source recheck: {len(ref_nums)} reference(s) rechecked via API."
        )

    return {"literature_review_doc": LiteratureReviewDoc(**review_doc)}


# ---------------------------------------------------------------------------
# CP3 — Research Design nodes
# ---------------------------------------------------------------------------


def cp3_router(_state: ResearchState) -> dict[str, Any]:
    """Entry-point no-op for CP3; routing logic lives in route_cp3_start.

    Args:
        state: Current ResearchState (unused — routing done in conditional edge).

    Returns:
        Empty dict (no state changes).
    """
    return {}


def route_cp3_start(state: ResearchState) -> str:
    """Conditional edge from cp3_router: fresh start vs. correction cycle.

    Args:
        state: Current ResearchState.

    Returns:
        ``"refine_design"`` when a design doc and user feedback already exist
        (researcher correction cycle); ``"read_attachments"`` otherwise
        (fresh CP3 run — Phase 1 starts).
    """
    if state.get("research_design_doc") and state.get("user_feedback"):
        return "refine_design"
    return "read_attachments"


def compile_design(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Synthesise methodology research findings into the Research Design document.

    Called when the evaluate node converges. Reads the top-40 findings plus
    the CP1 and CP2 approved documents, then calls the LLM to produce a fully
    structured ``ResearchDesignDoc`` covering all 9 mandatory sections.

    Args:
        state: Current ResearchState with converged findings.
        llm: Optional LLMClient override (uses default when None).

    Returns:
        Partial state update with ``research_design_doc`` and ``active_checkpoint=3``.
    """
    _configure_history(state, "compile_design")
    _llm = llm or LLMClient()

    workspace_path = state.get("workspace_path", "")

    # Build context from CP1 and CP2 approved documents
    charter_text = state.get("charter_document_text") or ""

    # Extract CP2 text from cp3_context if available, else stringify the doc
    cp3_context = state.get("cp3_context") or {}
    literature_text = cp3_context.get("literature_summary") or ""
    if not literature_text:
        review_doc = state.get("literature_review_doc") or {}
        sections = review_doc.get("sections", [])
        literature_text = "\n\n".join(
            f"## {s.get('section_title', '')}\n{s.get('content', '')}"
            for s in sections
        )

    # Take top 40 findings by confidence
    all_findings = state.get("findings") or []
    sorted_findings = sorted(
        all_findings,
        key=lambda f: float(f.get("confidence", 0.0)),
        reverse=True,
    )
    top_findings = sorted_findings[:40]

    # Build a compact representation for each finding
    compact_findings: list[dict[str, Any]] = []
    for f in top_findings:
        result = f.get("result") or {}
        skill = f.get("skill_name", "")
        summary: dict[str, Any] = {"skill": skill, "confidence": f.get("confidence", 0.0)}
        if skill == "article_search":
            papers = result.get("papers", [])[:5]
            summary["papers"] = [
                {"title": p.get("title"), "year": p.get("year"), "abstract": (p.get("abstract") or "")[:300]}
                for p in papers
            ]
        elif skill == "content_summarizer":
            summary["summary"] = (result.get("summary") or "")[:400]
            summary["key_points"] = result.get("key_points", [])[:5]
        else:
            # Generic: first 5 keys truncated
            for k, v in list(result.items())[:5]:
                summary[k] = str(v)[:400] if isinstance(v, str) else v
        compact_findings.append(summary)

    system, messages = build_compile_design_messages(
        charter_document_text=charter_text,
        literature_document_text=literature_text,
        cp3_context=cp3_context,
        findings=compact_findings,
    )

    try:
        design_obj: _ResearchDesignDocLLM = _llm.complete_structured(
            messages=messages,
            response_model=_ResearchDesignDocLLM,
            system=system,
            max_tokens=32768,
        )
    except LLMClientError as exc:
        logger.error("compile_design failed: %s", exc)
        return {"status": ResearchStatus.FAILED}

    sections: list[ResearchDesignSection] = [
        ResearchDesignSection(
            section_title=s.section_title,
            content=s.content,
        )
        for s in design_obj.sections
    ]

    design_doc = ResearchDesignDoc(
        sections=sections,
        study_type=design_obj.study_type,
        research_paradigm=design_obj.research_paradigm,
        epistemological_stance=design_obj.epistemological_stance,
        hypotheses=design_obj.hypotheses,
        research_questions=design_obj.research_questions,
        variables=[v.model_dump() for v in design_obj.variables],
        instruments=design_obj.instruments,
        sampling_strategy=design_obj.sampling_strategy,
        sample_size_justification=design_obj.sample_size_justification,
        ethical_considerations=design_obj.ethical_considerations,
        validity_threats=design_obj.validity_threats,
        mitigation_strategies=[m.model_dump() for m in design_obj.mitigation_strategies],
        data_management_plan=design_obj.data_management_plan,
        metrics_and_kpis=design_obj.metrics_and_kpis,
        data_sources=design_obj.data_sources,
        collection_protocol=design_obj.collection_protocol,
        methodology_timeline=design_obj.methodology_timeline,
        reporting_standard=design_obj.reporting_standard,
        target_journal_tier=design_obj.target_journal_tier,
    )

    if workspace_path:
        ResearchWorkspace(Path(workspace_path)).log(
            f"Research design compiled: {len(sections)} sections, "
            f"{len(design_obj.hypotheses)} hypotheses, "
            f"{len(design_obj.variables)} variables."
        )

    return {
        "research_design_doc": design_doc,
        "active_checkpoint": 3,
        "status": ResearchStatus.EXECUTING,
    }


def deliver_design(state: ResearchState) -> dict[str, Any]:
    """Generate the Research Design .docx checkpoint preview.

    Converts the ``research_design_doc`` to a formatted Word document and
    saves it as ``Checkpoint 3 - Research Design [preview_N].docx``.

    Args:
        state: Current ResearchState with a populated ``research_design_doc``.

    Returns:
        Partial state update with checkpoint label and ``design_approved=False``.
    """
    _configure_history(state, "deliver_design")

    design_doc = state.get("research_design_doc") or {}
    workspace_path = state.get("workspace_path", "")

    docx_bytes = _research_design_to_docx(design_doc)

    checkpoint_label = ""
    if workspace_path:
        pw = _get_project_workspace(workspace_path)
        if pw and docx_bytes:
            checkpoint_path = pw.save_checkpoint_preview(3, docx_bytes)
            checkpoint_label = str(checkpoint_path)
            logger.info("Research design checkpoint saved: %s", checkpoint_label)

    return {
        "checkpoint_label": checkpoint_label,
        "design_approved": False,
        "status": ResearchStatus.PLANNING,
    }


def review_design(state: ResearchState) -> dict[str, Any]:
    """Gate node: check whether the researcher approved the Research Design.

    This node runs after the ``interrupt_before`` checkpoint fires and the user
    resumes the graph. If ``user_feedback`` is present, the design needs
    revision; otherwise it is considered approved.

    Args:
        state: Current ResearchState, resumed by the user.

    Returns:
        Partial state update setting ``design_approved`` and ``status``.
    """
    _configure_history(state, "review_design")

    if state.get("user_feedback"):
        return {
            "design_approved": False,
            "status": ResearchStatus.PLANNING,
        }
    return {
        "design_approved": True,
        "status": ResearchStatus.COMPLETED,
    }


def refine_design(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Apply researcher corrections to the Research Design document.

    Called when the user provides feedback at the Checkpoint 3 review gate.
    Applies surgical corrections while preserving all unmarked content.

    Args:
        state: Current ResearchState with ``research_design_doc`` and ``user_feedback``.
        llm: Optional LLMClient override.

    Returns:
        Partial state update with the refined ``research_design_doc`` and
        ``user_feedback=None``.
    """
    _configure_history(state, "refine_design")
    _llm = llm or LLMClient()

    design_doc = state.get("research_design_doc") or {}
    feedback = state.get("user_feedback") or ""
    workspace_path = state.get("workspace_path", "")

    system, messages = build_refine_design_messages(
        design_doc=dict(design_doc),
        feedback=feedback,
    )

    try:
        design_obj: _ResearchDesignDocLLM = _llm.complete_structured(
            messages=messages,
            response_model=_ResearchDesignDocLLM,
            system=system,
            max_tokens=32768,
        )
    except LLMClientError as exc:
        logger.error("refine_design failed: %s", exc)
        return {"status": ResearchStatus.FAILED}

    sections: list[ResearchDesignSection] = [
        ResearchDesignSection(
            section_title=s.section_title,
            content=s.content,
        )
        for s in design_obj.sections
    ]

    updated_doc = ResearchDesignDoc(
        sections=sections,
        study_type=design_obj.study_type,
        research_paradigm=design_obj.research_paradigm,
        epistemological_stance=design_obj.epistemological_stance,
        hypotheses=design_obj.hypotheses,
        research_questions=design_obj.research_questions,
        variables=[v.model_dump() for v in design_obj.variables],
        instruments=design_obj.instruments,
        sampling_strategy=design_obj.sampling_strategy,
        sample_size_justification=design_obj.sample_size_justification,
        ethical_considerations=design_obj.ethical_considerations,
        validity_threats=design_obj.validity_threats,
        mitigation_strategies=[m.model_dump() for m in design_obj.mitigation_strategies],
        data_management_plan=design_obj.data_management_plan,
        metrics_and_kpis=design_obj.metrics_and_kpis,
        data_sources=design_obj.data_sources,
        collection_protocol=design_obj.collection_protocol,
        methodology_timeline=design_obj.methodology_timeline,
        reporting_standard=design_obj.reporting_standard,
        target_journal_tier=design_obj.target_journal_tier,
    )

    if workspace_path:
        ResearchWorkspace(Path(workspace_path)).log(
            "Research design refined based on researcher corrections."
        )

    return {
        "research_design_doc": updated_doc,
        "user_feedback": None,
        "status": ResearchStatus.EXECUTING,
    }


def route_after_review_design(state: ResearchState) -> str:
    """Conditional edge: decide next node after research design review gate.

    Args:
        state: Current ResearchState after ``review_design`` ran.

    Returns:
        ``"END"`` if approved, ``"refine_design"`` if revision needed.
    """
    return "END" if state.get("design_approved") else "refine_design"


# ---------------------------------------------------------------------------
# Edge routing functions
# ---------------------------------------------------------------------------

def route_after_evaluate(state: ResearchState) -> str:
    """Conditional edge: decide next node after evaluation.

    Args:
        state: Current ResearchState with evaluation result.

    Returns:
        Name of the next node: "compile_literature", "plan", or "request_support".
    """
    evaluation = state.get("evaluation")
    attempt = state.get("attempt", 0)
    max_retries = int(os.environ.get("AI_SKILL_MAX_RETRIES", "5"))

    if evaluation and evaluation.get("converged"):
        return "compile_literature"
    if attempt >= max_retries:
        return "request_support"
    return "plan"


def route_after_evaluate_cp3(state: ResearchState) -> str:
    """Conditional edge: decide next node after evaluation in CP3.

    Thin wrapper around the same logic as ``route_after_evaluate`` but routes
    to ``"compile_design"`` on convergence instead of ``"compile_literature"``.
    Kept separate to avoid modifying the CP2 routing function which is in
    production use.

    Args:
        state: Current ResearchState with evaluation result.

    Returns:
        Name of the next node: "compile_design", "plan", or "request_support".
    """
    evaluation = state.get("evaluation")
    attempt = state.get("attempt", 0)
    max_retries = int(os.environ.get("AI_SKILL_MAX_RETRIES", "5"))

    if evaluation and evaluation.get("converged"):
        return "compile_design"
    if attempt >= max_retries:
        return "request_support"
    return "plan"


# ---------------------------------------------------------------------------
# CP3 4-phase nodes
# ---------------------------------------------------------------------------


def read_attachments(
    state: ResearchState,
    registry: SkillRegistry | None = None,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Phase 1 of CP3: extract methodology-relevant content from each PDF.

    Processes PDFs in attachments/ one at a time using a RAG approach:
      1. Read PDF text with pdf_reader skill.
      2. Call LLM to extract only the methodology-relevant content.
      3. Save the result as ``{name}.md`` beside the PDF (persistent cache).

    If ``{name}.md`` already exists it is loaded directly — the PDF is not
    re-read on retries or subsequent runs.  This keeps the context window
    under control regardless of how many PDFs are in attachments/.

    The findings list contains one SkillOutput per PDF with the .md path and
    a short excerpt so ideate_design can reference specific files by name.

    Args:
        state: Current ResearchState.
        registry: Optional injected SkillRegistry (for testing).
        llm: Optional LLMClient override (for testing).

    Returns:
        Partial state with pdf findings, cp3_phase1_complete=True.
    """
    _configure_history(state, "read_attachments")
    _registry = registry or _get_default_registry()
    _llm = llm or LLMClient()
    workspace_path = state.get("workspace_path", "")
    ws = Path(workspace_path)
    attachments_dir = (
        ws.parent / "attachments" if ws.name == ".state" else ws / "attachments"
    )

    pdfs: list[Path] = sorted(attachments_dir.glob("*.pdf")) if attachments_dir.is_dir() else []
    if not pdfs:
        logger.warning("read_attachments: no PDFs found in %s", attachments_dir)

    # Namespace-isolated BM25 index for the CP3 methodology view.
    # Files produced: {stem}__cp3_methodology.md + .rag_index__cp3_methodology.json
    _CP3_NAMESPACE = "cp3_methodology"
    rag = RagIndex(attachments_dir, _CP3_NAMESPACE) if attachments_dir.is_dir() else None

    pdf_skill = _registry.get("pdf_reader")
    findings_current: list[SkillOutput] = []
    stage_str = PipelineStage.RESEARCH_DESIGN.value

    for pdf_path in pdfs:
        md_path = rag.md_path_for(pdf_path.stem) if rag is not None else pdf_path.with_suffix(".md")

        # ── Cache hit: .md already extracted for this namespace ───────────────
        if md_path.exists():
            try:
                md_content = md_path.read_text(encoding="utf-8", errors="ignore")
                # Ensure it is in the index even if created in a prior run
                if rag is not None and not rag.contains(pdf_path.stem):
                    rag.add(pdf_path.stem, md_content)
                findings_current.append(SkillOutput(
                    skill_name="read_attachments",
                    result={
                        "file_path": str(pdf_path),
                        "md_path": str(md_path),
                        "namespace": _CP3_NAMESPACE,
                        "excerpt": md_content[:500],
                        "cached": True,
                    },
                    confidence=0.8,
                    sources=[str(pdf_path)],
                    error=None,
                    cached=True,
                ))
                logger.debug("read_attachments[%s]: cache hit for %s", _CP3_NAMESPACE, pdf_path.name)
                continue
            except Exception as exc:
                logger.debug("read_attachments: cache read failed for %s: %s", md_path, exc)

        # ── Step 1: extract PDF text ──────────────────────────────────────────
        raw_text = ""
        if pdf_skill is not None:
            skill_input = SkillInput({
                "parameters": {"file_path": str(pdf_path), "_skill_name": "pdf_reader"},
                "objective": state.get("objective"),
                "stage": stage_str,
                "attempt": 0,
            })
            pdf_out = pdf_skill.run(skill_input)
            raw_text = (pdf_out.get("result") or {}).get("text") or ""
            if pdf_out.get("error"):
                logger.warning(
                    "read_attachments: pdf_reader error for %s: %s",
                    pdf_path.name, pdf_out["error"],
                )

        if not raw_text.strip():
            findings_current.append(SkillOutput(
                skill_name="read_attachments",
                result={"file_path": str(pdf_path), "error": "empty text after extraction"},
                confidence=0.0,
                sources=[str(pdf_path)],
                error=f"Could not extract text from {pdf_path.name}",
                cached=False,
            ))
            continue

        # ── Step 2: LLM extracts methodology-relevant content ─────────────────
        system, messages = build_extract_pdf_methodology_messages(
            filename=pdf_path.name,
            text=raw_text,
        )
        try:
            md_content: str = _llm.complete(
                messages=messages,
                system=system,
                max_tokens=1024,
            )
        except LLMClientError as exc:
            logger.warning(
                "read_attachments: LLM extraction failed for %s: %s", pdf_path.name, exc
            )
            md_content = f"## {pdf_path.name}\n\n*Extração automática falhou: {exc}*\n"

        # ── Step 3: save namespaced .md and update index ──────────────────────
        try:
            md_path.write_text(md_content, encoding="utf-8")
            logger.debug("read_attachments[%s]: saved %s", _CP3_NAMESPACE, md_path.name)
        except Exception as exc:
            logger.warning("read_attachments: could not save %s: %s", md_path, exc)

        if rag is not None:
            rag.add(pdf_path.stem, md_content)

        _history.log_skill_call(
            skill="read_attachments",
            stage=stage_str,
            attempt=0,
            step_id=-1,
            step_rationale=f"Phase 1 RAG extraction: {pdf_path.name}",
            request_sent={"file_path": str(pdf_path)},
            response_received={"md_path": str(md_path), "md_chars": len(md_content)},
        )

        findings_current.append(SkillOutput(
            skill_name="read_attachments",
            result={
                "file_path": str(pdf_path),
                "md_path": str(md_path),
                "namespace": _CP3_NAMESPACE,
                "excerpt": md_content[:500],
                "cached": False,
            },
            confidence=0.8,
            sources=[str(pdf_path)],
            error=None,
            cached=False,
        ))

    existing = state.get("findings", [])
    return {
        "findings": existing + findings_current,
        "findings_current": findings_current,
        "cp3_phase1_complete": True,
        "status": ResearchStatus.EXECUTING,
    }


def ideate_design(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Phase 2 of CP3: propose the full Research Design document.

    Uses CP1 charter + CP2 literature review + Phase 1 PDF context to generate
    an original research design proposal. On retry iterations, preserved sections
    (score >= 0.85) are injected verbatim and gaps from the previous evaluation
    are addressed.

    Args:
        state: Current ResearchState.
        llm: Optional LLMClient override.

    Returns:
        Partial state with research_design_doc (draft), status=EVALUATING.
    """
    _configure_history(state, "ideate_design")
    _llm = llm or LLMClient()

    charter_text = state.get("charter_document_text") or ""
    cp3_context = state.get("cp3_context") or {}

    literature_text = cp3_context.get("literature_summary") or ""
    if not literature_text:
        review_doc = state.get("literature_review_doc") or {}
        sections = review_doc.get("sections", [])
        literature_text = "\n\n".join(
            f"## {s.get('section_title', '')}\n{s.get('content', '')}"
            for s in sections
        )

    # Build PDF context via BM25 retrieval (RAG).
    # For each of the 9 mandatory CP3 sections, query the RagIndex and collect
    # the top-3 most relevant .md excerpts.  This caps context regardless of
    # how many PDFs exist in attachments/.
    workspace_path_str = state.get("workspace_path", "")
    ws_path = Path(workspace_path_str)
    attachments_dir_for_rag = (
        ws_path.parent / "attachments" if ws_path.name == ".state"
        else ws_path / "attachments"
    )
    # Use the same namespace that read_attachments wrote into
    _CP3_NAMESPACE = "cp3_methodology"
    rag = (
        RagIndex(attachments_dir_for_rag, _CP3_NAMESPACE)
        if attachments_dir_for_rag.is_dir()
        else None
    )

    # One query per mandatory CP3 section — captures the vocabulary each section needs
    _SECTION_QUERIES = [
        "research method experimental quasi-experimental observational study design type",
        "hypothesis research question testable falsifiable H0 H1",
        "variables independent dependent confounding operationalization measurement scale",
        "KPI metrics acceptance threshold statistical power sample size calculation",
        "data sources FAIR findable accessible interoperable reusable primary secondary",
        "instruments data collection protocol questionnaire survey interview sensor",
        "validity reliability internal external threats mitigation replication",
        "ethics CEP IRB LGPD consent privacy conflict of interest",
        "timeline schedule milestones PMBOK planning executing",
    ]

    # ── Context budget control ────────────────────────────────────────────────
    # The user message must stay well below the model's context window.
    # Rough budget (chars → tokens ≈ ÷ 4):
    #   charter_text:      ≤ 8 000 chars  (~2 000 tokens)
    #   literature_text:   ≤ 16 000 chars (~4 000 tokens)
    #   pdf_context total: ≤ 12 000 chars (~3 000 tokens) — top-8 unique docs × 1 500 chars
    #   preserved + gaps:  ≤  4 000 chars (~1 000 tokens)
    #   system prompt:     ~  3 000 tokens
    #   TOTAL INPUT:       ~13 000 tokens  →  leaves ~50 000 tokens for output (ample)
    _CHARTER_CHAR_LIMIT   = 8_000
    _LITERATURE_CHAR_LIMIT = 16_000
    _PDF_EXCERPT_LIMIT    = 1_500   # chars per doc in pdf_context
    _PDF_DOCS_LIMIT       = 8       # max unique docs included

    charter_text  = (charter_text or "")[:_CHARTER_CHAR_LIMIT]
    literature_text = (literature_text or "")[:_LITERATURE_CHAR_LIMIT]

    pdf_context: list[dict[str, Any]] = []
    seen_in_context: set[str] = set()

    if rag is not None and rag.size() > 0:
        for query in _SECTION_QUERIES:
            if len(seen_in_context) >= _PDF_DOCS_LIMIT:
                break
            for doc_id, excerpt, score in rag.search(query, top_k=2):
                if doc_id not in seen_in_context and len(seen_in_context) < _PDF_DOCS_LIMIT:
                    seen_in_context.add(doc_id)
                    pdf_context.append({
                        "file": doc_id,
                        "relevance_score": round(score, 3),
                        "methodology_notes": excerpt[:_PDF_EXCERPT_LIMIT],
                    })
    else:
        # Fallback: read namespaced .md files from findings if index not available
        all_findings = state.get("findings", [])
        seen_md: set[str] = set()
        for f in all_findings:
            if len(pdf_context) >= _PDF_DOCS_LIMIT:
                break
            if f.get("skill_name") != "read_attachments":
                continue
            result = f.get("result") or {}
            # Only use extracts from the cp3_methodology namespace
            if result.get("namespace", _CP3_NAMESPACE) != _CP3_NAMESPACE:
                continue
            md_path_str = result.get("md_path", "")
            if not md_path_str or md_path_str in seen_md:
                continue
            seen_md.add(md_path_str)
            md_path = Path(md_path_str)
            if md_path.exists():
                try:
                    md_content = md_path.read_text(encoding="utf-8", errors="ignore")
                    pdf_context.append({
                        "file": md_path.stem,
                        "methodology_notes": md_content[:_PDF_EXCERPT_LIMIT],
                    })
                except Exception:
                    pass

    preserved_sections: dict[str, Any] = state.get("cp3_preserved_sections") or {}
    evaluation = state.get("evaluation")
    gaps: list[str] = evaluation.get("gaps", []) if evaluation else []

    system, messages = build_ideate_design_messages(
        charter_document_text=charter_text,
        literature_document_text=literature_text,
        cp3_context=cp3_context,
        pdf_context=pdf_context,
        preserved_sections=preserved_sections,
        gaps=gaps,
    )

    try:
        design_obj: _ResearchDesignDocLLM = _llm.complete_structured(
            messages=messages,
            response_model=_ResearchDesignDocLLM,
            system=system,
            max_tokens=32768,   # raised: 9 sections × prose + structured fields needs room
        )
    except LLMClientError as exc:
        logger.error("ideate_design failed: %s", exc)
        return {"status": ResearchStatus.FAILED}

    sections: list[ResearchDesignSection] = [
        ResearchDesignSection(section_title=s.section_title, content=s.content)
        for s in design_obj.sections
    ]
    design_doc = ResearchDesignDoc(
        sections=sections,
        study_type=design_obj.study_type,
        research_paradigm=design_obj.research_paradigm,
        epistemological_stance=design_obj.epistemological_stance,
        hypotheses=design_obj.hypotheses,
        research_questions=design_obj.research_questions,
        variables=[v.model_dump() for v in design_obj.variables],
        instruments=design_obj.instruments,
        sampling_strategy=design_obj.sampling_strategy,
        sample_size_justification=design_obj.sample_size_justification,
        ethical_considerations=design_obj.ethical_considerations,
        validity_threats=design_obj.validity_threats,
        mitigation_strategies=[m.model_dump() for m in design_obj.mitigation_strategies],
        data_management_plan=design_obj.data_management_plan,
        metrics_and_kpis=design_obj.metrics_and_kpis,
        data_sources=design_obj.data_sources,
        collection_protocol=design_obj.collection_protocol,
        methodology_timeline=design_obj.methodology_timeline,
        reporting_standard=design_obj.reporting_standard,
        target_journal_tier=design_obj.target_journal_tier,
    )

    return {
        "research_design_doc": design_doc,
        "status": ResearchStatus.EVALUATING,
    }


def review_frameworks(
    state: ResearchState,
    registry: SkillRegistry | None = None,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Phase 3 of CP3: critique and adjust the design against formal standards.

    Searches the web for PRISMA-trAIce 2025 and other standard specifics, then
    calls the LLM to review the design against OR, PMBOK, PRISMA, PRISMA-trAIce
    2025, FAIR, SJR, ABNT NBR, and ISO 9001:2015. Targeted corrections are
    applied to produce an improved design document.

    Web searches for methodology standards require no user authorization.

    Args:
        state: Current ResearchState with research_design_doc from Phase 2.
        registry: Optional injected SkillRegistry (for testing).
        llm: Optional LLMClient override.

    Returns:
        Partial state with (possibly corrected) research_design_doc.
    """
    _configure_history(state, "review_frameworks")
    _registry = registry or _get_default_registry()
    _llm = llm or LLMClient()

    design_doc = state.get("research_design_doc")
    if not design_doc:
        logger.warning("review_frameworks: no research_design_doc — skipping.")
        return {}

    objective = state.get("cp3_context") or state.get("objective") or {}
    charter_summary = (
        f"Tópico: {objective.get('topic', '')}\n"
        f"Objetivos: {'; '.join(objective.get('goals') or [])}"
    )

    # Web searches for standards (always allowed in Phase 3)
    web_skill = _registry.get("web_search")
    search_snippets: list[str] = []
    prisma_traice_context = ""

    _framework_queries = [
        ("PRISMA-trAIce 2025 checklist AI research reporting", "en"),
        ("PRISMA-trAIce 2025 guidelines LLM studies", "en"),
        ("FAIR data principles research design compliance checklist", "en"),
        ("SJR Q1 Q2 methodology requirements academic research", "en"),
    ]

    if web_skill is not None:
        for query, lang in _framework_queries:
            skill_input = SkillInput({
                "parameters": {
                    "query": query,
                    "max_results": 5,
                    "language": lang,
                    "_skill_name": "web_search",
                },
                "objective": state.get("objective"),
                "stage": PipelineStage.RESEARCH_DESIGN.value,
                "attempt": state.get("attempt", 0),
            })
            out = web_skill.run(skill_input)
            if not out.get("error"):
                results = (out.get("result") or {}).get("results", [])
                for r in results[:3]:
                    snippet = r.get("snippet") or r.get("description") or ""
                    title = r.get("title") or ""
                    if snippet:
                        search_snippets.append(f"[{title}] {snippet}")
                        if "trAIce" in title or "trAIce" in snippet:
                            prisma_traice_context += f"\n{snippet}"

    framework_search_results = "\n".join(search_snippets) if search_snippets else ""

    system, messages = build_review_frameworks_messages(
        design_doc=dict(design_doc),
        charter_summary=charter_summary,
        framework_search_results=framework_search_results,
        prisma_traice_context=prisma_traice_context.strip(),
    )

    try:
        revised_obj: _ResearchDesignDocLLM = _llm.complete_structured(
            messages=messages,
            response_model=_ResearchDesignDocLLM,
            system=system,
            max_tokens=32768,
        )
    except LLMClientError as exc:
        logger.error("review_frameworks failed: %s — keeping Phase 2 draft", exc)
        return {}

    sections: list[ResearchDesignSection] = [
        ResearchDesignSection(section_title=s.section_title, content=s.content)
        for s in revised_obj.sections
    ]
    revised_doc = ResearchDesignDoc(
        sections=sections,
        study_type=revised_obj.study_type,
        research_paradigm=revised_obj.research_paradigm,
        epistemological_stance=revised_obj.epistemological_stance,
        hypotheses=revised_obj.hypotheses,
        research_questions=revised_obj.research_questions,
        variables=[v.model_dump() for v in revised_obj.variables],
        instruments=revised_obj.instruments,
        sampling_strategy=revised_obj.sampling_strategy,
        sample_size_justification=revised_obj.sample_size_justification,
        ethical_considerations=revised_obj.ethical_considerations,
        validity_threats=revised_obj.validity_threats,
        mitigation_strategies=[m.model_dump() for m in revised_obj.mitigation_strategies],
        data_management_plan=revised_obj.data_management_plan,
        metrics_and_kpis=revised_obj.metrics_and_kpis,
        data_sources=revised_obj.data_sources,
        collection_protocol=revised_obj.collection_protocol,
        methodology_timeline=revised_obj.methodology_timeline,
        reporting_standard=revised_obj.reporting_standard,
        target_journal_tier=revised_obj.target_journal_tier,
    )
    return {"research_design_doc": revised_doc}


class _SectionScoreLLM(BaseModel):
    """Score for a single design section."""

    section_title: str = Field(description="Title of the section being scored.")
    score: float = Field(description="Quality score 0.0–1.0.", ge=0.0, le=1.0)
    rationale: str = Field(description="One-sentence justification for the score.")
    gaps: list[str] = Field(
        default_factory=list,
        description="Specific gaps vs. CP1 objectives (empty when score >= 0.85).",
    )


class _EvaluateObjectivesLLM(BaseModel):
    """Per-section scores and aggregated gaps from evaluate_objectives."""

    sections: list[_SectionScoreLLM] = Field(
        description="One score entry per mandatory design section."
    )
    gaps: list[str] = Field(
        description="Aggregated list of gaps across all sections with score < 0.85."
    )


def evaluate_objectives(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Phase 4 of CP3: evaluate the design document against CP1 research objectives.

    Scores each of the 9 mandatory sections individually against the research
    objectives from CP1. Sections scoring >= 0.85 are preserved verbatim in
    subsequent ideate_design calls. Computes total_score as the mean of section
    scores and sets converged = total_score >= _CONVERGENCE_THRESHOLD.

    Args:
        state: Current ResearchState with research_design_doc from Phase 3.
        llm: Optional LLMClient override.

    Returns:
        Partial state with evaluation, cp3_preserved_sections, attempt+1,
        and quality_history snapshot.
    """
    _configure_history(state, "evaluate_objectives")
    _llm = llm or LLMClient()

    design_doc = state.get("research_design_doc")
    if not design_doc:
        logger.warning("evaluate_objectives: no research_design_doc — failing.")
        return {
            "evaluation": EvaluationResult(
                per_metric=[], total_score=0.0, converged=False,
                gaps=["research_design_doc ausente"], regression=False,
            ),
            "attempt": state.get("attempt", 0) + 1,
        }

    objective = state.get("cp3_context") or state.get("objective") or {}
    attempt = state.get("attempt", 0)
    quality_history = state.get("quality_history", [])
    stage_str = PipelineStage.RESEARCH_DESIGN.value

    system, messages = build_evaluate_objectives_messages(
        objective=objective,
        design_doc=dict(design_doc),
        threshold=_CONVERGENCE_THRESHOLD,
    )

    try:
        eval_obj: _EvaluateObjectivesLLM = _llm.complete_structured(
            messages=messages,
            response_model=_EvaluateObjectivesLLM,
            system=system,
        )
    except LLMClientError as exc:
        logger.error("evaluate_objectives failed: %s", exc)
        eval_obj = _EvaluateObjectivesLLM(
            sections=[],
            gaps=[f"Evaluation error: {exc}"],
        )

    # Compute total_score deterministically
    scores = [s.score for s in eval_obj.sections]
    total_score = sum(scores) / len(scores) if scores else 0.0
    converged = total_score >= _CONVERGENCE_THRESHOLD

    # Regression detection
    previous_score = quality_history[-1]["total_score"] if quality_history else None
    regression = (
        previous_score is not None and (previous_score - total_score) > 0.05
    )

    per_metric: list[MetricScore] = [
        MetricScore(
            metric=s.section_title,
            score=s.score,
            rationale=s.rationale,
            gaps=s.gaps,
        )
        for s in eval_obj.sections
    ]

    evaluation = EvaluationResult(
        per_metric=per_metric,
        total_score=total_score,
        converged=converged,
        gaps=eval_obj.gaps,
        regression=regression,
    )

    # Build section content lookup from the current design doc
    design_sections: list[dict[str, Any]] = (design_doc.get("sections") or [])
    section_content_by_title: dict[str, str] = {
        s.get("section_title", ""): s.get("content", "")
        for s in design_sections
    }

    # ── CP3 anti-regression + per-section historical best tracking ────────
    #
    # cp3_best_section_scores tracks the highest score ever achieved for each
    # section title across all attempts. cp3_preserved_sections is rebuilt from
    # these historical bests every evaluation, so:
    #   (a) a section that scores 0.83 in attempt 1 and 0.70 in attempt 2 keeps
    #       its 0.83 content instead of being silently downgraded, and
    #   (b) the threshold for preservation shifts from the hard ≥0.85 cutoff to
    #       "any improvement over the historical best".
    #
    # cp3_best_doc / cp3_best_score track the document with the highest total_score.
    # When a retry regresses, research_design_doc is reverted to the best doc so
    # ideate_design always starts from the strongest known base.

    best_section_scores: dict[str, Any] = dict(state.get("cp3_best_section_scores") or {})

    for s in eval_obj.sections:
        current_score = s.score
        prev_best = best_section_scores.get(s.section_title, {})
        prev_best_score = float(prev_best.get("score", 0.0))
        if current_score > prev_best_score:
            content = section_content_by_title.get(s.section_title, "")
            if content:
                best_section_scores[s.section_title] = {
                    "content": content,
                    "score": current_score,
                }
    # Rebuild cp3_preserved_sections from historical bests.
    # Preserve every section that has ever reached ≥0.85 (verbatim in next call)
    # AND any section whose historical best is the best we have seen
    # (even if below 0.85 — prevents silent regressions on partially-good sections).
    preserved: dict[str, Any] = {}
    for title, best in best_section_scores.items():
        if best.get("score", 0.0) >= 0.85:
            preserved[title] = best

    # ── Total-doc anti-regression ─────────────────────────────────────────
    cp3_best_score = float(state.get("cp3_best_score") or 0.0)
    extra_updates: dict[str, Any] = {"cp3_best_section_scores": best_section_scores}

    if total_score > cp3_best_score and design_doc.get("sections"):
        extra_updates["cp3_best_doc"] = design_doc
        extra_updates["cp3_best_score"] = total_score
        logger.info(
            "CP3 new best total score: %.2f → %.2f (attempt %d)",
            cp3_best_score, total_score, attempt,
        )
    elif state.get("cp3_best_doc") and total_score < cp3_best_score - 0.01:
        # Revert the design doc to the best known version so the next
        # ideate_design call starts from a stronger base.
        logger.warning(
            "CP3 regression %.2f → %.2f — reverting research_design_doc to best (%.2f)",
            previous_score or 0.0, total_score, cp3_best_score,
        )
        extra_updates["research_design_doc"] = state["cp3_best_doc"]

    snapshot = QualitySnapshot(
        attempt=attempt,
        stage=stage_str,
        total_score=total_score,
        per_metric_scores={s.section_title: s.score for s in eval_obj.sections},
        skills_used=["ideate_design", "review_frameworks"],
        cache_hit_rate=0.0,
    )

    new_quality_history = list(quality_history) + [snapshot]
    stage_quality_history = dict(state.get("stage_quality_history") or {})
    stage_quality_history.setdefault(stage_str, [])
    stage_quality_history[stage_str] = list(stage_quality_history[stage_str]) + [snapshot]

    if regression:
        logger.warning(
            "Quality regression in CP3: %.2f → %.2f (attempt %d)",
            previous_score, total_score, attempt,
        )

    return {
        "evaluation": evaluation,
        "cp3_preserved_sections": preserved,
        "attempt": attempt + 1,
        "quality_history": new_quality_history,
        "stage_quality_history": stage_quality_history,
        "status": ResearchStatus.PLANNING,
        **extra_updates,
    }


def route_after_evaluate_objectives(state: ResearchState) -> str:
    """Conditional edge after evaluate_objectives in CP3.

    Routes to deliver_design on convergence, ideate_design for retry,
    or request_support when max retries are exhausted.

    Args:
        state: Current ResearchState with evaluation result.

    Returns:
        Name of the next node.
    """
    evaluation = state.get("evaluation")
    attempt = state.get("attempt", 0)
    max_retries = int(os.environ.get("AI_SKILL_MAX_RETRIES", "5"))

    if evaluation and evaluation.get("converged"):
        return "deliver_design"
    if attempt >= max_retries:
        return "request_support"
    return "ideate_design"


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------


def _get_project_workspace(workspace_path: str) -> ProjectWorkspace | None:
    """Return the ProjectWorkspace linked to an internal state path, if any.

    Convention: the ResearchWorkspace lives at ``<project>/.state/``.
    If the parent of *workspace_path* contains ``workspace.yaml``, it is a
    ProjectWorkspace and is returned; otherwise returns None.

    Args:
        workspace_path: Absolute path to the ResearchWorkspace directory.

    Returns:
        The :class:`ProjectWorkspace` instance, or None.
    """
    if not workspace_path:
        return None
    parent = Path(workspace_path).parent
    if (parent / "workspace.yaml").exists():
        return ProjectWorkspace.from_path(parent)
    return None


def _charter_to_docx(objective: ResearchObjective) -> bytes:
    """Serialise a ResearchObjective into a minimal .docx file.

    Args:
        objective: The research objective TypedDict.

    Returns:
        Raw bytes of the generated .docx file.
    """
    try:
        import io
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        logger.warning("python-docx not available — skipping charter .docx export.")
        return b""

    doc = Document()
    doc.add_heading("Research Charter", level=1)

    topic = objective.get("topic", "")
    doc.add_heading("Tópico de Pesquisa", level=2)
    doc.add_paragraph(topic)

    goals = objective.get("goals", [])
    if goals:
        doc.add_heading("Objetivos", level=2)
        for goal in goals:
            doc.add_paragraph(goal, style="List Bullet")

    metrics = objective.get("success_metrics", [])
    if metrics:
        doc.add_heading("Métricas de Sucesso", level=2)
        for m in metrics:
            doc.add_paragraph(m, style="List Bullet")

    constraints = objective.get("scope_constraints", [])
    if constraints:
        doc.add_heading("Restrições de Escopo", level=2)
        for c in constraints:
            doc.add_paragraph(c, style="List Bullet")

    methodology = objective.get("methodology_preference", "")
    if methodology:
        doc.add_heading("Preferência Metodológica", level=2)
        doc.add_paragraph(methodology)

    # Render stage-specific guidelines — one section per CP stage
    _STAGE_DISPLAY_NAMES: dict[str, str] = {
        "literature_review":      "CP2 — Diretrizes da Revisão Bibliográfica",
        "research_design":        "CP3 — Diretrizes do Design de Pesquisa",
        "data_collection_guide":  "CP4 — Diretrizes da Coleta de Dados",
        "analysis_guide":         "CP5 — Diretrizes da Análise",
        "results_interpretation": "CP6/7 — Diretrizes de Interpretação de Resultados",
        "paper_composition":      "CP8 — Diretrizes para Redação do Artigo",
        "publication":            "CP8+ — Diretrizes para Publicação",
    }
    stage_guidelines: dict[str, list[str]] = objective.get("stage_guidelines") or {}
    for stage_key, display_name in _STAGE_DISPLAY_NAMES.items():
        directives = stage_guidelines.get(stage_key, [])
        if directives:
            doc.add_heading(display_name, level=2)
            for directive in directives:
                doc.add_paragraph(directive, style="List Bullet")

    bib_style = objective.get("bibliography_style", "abnt")
    language = objective.get("language", "pt-BR")
    generated_at = objective.get("generated_at", "")
    meta_para = doc.add_paragraph()
    meta_line = f"Estilo bibliográfico: {bib_style}  |  Idioma: {language}"
    if generated_at:
        meta_line += f"  |  Gerado em: {generated_at}"
    meta_para.add_run(meta_line).font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()




def _configure_history(state: ResearchState, node_name: str) -> None:
    """Point the history logger at the current workspace and set node + stage names."""
    workspace_path = state.get("workspace_path", "")
    stage = state.get("stage")
    stage_str = stage.value if hasattr(stage, "value") else str(stage or "")
    if workspace_path:
        _history.configure(workspace_path)
    _history.set_current_stage(stage_str)
    set_current_node(node_name)


def _default_literature_review_guidelines() -> list[str]:
    """Return generic bibliographic-only guidelines for CP2.

    Used as fallback when the charter was created before stage_guidelines
    were introduced, or when the researcher did not customise them.
    These criteria are intentionally restricted to the bibliographic review
    scope — they never reference implementation, experiments, or publication.
    """
    return [
        "Realizar pelo menos 4 buscas distintas com queries diferentes cobrindo aspectos "
        "teóricos, metodológicos e aplicados do tema de pesquisa.",
        "Cada busca deve retornar resultados relevantes (total_found > 0); consultas em "
        "arXiv e Semantic Scholar são obrigatórias.",
        "O conjunto de buscas deve cobrir pelo menos 2 sub-áreas temáticas distintas "
        "derivadas do tópico de pesquisa.",
        "Ao menos 30% dos papers retornados devem ter sido publicados nos últimos 5 anos "
        "(campo year >= ano_atual - 5).",
        "Para ao menos 3 papers de alta relevância, executar summarização completa via "
        "content_summarizer para extrair contribuições detalhadas ao campo.",
        "As queries devem ser específicas o suficiente para retornar papers do domínio "
        "correto — queries genéricas demais (total_found > 200) devem ser refinadas.",
        "Cobrir tanto referências seminais (> 10 anos) quanto referências recentes "
        "(< 3 anos) para mostrar evolução do campo.",
        "Documentar os termos de busca e bases consultadas em cada step do plano.",
    ]


def _default_research_design_guidelines() -> list[str]:
    """Return PMBOK/ISO/OR-aligned methodology guidelines for CP3 (Research Design).

    Used instead of any charter-generated version to prevent the planner from
    broadening CP3 scope into data collection, analysis, or article writing.
    All 8 directives are scoped to design-phase deliverables only.
    """
    return [
        # OR Stage 2: model construction — method selection
        "Selecionar e justificar o método de pesquisa (experimental/quasi-experimental/"
        "observacional/estudo-de-caso/pesquisa-ação/revisão-sistemática/misto) com base "
        "explícita na literatura revisada (CP2) e no objetivo de pesquisa (CP1).",

        # OR Stage 2: hypothesis formulation
        "Formular ao menos 2 hipóteses (H1, H2, …) ou questões de pesquisa testáveis, "
        "com definições operacionais precisas e formato falsificável.",

        # Variable identification with measurement scales
        "Identificar e tabular todas as variáveis independentes (VI), dependentes (VD) "
        "e confundidoras com: nome, tipo, operacionalização e escala de medição "
        "(nominal/ordinal/intervalar/racional).",

        # Metrics and data goals (ISO 9001:2015 acceptance criteria)
        "Definir KPIs, limiares de aceitação, poder estatístico e tamanho amostral "
        "mínimo com justificativa (análise de poder ou critério de saturação qualitativa).",

        # Data sources with FAIR compliance
        "Especificar fontes de dados primárias e secundárias com critérios de qualidade "
        "(credibilidade, recência, acessibilidade) e conformidade FAIR "
        "(Findable, Accessible, Interoperable, Reusable).",

        # Validity and reliability (ISO 9001:2015)
        "Identificar ameaças à validade interna e externa, confiabilidade e "
        "replicabilidade; propor estratégias de mitigação e protocolo de replicação.",

        # Ethics (LGPD / IRB)
        "Enumerar considerações éticas: requisitos de CEP/IRB, conformidade LGPD, "
        "consentimento informado e conflitos de interesse.",

        # PMBOK Planning: methodology timeline
        "Propor cronograma metodológico com marcos PMBOK (Planning → Executing) "
        "alinhando CP3 (design) → CP4 (coleta) → CP5 (análise), com estimativas de prazo.",
    ]


def _default_stage_guidelines(stage_str: str) -> list[str]:
    """Return generic stage-specific guidelines for stages without custom directives.

    Used as fallback when the charter was created without stage_guidelines for
    a given stage. Criteria are scoped to each stage's specific deliverable and
    never reference global project goals (implementation, publication, etc.).

    Args:
        stage_str: PipelineStage value string (e.g. "research_design").

    Returns:
        List of guideline strings, or empty list if stage_str is unrecognised.
    """
    defaults: dict[str, list[str]] = {
        "research_design": [
            "Identificar e justificar o método de pesquisa mais adequado ao objetivo.",
            "Formular pelo menos 2 hipóteses ou questões de pesquisa testáveis.",
            "Definir variáveis dependentes e independentes com precisão operacional.",
            "Especificar os critérios de validade e confiabilidade da abordagem escolhida.",
            "Relacionar a metodologia escolhida com precedentes na literatura revisada.",
        ],
        "data_collection_guide": [
            "Produzir protocolo de coleta com instrumento, passos e critérios de aceitação.",
            "Definir tamanho amostral mínimo com justificativa estatística.",
            "Especificar formato de entrega dos dados (estrutura, unidades, codificação).",
            "Incluir checklist de verificação pré-coleta para o pesquisador.",
        ],
        "analysis_guide": [
            "Descrever passo a passo do método analítico com fórmulas em LaTeX.",
            "Especificar softwares ou bibliotecas recomendados com versões.",
            "Incluir critérios de interpretação dos resultados (ex: p-value, effect size).",
            "Antecipar possíveis limitações ou fontes de erro na análise.",
        ],
        "results_interpretation": [
            "Descrever resultados com base exclusivamente nos dados fornecidos pelo usuário.",
            "Relacionar cada resultado com as hipóteses formuladas no design de pesquisa.",
            "Identificar resultados inesperados e propor explicações plausíveis.",
            "Discutir limitações dos resultados obtidos.",
        ],
        "paper_composition": [
            "Produzir todas as seções: Abstract, Introdução, Metodologia, Resultados, "
            "Discussão e Conclusão.",
            "Garantir coerência narrativa entre seções e consistência das citações inline.",
            "Aplicar formatação ABNT NBR 6023:2018 nas referências.",
            "Revisar que cada afirmação factual está suportada por uma citação.",
        ],
        "publication": [
            "Validar conformidade com TOP Guidelines (transparência, abertura).",
            "Verificar FAIR principles para dados e materiais de pesquisa.",
            "Exportar documento em .docx, Markdown e PDF.",
            "Revisar metadados do documento (título, autores, palavras-chave, resumo).",
        ],
    }
    return defaults.get(stage_str, [])


def _today_dd_mm_yyyy() -> str:
    """Return today's date in dd/mm/yyyy format for charter generation stamps.

    Returns:
        String like "22/03/2026".
    """
    from datetime import datetime
    now = datetime.now()
    return f"{now.day:02d}/{now.month:02d}/{now.year}"


def _abnt_access_date() -> str:
    """Return today's date in ABNT access-date format.

    Returns:
        String like "Acesso em: 22 mar. 2026."
    """
    from datetime import datetime
    _month_abbr = [
        "jan.", "fev.", "mar.", "abr.", "maio", "jun.",
        "jul.", "ago.", "set.", "out.", "nov.", "dez.",
    ]
    now = datetime.now()
    return f"Acesso em: {now.day} {_month_abbr[now.month - 1]} {now.year}."


def _extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    """Extract text from a PDF byte string using PyMuPDF (first 5 pages).

    Returns up to 3000 characters of extracted text, or empty string on error.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        pages_text: list[str] = []
        for page_num in range(min(5, len(doc))):
            pages_text.append(doc[page_num].get_text())
        return "\n".join(pages_text)[:3000]
    except Exception as exc:
        logger.debug("PyMuPDF extraction failed: %s", exc)
        return ""


def _find_pdf_links_in_html(html: str, base_url: str) -> list[str]:
    """Parse an HTML page and return candidate PDF/download URLs.

    Looks for ``<a>`` tags whose href ends in ``.pdf`` or whose text/href
    contains keywords like ``pdf``, ``download``, ``full-text``, ``fulltext``,
    ``full_text``, ``preprint``, ``manuscript``.

    Returns up to 5 unique absolute URLs.
    """
    try:
        from bs4 import BeautifulSoup
        from urllib.parse import urljoin
    except ImportError:
        return []

    soup = BeautifulSoup(html, "html.parser")
    seen: set[str] = set()
    results: list[str] = []

    _LINK_KEYWORDS = (
        "pdf", "download", "full-text", "fulltext", "full_text",
        "preprint", "manuscript", "article", "paper",
    )

    for a in soup.find_all("a", href=True):
        href: str = a.get("href", "").strip()
        if not href:
            continue
        text = a.get_text(" ", strip=True).lower()
        href_lower = href.lower()

        is_candidate = (
            href_lower.endswith(".pdf")
            or any(kw in href_lower for kw in _LINK_KEYWORDS)
            or any(kw in text for kw in _LINK_KEYWORDS)
        )
        if not is_candidate:
            continue

        full_url = urljoin(base_url, href)
        if full_url.startswith(("http://", "https://")) and full_url not in seen:
            seen.add(full_url)
            results.append(full_url)
            if len(results) >= 5:
                break

    return results


_S2_PAPER_API = "https://api.semanticscholar.org/graph/v1/paper"
_S2_PAPER_FIELDS = "title,authors,abstract,year,openAccessPdf"


def _fetch_via_semantic_scholar_api(url: str, timeout: int = 15) -> tuple[bool, str, str] | None:
    """Fetch paper metadata from the Semantic Scholar API when the URL is recognised.

    Supports:
    - ``semanticscholar.org/paper/{id}`` — direct S2 paper page
    - ``doi.org/{doi}`` or any URL containing ``/doi/`` — DOI lookup
    - ``arxiv.org/abs/{id}`` or ``arxiv.org/pdf/{id}`` — arXiv lookup

    Args:
        url: The reference URL to try.
        timeout: Per-request timeout in seconds.

    Returns:
        ``(True, abstract_text, access_method)`` on success, or ``None`` if
        the URL is not a recognised S2-resolvable source.
    """
    import os
    import re
    import time
    import urllib.request
    import urllib.error
    from urllib.parse import urlparse, unquote

    paper_id: str | None = None

    # ── Semantic Scholar paper page ──────────────────────────────────────────
    # URL format: /paper/{optional-title-slug}/{40-char-hex-id}
    # The paper ID is the LAST pure-alphanumeric path segment (≥20 chars).
    if "semanticscholar.org" in url:
        path_segments = [s for s in urlparse(url).path.split("/") if s]
        # Paper ID is hex-only, ≥20 chars; title slug segments contain hyphens
        s2_candidates = [s for s in path_segments if re.match(r"^[A-Za-z0-9]{20,}$", s)]
        if s2_candidates:
            paper_id = s2_candidates[-1]

    # ── DOI (doi.org/10.xxxx/... or /doi/ path) ──────────────────────────────
    if paper_id is None:
        m = re.search(r"(?:^|/)(?:doi\.org/|doi/)(10\.\d{4,}/[^\s?#]+)", url)
        if m:
            paper_id = f"DOI:{unquote(m.group(1)).rstrip('/')}"

    # ── arXiv ────────────────────────────────────────────────────────────────
    if paper_id is None:
        m = re.search(r"arxiv\.org/(?:abs|pdf)/([0-9]{4}\.[0-9]{4,5}(?:v\d+)?)", url)
        if m:
            paper_id = f"arXiv:{m.group(1).split('v')[0]}"

    if paper_id is None:
        return None  # not a recognised S2-resolvable source

    api_url = f"{_S2_PAPER_API}/{paper_id}?fields={_S2_PAPER_FIELDS}"
    headers: dict[str, str] = {"Accept": "application/json"}
    api_key = os.environ.get("SEMANTIC_SCHOLAR_API_KEY", "")
    if api_key:
        headers["x-api-key"] = api_key

    # Rate limit: 1 req/s without API key, 10 req/s with key
    _interval = 0.1 if api_key else 1.1
    time.sleep(_interval)

    try:
        req = urllib.request.Request(api_url, headers=headers)
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            import json as _json
            data: dict = _json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        logger.debug(
            "S2 API HTTP error for %s (paper_id=%s): %s — falling back to HTTP/Tavily/Firecrawl.",
            url, paper_id, exc,
        )
        return None  # let the caller fall back to HTTP/Tavily/Firecrawl/Playwright
    except Exception as exc:
        logger.debug("S2 API fetch failed for %s: %s", url, exc)
        return None  # let the caller fall back to HTTP

    abstract = data.get("abstract") or ""
    title = data.get("title") or ""
    authors_raw = data.get("authors") or []
    authors_str = ", ".join(a.get("name", "") for a in authors_raw[:3])
    year = data.get("year") or ""

    text = f"Title: {title}\nAuthors: {authors_str}\nYear: {year}\n\nAbstract:\n{abstract}"
    return True, text[:3000], "Semantic_Scholar_API"


def _fetch_via_tavily(url: str) -> tuple[bool, str, str] | None:
    """Fetch page content via Tavily's /extract endpoint.

    Lighter-weight than Firecrawl (no headless browser), but can bypass some
    bot-blocked pages that urllib cannot reach. Falls through (returns None) when
    the API key is absent or the extraction fails, so the caller can escalate to
    Firecrawl or Playwright.

    Args:
        url: The URL to extract.

    Returns:
        ``(True, text_snippet, "Tavily")`` on success, or ``None`` on failure.
    """
    import os

    api_key = os.environ.get("TAVILY_API_KEY", "")
    if not api_key:
        return None

    try:
        from tavily import TavilyClient  # type: ignore[import-untyped]

        client = TavilyClient(api_key=api_key)
        response = client.extract(urls=[url])
        results = response.get("results", [])
        if not results:
            return None

        text: str = results[0].get("raw_content", "") or ""
        if not text:
            return None

        logger.debug("Tavily extract succeeded for %s (%d chars).", url, len(text))
        return True, text[:3000], "Tavily"
    except Exception as exc:
        logger.debug("Tavily extract failed for %s: %s", url, exc)
        return None


def _fetch_via_firecrawl(url: str) -> tuple[bool, str, str] | None:
    """Scrape a URL via the Firecrawl API as a fallback for bot-blocked pages.

    Args:
        url: The URL to scrape.

    Returns:
        ``(accessible, text_snippet, "Firecrawl")`` on success, or ``None`` when
        the API key is not configured or the scrape fails (caller should fall back
        to returning ``False``).
    """
    import os

    api_key = os.environ.get("FIRECRAWL_API_KEY", "")
    if not api_key:
        return None

    try:
        from firecrawl import FirecrawlApp  # type: ignore[import-untyped]

        app = FirecrawlApp(api_key=api_key)
        # firecrawl-py ≥ 1.0 uses .scrape(); older versions used .scrape_url()
        scrape_fn = getattr(app, "scrape", None) or getattr(app, "scrape_url")
        result = scrape_fn(url, formats=["markdown"])

        text = ""
        if hasattr(result, "markdown") and result.markdown:
            text = result.markdown
        elif isinstance(result, dict):
            text = result.get("markdown", "") or result.get("content", "")

        logger.debug("Firecrawl fetch succeeded for %s (%d chars).", url, len(text))
        return True, text[:3000], "Firecrawl"
    except Exception as exc:
        logger.debug("Firecrawl fetch failed for %s: %s", url, exc)
        return None


def _fetch_via_playwright(url: str, timeout_ms: int = 30000) -> tuple[bool, str, str] | None:
    """Render a page with a headless Chromium browser via Playwright.

    Last-resort fallback for pages that require full JavaScript execution and
    cannot be handled by urllib (static HTTP) or Firecrawl (managed headless).
    Requires ``playwright`` to be installed and browsers to be downloaded:
    ``uv add playwright && playwright install chromium``.

    Args:
        url: The URL to render.
        timeout_ms: Navigation timeout in milliseconds. Default: 30000.

    Returns:
        ``(True, text_snippet, "Playwright")`` on success, or ``None`` when
        Playwright is not installed or navigation fails.
    """
    try:
        from playwright.sync_api import sync_playwright  # type: ignore[import-untyped]
    except ImportError:
        logger.debug("Playwright not installed — skipping Playwright fallback for %s.", url)
        return None

    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            try:
                context = browser.new_context(
                    user_agent=(
                        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                        "AppleWebKit/537.36 (KHTML, like Gecko) "
                        "Chrome/124.0.0.0 Safari/537.36"
                    ),
                    locale="pt-BR",
                )
                page = context.new_page()
                page.goto(url, wait_until="networkidle", timeout=timeout_ms)
                # Extract body text — cleaner than inner_html
                text = page.evaluate(
                    """() => {
                        const remove = ['script','style','nav','footer','header','aside'];
                        remove.forEach(t => document.querySelectorAll(t)
                            .forEach(el => el.remove()));
                        return document.body ? document.body.innerText : '';
                    }"""
                )
                text = (text or "").strip()
            finally:
                browser.close()

        if not text:
            logger.debug("Playwright rendered empty body for %s.", url)
            return None

        logger.debug("Playwright fetch succeeded for %s (%d chars).", url, len(text))
        return True, text[:3000], "Playwright"

    except Exception as exc:
        logger.debug("Playwright fetch failed for %s: %s", url, exc)
        return None


def _fetch_url_content(url: str, timeout: int = 15) -> tuple[bool, str, str]:
    """Fetch a URL and return (accessible, text_snippet, access_method).

    Strategy (in order):
    0. For Semantic Scholar / DOI / arXiv URLs: use the S2 Graph API to obtain
       the paper abstract — avoids bot-blocking on journal landing pages.
    1. Direct HTTP fetch with retry on 429 (up to 2 retries, 5 s wait each).
       On HTTP error or exception: try Firecrawl before giving up.
    2. If the response is a PDF (Content-Type or .pdf URL), extract text via
       PyMuPDF instead of returning an empty snippet.
    3. If the HTML page contains PDF/download links, follow the first one that
       yields extractable PDF content — marked as "PDF_fallback".
    4. If the HTTP fetch returns empty bytes: try Firecrawl before returning False.

    Args:
        url: The URL to fetch.
        timeout: Per-request timeout in seconds.

    Returns:
        Tuple of (accessible, text_snippet, access_method).
        ``accessible`` — True when any 2xx response was obtained.
        ``text_snippet`` — up to 3000 chars of body text (may be empty).
        ``access_method`` — "HTTP" | "PDF_direto" | "PDF_fallback" | "Firecrawl".
    """
    import time
    import urllib.error
    import urllib.request

    if not url or not url.startswith(("http://", "https://")):
        return False, "", "HTTP"

    # Strategy 0: use S2 API for recognised academic sources
    api_result = _fetch_via_semantic_scholar_api(url, timeout=timeout)
    if api_result is not None:
        return api_result

    _UA = (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )

    def _build_req(target_url: str) -> urllib.request.Request:
        return urllib.request.Request(
            target_url,
            headers={
                "User-Agent": _UA,
                "Accept": (
                    "text/html,application/xhtml+xml,application/xml;q=0.9,"
                    "application/pdf,*/*;q=0.8"
                ),
                "Accept-Language": "pt-BR,pt;q=0.9,en;q=0.8",
            },
        )

    raw_bytes: bytes = b""
    content_type: str = ""
    max_retries = 2

    for attempt in range(max_retries + 1):
        try:
            req = _build_req(url)
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                content_type = resp.headers.get("Content-Type", "").lower()
                raw_bytes = resp.read(5 * 1024 * 1024)  # max 5 MB
            break  # success — exit retry loop
        except urllib.error.HTTPError as exc:
            if exc.code == 429 and attempt < max_retries:
                wait = 5
                retry_after = exc.headers.get("Retry-After", "")
                try:
                    wait = min(int(retry_after), 5)
                except (ValueError, TypeError):
                    pass
                logger.debug(
                    "429 rate-limit on %s — waiting %ds (attempt %d/%d)",
                    url, wait, attempt + 1, max_retries,
                )
                time.sleep(wait)
                continue
            logger.debug("HTTP error fetching %s: %s", url, exc)
            tv = _fetch_via_tavily(url)
            if tv is not None:
                return tv
            fc = _fetch_via_firecrawl(url)
            if fc is not None:
                return fc
            pw = _fetch_via_playwright(url)
            if pw is not None:
                return pw
            return False, "", "HTTP"
        except Exception as exc:
            logger.debug("Failed to fetch %s: %s", url, exc)
            tv = _fetch_via_tavily(url)
            if tv is not None:
                return tv
            fc = _fetch_via_firecrawl(url)
            if fc is not None:
                return fc
            pw = _fetch_via_playwright(url)
            if pw is not None:
                return pw
            return False, "", "HTTP"

    if not raw_bytes:
        tv = _fetch_via_tavily(url)
        if tv is not None:
            return tv
        fc = _fetch_via_firecrawl(url)
        if fc is not None:
            return fc
        pw = _fetch_via_playwright(url)
        if pw is not None:
            return pw
        return False, "", "HTTP"

    # ── PDF: direct URL or PDF Content-Type ──────────────────────────────────
    is_pdf_url = url.lower().split("?")[0].endswith(".pdf")
    is_pdf_ct = "pdf" in content_type
    if is_pdf_url or is_pdf_ct:
        text = _extract_text_from_pdf_bytes(raw_bytes)
        if text:
            return True, text, "PDF_direto"
        # Readable but empty extract (scanned PDF etc.) — still accessible
        return True, "", "PDF_direto"

    # ── HTML: decode and try PDF link fallback ────────────────────────────────
    try:
        html_text = raw_bytes.decode("utf-8", errors="ignore")
    except Exception:
        html_text = ""

    snippet = html_text[:3000]

    # Try to find and follow PDF/download links embedded in the page
    pdf_links = _find_pdf_links_in_html(html_text, url)
    for pdf_url in pdf_links:
        try:
            pdf_req = _build_req(pdf_url)
            with urllib.request.urlopen(pdf_req, timeout=timeout) as pdf_resp:
                pdf_ct = pdf_resp.headers.get("Content-Type", "").lower()
                pdf_bytes = pdf_resp.read(5 * 1024 * 1024)

            if "pdf" in pdf_ct or pdf_url.lower().split("?")[0].endswith(".pdf"):
                extracted = _extract_text_from_pdf_bytes(pdf_bytes)
                if extracted:
                    logger.debug("PDF fallback succeeded for %s via %s", url, pdf_url)
                    return True, extracted, "PDF_fallback"
        except Exception as exc:
            logger.debug("PDF fallback failed for %s → %s: %s", url, pdf_url, exc)
            continue

    # If the snippet looks like a JS-only SPA shell (< 300 meaningful chars after
    # stripping tags), escalate through Tavily → Firecrawl → Playwright before
    # returning the near-empty shell as the verified content.
    import re as _re_snap
    _visible = _re_snap.sub(r"<[^>]+>", " ", snippet)
    _visible = _re_snap.sub(r"\s{2,}", " ", _visible).strip()
    if len(_visible) < 300:
        logger.debug(
            "_fetch_url_content: near-empty HTML for %s (%d visible chars) — "
            "trying Tavily/Firecrawl/Playwright.",
            url,
            len(_visible),
        )
        tv = _fetch_via_tavily(url)
        if tv is not None:
            return tv
        fc = _fetch_via_firecrawl(url)
        if fc is not None:
            return fc
        pw = _fetch_via_playwright(url)
        if pw is not None:
            return pw

    return True, snippet, "HTTP"


def _literature_review_to_docx(review_doc: dict[str, Any]) -> bytes:
    """Convert a LiteratureReviewDoc dict to a formatted .docx file.

    Renders:
    - A title heading "Revisão Bibliográfica"
    - Each section as Heading 2 + body paragraphs (Markdown-lite rendering)
    - A "Referências Bibliográficas" section with colour-coded entries:
        ✓ green  — accessible and content verified
        ⚠ orange — accessible but content questionable
        ✗ red    — inaccessible URL

    Args:
        review_doc: A dict matching the LiteratureReviewDoc shape.

    Returns:
        Raw bytes of the generated .docx file, or empty bytes on error.
    """
    try:
        import io
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        logger.warning("python-docx not available — skipping literature review .docx export.")
        return b""

    doc = Document()
    doc.add_heading("Revisão Bibliográfica", level=1)

    sections = review_doc.get("sections", [])
    for section in sections:
        heading = section.get("section_title", "")
        content = section.get("content", "")
        if heading:
            doc.add_heading(heading, level=2)
        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    # Build verification lookup
    verified_map: dict[int, dict[str, Any]] = {}
    for v in review_doc.get("verified_sources", []):
        ref_num = v.get("reference_number", 0)
        verified_map[ref_num] = v  # type: ignore[assignment]

    references = review_doc.get("references", [])
    if references:
        doc.add_heading("Referências Bibliográficas", level=2)
        for ref in references:
            ref_num = ref.get("reference_number", 0)
            abnt_entry = ref.get("abnt_entry", "")
            v_info = verified_map.get(ref_num, {})
            accessible = v_info.get("accessible", False)
            matches = v_info.get("content_matches", False)
            note = v_info.get("verification_note", "")

            if accessible and matches:
                mark = "✓"
                color = RGBColor(0, 128, 0)  # green
            elif accessible:
                mark = "⚠"
                color = RGBColor(204, 102, 0)  # orange
            else:
                mark = "✗"
                color = RGBColor(180, 0, 0)  # red

            para = doc.add_paragraph()
            mark_run = para.add_run(f"[{mark}] ")
            mark_run.font.color.rgb = color
            mark_run.font.bold = True
            mark_run.font.size = Pt(9)

            # Reference number [N] for round-trip lookup
            num_run = para.add_run(f"[{ref_num}] ")
            num_run.font.bold = True
            num_run.font.size = Pt(9)

            entry_run = para.add_run(abnt_entry)
            entry_run.font.size = Pt(9)

            # Access method badge (shown when raw PDF/HTTP fallback was used instead of an API)
            access_method = v_info.get("access_method", "HTTP")
            # Methods that do NOT show the "Acesso direto (sem API)" badge.
            # Playwright is browser automation (not an API) — badge shown.
            _api_methods = {"HTTP", "Semantic_Scholar_API", "Firecrawl", "Tavily"}
            if access_method and access_method not in _api_methods:
                badge_run = para.add_run(f"  [Acesso direto (sem API) — {access_method}]")
                badge_run.font.size = Pt(8)
                badge_run.font.italic = True
                badge_run.font.color.rgb = RGBColor(0, 80, 160)

            if note:
                note_run = para.add_run(f"\n        [{note}]")
                note_run.font.size = Pt(8)
                note_run.font.italic = True
                note_run.font.color.rgb = RGBColor(100, 100, 100)

            # Article summary (up to 200 words)
            summary = ref.get("summary", "")
            if summary:
                summary_para = doc.add_paragraph()
                summary_run = summary_para.add_run(summary)
                summary_run.font.size = Pt(8)
                summary_run.font.italic = True
                summary_run.font.color.rgb = RGBColor(80, 80, 80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _research_design_to_docx(design_doc: dict[str, Any]) -> bytes:
    """Convert a ResearchDesignDoc dict to a formatted .docx file.

    Renders:
    - A title heading "Design de Pesquisa"
    - Each section as Heading 2 + body paragraphs (Markdown-lite rendering)
    - A structured summary block: study type, paradigm, hypotheses (list),
      variables table (nome | tipo | escala)

    Args:
        design_doc: A dict matching the ResearchDesignDoc shape.

    Returns:
        Raw bytes of the generated .docx file, or empty bytes on error.
    """
    try:
        import io
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        logger.warning("python-docx not available — skipping research design .docx export.")
        return b""

    doc = Document()
    doc.add_heading("Design de Pesquisa", level=1)

    # --- Prose sections ---
    sections = design_doc.get("sections", [])
    for section in sections:
        heading = section.get("section_title", "")
        content = section.get("content", "")
        if heading:
            doc.add_heading(heading, level=2)
        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    # --- Structured summary block ---
    doc.add_heading("Resumo Estruturado", level=2)

    study_type = design_doc.get("study_type", "")
    paradigm = design_doc.get("research_paradigm", "")
    stance = design_doc.get("epistemological_stance", "")
    if study_type or paradigm:
        meta_para = doc.add_paragraph()
        meta_para.add_run("Tipo de estudo: ").font.bold = True
        meta_para.add_run(study_type or "—")
        meta_para.add_run("   Paradigma: ").font.bold = True
        meta_para.add_run(paradigm or "—")
        if stance:
            meta_para.add_run("   Epistemologia: ").font.bold = True
            meta_para.add_run(stance)

    hypotheses = design_doc.get("hypotheses", [])
    if hypotheses:
        doc.add_heading("Hipóteses", level=3)
        for h in hypotheses:
            doc.add_paragraph(h, style="List Bullet")

    research_questions = design_doc.get("research_questions", [])
    if research_questions:
        doc.add_heading("Questões de Pesquisa", level=3)
        for q in research_questions:
            doc.add_paragraph(q, style="List Bullet")

    metrics = design_doc.get("metrics_and_kpis", [])
    if metrics:
        doc.add_heading("Métricas e KPIs", level=3)
        for m in metrics:
            doc.add_paragraph(m, style="List Bullet")

    data_sources = design_doc.get("data_sources", [])
    if data_sources:
        doc.add_heading("Fontes de Dados", level=3)
        for ds in data_sources:
            doc.add_paragraph(ds, style="List Bullet")

    # Variables table
    variables = design_doc.get("variables", [])
    if variables:
        doc.add_heading("Tabela de Variáveis", level=3)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, label in enumerate(["Nome", "Tipo", "Escala", "Operacionalização"]):
            hdr[i].text = label
            run = hdr[i].paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(9)
        for var in variables:
            row = tbl.add_row().cells
            row[0].text = var.get("name", "")
            row[1].text = var.get("variable_type", "")
            row[2].text = var.get("measurement_scale", "")
            row[3].text = var.get("operationalization", "")
            for cell in row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    # Methodology timeline
    timeline = design_doc.get("methodology_timeline", "")
    if timeline:
        doc.add_heading("Cronograma Metodológico", level=3)
        for line in timeline.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    reporting = design_doc.get("reporting_standard", "")
    journal_tier = design_doc.get("target_journal_tier", "")
    if reporting or journal_tier:
        meta2 = doc.add_paragraph()
        if reporting:
            meta2.add_run("Norma de relato: ").font.bold = True
            meta2.add_run(reporting)
        if journal_tier:
            meta2.add_run("   Alvo de publicação: ").font.bold = True
            meta2.add_run(journal_tier)
        for run in meta2.runs:
            run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _get_default_registry() -> SkillRegistry:
    """Create and auto-discover a SkillRegistry instance.

    Returns:
        A SkillRegistry with all available built-in skills registered.
    """
    from ai_skill.skills.registry import SkillRegistry
    registry = SkillRegistry()
    registry.auto_discover()
    return registry


def _topological_batches(steps: list[PlanStep]) -> list[list[PlanStep]]:
    """Group plan steps into batches that can run concurrently.

    Steps with no unresolved dependencies form a batch. Each batch runs
    after the previous one completes.

    Args:
        steps: List of PlanStep with step_id and depends_on fields.

    Returns:
        List of batches (each batch is a list of steps to run concurrently).
    """
    completed: set[int] = set()
    remaining = list(steps)
    batches: list[list[PlanStep]] = []

    max_iterations = len(steps) + 1  # safety valve against cycles
    iterations = 0

    while remaining and iterations < max_iterations:
        batch = [
            s for s in remaining
            if all(dep in completed for dep in s.get("depends_on", []))
        ]
        if not batch:
            # Remaining steps have unresolvable dependencies — add them as-is
            batches.append(remaining)
            break
        batches.append(batch)
        for step in batch:
            completed.add(step["step_id"])
        remaining = [s for s in remaining if s["step_id"] not in completed]
        iterations += 1

    return batches


async def _run_batch_async(
    registry: SkillRegistry,
    batch_inputs: list[tuple[int, SkillInput]],
) -> dict[int, SkillOutput]:
    """Run a batch of skill steps concurrently using asyncio.gather.

    Args:
        registry: The skill registry to look up skill instances.
        batch_inputs: List of (step_id, SkillInput) pairs.

    Returns:
        Dict mapping step_id → SkillOutput.
    """
    async def run_one(step_id: int, skill_input: SkillInput) -> tuple[int, SkillOutput]:
        skill_name = skill_input.parameters.get("_skill_name", "")
        # The skill_name is stored in the parameters under a special key
        # set by the execute node; extract it before passing to the skill.
        clean_params = {k: v for k, v in skill_input.parameters.items() if k != "_skill_name"}
        skill_input_clean = SkillInput({**dict(skill_input), "parameters": clean_params})

        try:
            skill = registry.get(skill_name)
            output = await skill.arun(skill_input_clean)
        except Exception as exc:
            logger.error("Skill '%s' raised an exception: %s", skill_name, exc)
            output = SkillOutput(
                skill_name=skill_name,
                result={},
                confidence=0.0,
                sources=[],
                error=str(exc),
                cached=False,
            )
        return step_id, output

    tasks = [run_one(sid, sinput) for sid, sinput in batch_inputs]
    results = await asyncio.gather(*tasks)
    return dict(results)


# ===========================================================================
# CP4 — Data Collection Guide
# ===========================================================================
# Three-phase pipeline:
#   Phase 1 — draft_collection_guide      (operationalise CP3 into 8-section guide)
#   Phase 2 — review_collection_standards (web critique: FAIR/ISO/PMBOK/PRISMA/ABNT)
#   Phase 3 — evaluate_collection_objectives (score 8 sections vs CP1 + CP3)
#
# Principle: agent is methodologist only. Never generates/fabricates research data.
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Pydantic structured-output models for CP4
# ---------------------------------------------------------------------------


class _CollectionSectionLLM(BaseModel):
    """One section of the Data Collection Guide."""

    section_title: str = Field(description="Section heading in pt-BR.")
    content: str = Field(description="Body text in pt-BR Markdown (prose, templates, checklists).")


class _CollectionStepLLM(BaseModel):
    """One step in the step-by-step collection protocol."""

    step_id: int = Field(description="Sequential step number (1-based).")
    description: str = Field(description="What to do in this step.")
    responsible: str = Field(description="Role or person responsible.")
    tool: str = Field(description="Tool, platform, or instrument used.")
    acceptance_criterion: str = Field(
        description="Observable evidence that this step was completed successfully."
    )


class _DataDictionaryEntryLLM(BaseModel):
    """One variable definition in the data dictionary."""

    variable: str = Field(description="Variable/column name (snake_case).")
    type: str = Field(description="Data type: string | integer | float | boolean | date | categorical.")
    unit: str = Field(description="Unit of measurement, or 'N/A'.")
    encoding: str = Field(description="Value encoding (e.g. 0=Não, 1=Sim) or 'free text'.")
    nullable: bool = Field(description="Whether the field may be missing/null.")


class _ContingencyPlanLLM(BaseModel):
    """One contingency scenario."""

    trigger: str = Field(description="Observable condition that activates this contingency.")
    action: str = Field(description="Steps to take when triggered.")
    responsible: str = Field(description="Role or person responsible for executing the action.")


class _DataCollectionGuideDocLLM(BaseModel):
    """Full structured Data Collection Guide — enforces consistent output format for CP4."""

    sections: list[_CollectionSectionLLM] = Field(
        min_length=1,
        description="Ordered sections covering all 8 mandatory headings.",
    )
    instruments: list[str] = Field(
        default_factory=list,
        description="List of collection instruments (name, type, version, reference).",
    )
    collection_steps: list[_CollectionStepLLM] = Field(
        default_factory=list,
        description="Ordered step-by-step protocol.",
    )
    sampling_strategy: str = Field(description="Sampling approach description.")
    min_sample_size: int = Field(default=0, description="Minimum required sample size (n).")
    sample_size_rationale: str = Field(description="Power analysis or saturation rationale.")
    inclusion_criteria: list[str] = Field(
        default_factory=list, description="Participant/data inclusion criteria."
    )
    exclusion_criteria: list[str] = Field(
        default_factory=list, description="Participant/data exclusion criteria."
    )
    data_dictionary: list[_DataDictionaryEntryLLM] = Field(
        default_factory=list, description="Variable definitions."
    )
    data_format: str = Field(
        description="Expected delivery format: CSV | JSON | XLSX | SQL | outro."
    )
    acceptance_criteria_per_step: list[str] = Field(
        default_factory=list,
        description="Quality gate per collection step (same order as collection_steps).",
    )
    tcle_elements: list[str] = Field(
        default_factory=list,
        description="Minimum TCLE elements per CNS 466/2012.",
    )
    lgpd_measures: list[str] = Field(
        default_factory=list,
        description="LGPD compliance measures (pseudonymisation, retention, etc.).",
    )
    cep_required: bool = Field(
        default=False, description="Whether CEP/CONEP submission is required."
    )
    pre_collection_checklist: list[str] = Field(
        default_factory=list, description="Ordered checklist items before starting collection."
    )
    contingency_plans: list[_ContingencyPlanLLM] = Field(
        default_factory=list, description="Contingency scenarios."
    )


# ---------------------------------------------------------------------------
# CP4 router and entry-point
# ---------------------------------------------------------------------------


def cp4_router(_state: ResearchState) -> dict[str, Any]:
    """Entry-point no-op for CP4; routing logic lives in route_cp4_start.

    Returns:
        Empty dict (no state changes).
    """
    return {}


def route_cp4_start(state: ResearchState) -> str:
    """Conditional edge from cp4_router.

    Returns ``"refine_collection_guide"`` when a guide exists and the user has
    provided feedback (correction cycle).  Otherwise starts fresh from
    ``"draft_collection_guide"``.

    Args:
        state: Current ResearchState.

    Returns:
        Name of the next node.
    """
    if state.get("data_collection_guide_doc") and state.get("user_feedback"):
        return "refine_collection_guide"
    return "draft_collection_guide"


# ---------------------------------------------------------------------------
# Phase 1 — draft_collection_guide
# ---------------------------------------------------------------------------


def draft_collection_guide(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Phase 1 of CP4: operationalise CP3 into the full Data Collection Guide.

    Reads the approved ResearchDesignDoc (CP3) and the CP1 research objectives,
    then proposes a complete 8-section guide with instruments, protocol,
    sampling, data specification, acceptance criteria, ethics/legal compliance,
    pre-collection checklist, and contingency plan.

    The LLM is expected to PROPOSE — not merely summarise CP3.  It may go
    beyond the CP3 text wherever additional operational precision is needed.

    Args:
        state: Current ResearchState with research_design_doc (CP3 approved)
            and objective (CP1).
        llm: Optional LLMClient override (useful for testing).

    Returns:
        Partial state with data_collection_guide_doc and status=EVALUATING.
    """
    _configure_history(state, "draft_collection_guide")
    _llm = llm or LLMClient()

    design_doc = state.get("research_design_doc") or {}
    objective = state.get("cp4_context") or state.get("objective") or {}
    preserved_sections: dict[str, Any] = state.get("cp4_preserved_sections") or {}
    evaluation = state.get("evaluation")
    gaps: list[str] = evaluation.get("gaps", []) if evaluation else []

    system, messages = build_draft_collection_guide_messages(
        objective=objective,
        design_doc=dict(design_doc),
        preserved_sections=preserved_sections,
        gaps=gaps,
    )

    try:
        guide_obj: _DataCollectionGuideDocLLM = _llm.complete_structured(
            messages=messages,
            response_model=_DataCollectionGuideDocLLM,
            system=system,
            max_tokens=32768,
        )
    except LLMClientError as exc:
        logger.error("draft_collection_guide failed: %s", exc)
        return {"status": ResearchStatus.FAILED}

    sections: list[DataCollectionSection] = [
        DataCollectionSection(section_title=s.section_title, content=s.content)
        for s in guide_obj.sections
    ]
    guide_doc = DataCollectionGuideDoc(
        sections=sections,
        instruments=guide_obj.instruments,
        collection_steps=[step.model_dump() for step in guide_obj.collection_steps],
        sampling_strategy=guide_obj.sampling_strategy,
        min_sample_size=guide_obj.min_sample_size,
        sample_size_rationale=guide_obj.sample_size_rationale,
        inclusion_criteria=guide_obj.inclusion_criteria,
        exclusion_criteria=guide_obj.exclusion_criteria,
        data_dictionary=[e.model_dump() for e in guide_obj.data_dictionary],
        data_format=guide_obj.data_format,
        acceptance_criteria_per_step=guide_obj.acceptance_criteria_per_step,
        tcle_elements=guide_obj.tcle_elements,
        lgpd_measures=guide_obj.lgpd_measures,
        cep_required=guide_obj.cep_required,
        pre_collection_checklist=guide_obj.pre_collection_checklist,
        contingency_plans=[cp.model_dump() for cp in guide_obj.contingency_plans],
    )
    return {
        "data_collection_guide_doc": guide_doc,
        "status": ResearchStatus.EVALUATING,
    }


# ---------------------------------------------------------------------------
# Phase 2 — review_collection_standards
# ---------------------------------------------------------------------------


def review_collection_standards(
    state: ResearchState,
    registry: SkillRegistry | None = None,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Phase 2 of CP4: critique guide against FAIR, ISO 25012, PMBOK, PRISMA, ABNT.

    Runs web searches for the relevant standards (always authorised — no Y/n
    required) and calls the LLM to apply targeted corrections to the guide.
    Sections already compliant are returned verbatim.

    Args:
        state: Current ResearchState with data_collection_guide_doc from Phase 1
            and research_design_doc (CP3, used as reference to avoid contradictions).
        registry: Optional injected SkillRegistry (for testing).
        llm: Optional LLMClient override.

    Returns:
        Partial state with (possibly corrected) data_collection_guide_doc.
        Returns empty dict on LLM error (keeps Phase 1 draft).
    """
    _configure_history(state, "review_collection_standards")
    _registry = registry or _get_default_registry()
    _llm = llm or LLMClient()

    guide_doc = state.get("data_collection_guide_doc")
    if not guide_doc:
        logger.warning("review_collection_standards: no data_collection_guide_doc — skipping.")
        return {}

    design_doc = state.get("research_design_doc") or {}

    # Web searches for methodology standards (always allowed)
    web_skill = _registry.get("web_search")
    search_snippets: list[str] = []
    pmbok_context = ""
    prisma_traice_context = ""

    if web_skill is not None:
        for query in CP4_FRAMEWORK_QUERIES:
            skill_input = SkillInput({
                "parameters": {
                    "query": query,
                    "max_results": 5,
                    "language": "en",
                    "_skill_name": "web_search",
                },
                "objective": state.get("objective"),
                "stage": PipelineStage.DATA_COLLECTION_GUIDE.value,
                "attempt": state.get("attempt", 0),
            })
            out = web_skill.run(skill_input)
            if not out.get("error"):
                results = (out.get("result") or {}).get("results", [])
                for r in results[:3]:
                    snippet = r.get("snippet") or r.get("description") or ""
                    title = r.get("title") or ""
                    if snippet:
                        search_snippets.append(f"[{title}] {snippet}")
                        if "PMBOK" in title or "PMBOK" in snippet:
                            pmbok_context += f"\n{snippet}"
                        if "trAIce" in title or "trAIce" in snippet:
                            prisma_traice_context += f"\n{snippet}"

    framework_search_results = "\n".join(search_snippets)

    system, messages = build_review_collection_standards_messages(
        guide_doc=dict(guide_doc),
        design_doc=dict(design_doc),
        framework_search_results=framework_search_results,
        pmbok_context=pmbok_context.strip(),
        prisma_traice_context=prisma_traice_context.strip(),
    )

    try:
        revised_obj: _DataCollectionGuideDocLLM = _llm.complete_structured(
            messages=messages,
            response_model=_DataCollectionGuideDocLLM,
            system=system,
            max_tokens=32768,
        )
    except LLMClientError as exc:
        logger.error("review_collection_standards failed: %s — keeping Phase 1 draft", exc)
        return {}

    sections: list[DataCollectionSection] = [
        DataCollectionSection(section_title=s.section_title, content=s.content)
        for s in revised_obj.sections
    ]
    revised_doc = DataCollectionGuideDoc(
        sections=sections,
        instruments=revised_obj.instruments,
        collection_steps=[step.model_dump() for step in revised_obj.collection_steps],
        sampling_strategy=revised_obj.sampling_strategy,
        min_sample_size=revised_obj.min_sample_size,
        sample_size_rationale=revised_obj.sample_size_rationale,
        inclusion_criteria=revised_obj.inclusion_criteria,
        exclusion_criteria=revised_obj.exclusion_criteria,
        data_dictionary=[e.model_dump() for e in revised_obj.data_dictionary],
        data_format=revised_obj.data_format,
        acceptance_criteria_per_step=revised_obj.acceptance_criteria_per_step,
        tcle_elements=revised_obj.tcle_elements,
        lgpd_measures=revised_obj.lgpd_measures,
        cep_required=revised_obj.cep_required,
        pre_collection_checklist=revised_obj.pre_collection_checklist,
        contingency_plans=[cp.model_dump() for cp in revised_obj.contingency_plans],
    )
    return {"data_collection_guide_doc": revised_doc}


# ---------------------------------------------------------------------------
# Phase 3 — evaluate_collection_objectives
# ---------------------------------------------------------------------------


class _CollectionSectionScoreLLM(BaseModel):
    """Score for a single collection guide section."""

    section_title: str = Field(description="Title of the section being scored.")
    score: float = Field(description="Quality score 0.0–1.0.", ge=0.0, le=1.0)
    rationale: str = Field(description="One-sentence justification for the score.")
    gaps: list[str] = Field(
        default_factory=list,
        description="Operational gaps vs. CP1+CP3 (empty when score >= 0.85).",
    )


class _EvaluateCollectionObjectivesLLM(BaseModel):
    """Per-section scores and aggregated gaps from evaluate_collection_objectives."""

    sections: list[_CollectionSectionScoreLLM] = Field(
        description="One score entry per mandatory guide section."
    )
    gaps: list[str] = Field(
        description="Aggregated list of gaps across all sections with score < 0.85."
    )


def evaluate_collection_objectives(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Phase 3 of CP4: score the guide against CP1 objectives and CP3 design.

    Scores each of the 8 mandatory sections individually. Sections scoring
    >= 0.85 are preserved verbatim in subsequent draft_collection_guide calls.
    Computes total_score as the mean of section scores and sets
    converged = total_score >= _CONVERGENCE_THRESHOLD.

    Args:
        state: Current ResearchState with data_collection_guide_doc from Phase 2.
        llm: Optional LLMClient override.

    Returns:
        Partial state with evaluation, cp4_preserved_sections, attempt+1,
        and quality_history snapshot.
    """
    _configure_history(state, "evaluate_collection_objectives")
    _llm = llm or LLMClient()

    guide_doc = state.get("data_collection_guide_doc")
    if not guide_doc:
        logger.warning("evaluate_collection_objectives: no data_collection_guide_doc — failing.")
        return {
            "evaluation": EvaluationResult(
                per_metric=[], total_score=0.0, converged=False,
                gaps=["data_collection_guide_doc ausente"], regression=False,
            ),
            "attempt": state.get("attempt", 0) + 1,
        }

    design_doc = state.get("research_design_doc") or {}
    objective = state.get("cp4_context") or state.get("objective") or {}
    attempt = state.get("attempt", 0)
    quality_history = state.get("quality_history", [])
    stage_str = PipelineStage.DATA_COLLECTION_GUIDE.value

    system, messages = build_evaluate_collection_objectives_messages(
        objective=objective,
        design_doc=dict(design_doc),
        guide_doc=dict(guide_doc),
        threshold=_CONVERGENCE_THRESHOLD,
    )

    try:
        eval_obj: _EvaluateCollectionObjectivesLLM = _llm.complete_structured(
            messages=messages,
            response_model=_EvaluateCollectionObjectivesLLM,
            system=system,
        )
    except LLMClientError as exc:
        logger.error("evaluate_collection_objectives failed: %s", exc)
        eval_obj = _EvaluateCollectionObjectivesLLM(
            sections=[],
            gaps=[f"Evaluation error: {exc}"],
        )

    # Compute total_score deterministically
    scores = [s.score for s in eval_obj.sections]
    total_score = sum(scores) / len(scores) if scores else 0.0
    converged = total_score >= _CONVERGENCE_THRESHOLD

    # Regression detection
    previous_score = quality_history[-1]["total_score"] if quality_history else None
    regression = (
        previous_score is not None and (previous_score - total_score) > 0.05
    )

    per_metric: list[MetricScore] = [
        MetricScore(
            metric=s.section_title,
            score=s.score,
            rationale=s.rationale,
            gaps=s.gaps,
        )
        for s in eval_obj.sections
    ]

    evaluation = EvaluationResult(
        per_metric=per_metric,
        total_score=total_score,
        converged=converged,
        gaps=eval_obj.gaps,
        regression=regression,
    )

    # Preserve sections scoring >= 0.85 for next iteration
    guide_sections: list[dict[str, Any]] = guide_doc.get("sections") or []
    section_content_by_title: dict[str, str] = {
        s.get("section_title", ""): s.get("content", "")
        for s in guide_sections
    }
    preserved: dict[str, Any] = {}
    for s in eval_obj.sections:
        if s.score >= 0.85:
            content = section_content_by_title.get(s.section_title, "")
            if content:
                preserved[s.section_title] = {"content": content, "score": s.score}

    snapshot = QualitySnapshot(
        attempt=attempt,
        stage=stage_str,
        total_score=total_score,
        per_metric_scores={s.section_title: s.score for s in eval_obj.sections},
        skills_used=["draft_collection_guide", "review_collection_standards"],
        cache_hit_rate=0.0,
    )

    new_quality_history = list(quality_history) + [snapshot]
    stage_quality_history = dict(state.get("stage_quality_history") or {})
    stage_quality_history.setdefault(stage_str, [])
    stage_quality_history[stage_str] = list(stage_quality_history[stage_str]) + [snapshot]

    if regression:
        logger.warning(
            "Quality regression in CP4: %.2f → %.2f (attempt %d)",
            previous_score, total_score, attempt,
        )

    return {
        "evaluation": evaluation,
        "cp4_preserved_sections": preserved,
        "attempt": attempt + 1,
        "quality_history": new_quality_history,
        "stage_quality_history": stage_quality_history,
        "status": ResearchStatus.PLANNING,
    }


def route_after_evaluate_collection(state: ResearchState) -> str:
    """Conditional edge after evaluate_collection_objectives in CP4.

    Routes to deliver_collection_guide on convergence, draft_collection_guide
    for retry, or request_support when max retries are exhausted.

    Args:
        state: Current ResearchState with evaluation result.

    Returns:
        Name of the next node.
    """
    evaluation = state.get("evaluation")
    attempt = state.get("attempt", 0)
    max_retries = int(os.environ.get("AI_SKILL_MAX_RETRIES", "5"))

    if evaluation and evaluation.get("converged"):
        return "deliver_collection_guide"
    if attempt >= max_retries:
        return "request_support"
    return "draft_collection_guide"


# ---------------------------------------------------------------------------
# Deliver, review gate, refine — CP4
# ---------------------------------------------------------------------------


def deliver_collection_guide(state: ResearchState) -> dict[str, Any]:
    """Generate the Data Collection Guide .docx checkpoint preview.

    Converts the ``data_collection_guide_doc`` to a formatted Word document
    and saves it as ``Checkpoint 4 - Guia de Coleta [preview_N].docx``.

    Args:
        state: Current ResearchState with a populated data_collection_guide_doc.

    Returns:
        Partial state update with checkpoint label and collection_guide_approved=False.
    """
    _configure_history(state, "deliver_collection_guide")

    guide_doc = state.get("data_collection_guide_doc") or {}
    workspace_path = state.get("workspace_path", "")

    docx_bytes = _collection_guide_to_docx(guide_doc)

    checkpoint_label = ""
    if workspace_path:
        pw = _get_project_workspace(workspace_path)
        if pw and docx_bytes:
            checkpoint_path = pw.save_checkpoint_preview(4, docx_bytes)
            checkpoint_label = str(checkpoint_path)
            logger.info("Collection guide checkpoint saved: %s", checkpoint_label)

    return {
        "checkpoint_label": checkpoint_label,
        "collection_guide_approved": False,
        "status": ResearchStatus.PLANNING,
    }


def review_collection_guide(state: ResearchState) -> dict[str, Any]:
    """Gate node: check whether the researcher approved the Data Collection Guide.

    Runs after the interrupt fires and the user resumes the graph.  If
    ``user_feedback`` is present, the guide needs revision; otherwise approved.

    Args:
        state: Current ResearchState, resumed after user review.

    Returns:
        Partial state update setting collection_guide_approved and status.
    """
    _configure_history(state, "review_collection_guide")

    if state.get("user_feedback"):
        return {
            "collection_guide_approved": False,
            "status": ResearchStatus.PLANNING,
        }
    return {
        "collection_guide_approved": True,
        "status": ResearchStatus.COMPLETED,
    }


def route_after_review_collection(state: ResearchState) -> str:
    """Conditional edge after review_collection_guide.

    Returns:
        ``"END"`` if approved, ``"refine_collection_guide"`` otherwise.
    """
    return "END" if state.get("collection_guide_approved") else "refine_collection_guide"


def refine_collection_guide(
    state: ResearchState,
    llm: LLMClient | None = None,
) -> dict[str, Any]:
    """Apply researcher corrections to the Data Collection Guide.

    Called when the user provides feedback at the Checkpoint 4 review gate.
    Applies surgical corrections while preserving all unmarked content.

    Args:
        state: Current ResearchState with data_collection_guide_doc and user_feedback.
        llm: Optional LLMClient override.

    Returns:
        Partial state update with the refined data_collection_guide_doc and
        user_feedback=None.
    """
    _configure_history(state, "refine_collection_guide")
    _llm = llm or LLMClient()

    guide_doc = state.get("data_collection_guide_doc") or {}
    feedback = state.get("user_feedback") or ""

    system, messages = build_refine_collection_guide_messages(
        guide_doc=dict(guide_doc),
        feedback=feedback,
    )

    try:
        guide_obj: _DataCollectionGuideDocLLM = _llm.complete_structured(
            messages=messages,
            response_model=_DataCollectionGuideDocLLM,
            system=system,
            max_tokens=32768,
        )
    except LLMClientError as exc:
        logger.error("refine_collection_guide failed: %s", exc)
        return {"status": ResearchStatus.FAILED}

    sections: list[DataCollectionSection] = [
        DataCollectionSection(section_title=s.section_title, content=s.content)
        for s in guide_obj.sections
    ]
    updated_doc = DataCollectionGuideDoc(
        sections=sections,
        instruments=guide_obj.instruments,
        collection_steps=[step.model_dump() for step in guide_obj.collection_steps],
        sampling_strategy=guide_obj.sampling_strategy,
        min_sample_size=guide_obj.min_sample_size,
        sample_size_rationale=guide_obj.sample_size_rationale,
        inclusion_criteria=guide_obj.inclusion_criteria,
        exclusion_criteria=guide_obj.exclusion_criteria,
        data_dictionary=[e.model_dump() for e in guide_obj.data_dictionary],
        data_format=guide_obj.data_format,
        acceptance_criteria_per_step=guide_obj.acceptance_criteria_per_step,
        tcle_elements=guide_obj.tcle_elements,
        lgpd_measures=guide_obj.lgpd_measures,
        cep_required=guide_obj.cep_required,
        pre_collection_checklist=guide_obj.pre_collection_checklist,
        contingency_plans=[cp.model_dump() for cp in guide_obj.contingency_plans],
    )
    return {
        "data_collection_guide_doc": updated_doc,
        "user_feedback": None,
        "status": ResearchStatus.EXECUTING,
    }


# ---------------------------------------------------------------------------
# Private helper: convert DataCollectionGuideDoc → .docx bytes
# ---------------------------------------------------------------------------


def _collection_guide_to_docx(guide_doc: dict[str, Any]) -> bytes:
    """Convert a DataCollectionGuideDoc dict to a formatted .docx file.

    Renders:
    - Title heading "Guia de Coleta de Dados"
    - Each section as Heading 2 + body paragraphs (Markdown-lite rendering)
    - A structured summary block: data format, sample size, pre-collection checklist

    Args:
        guide_doc: A dict matching the DataCollectionGuideDoc shape.

    Returns:
        Raw bytes of the generated .docx file, or empty bytes on error.
    """
    try:
        import io
        from docx import Document
    except ImportError:
        logger.warning("python-docx not available — skipping collection guide .docx export.")
        return b""

    doc = Document()
    doc.add_heading("Guia de Coleta de Dados", level=1)

    # --- Prose sections ---
    sections = guide_doc.get("sections", [])
    for section in sections:
        heading = section.get("section_title", "")
        content = section.get("content", "")
        if heading:
            doc.add_heading(heading, level=2)
        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    # --- Structured summary block ---
    doc.add_heading("Resumo Estruturado", level=2)

    data_format = guide_doc.get("data_format", "")
    min_n = guide_doc.get("min_sample_size", 0)
    rationale = guide_doc.get("sample_size_rationale", "")
    if data_format or min_n:
        meta = doc.add_paragraph()
        meta.add_run("Formato dos dados: ").font.bold = True
        meta.add_run(data_format or "—")
        if min_n:
            meta.add_run(f"   N mínimo: ").font.bold = True
            meta.add_run(str(min_n))

    if rationale:
        p = doc.add_paragraph()
        p.add_run("Justificativa amostral: ").font.bold = True
        p.add_run(rationale)

    checklist = guide_doc.get("pre_collection_checklist", [])
    if checklist:
        doc.add_heading("Checklist Pré-Coleta", level=3)
        for item in checklist:
            doc.add_paragraph(item, style="List Bullet")

    contingencies = guide_doc.get("contingency_plans", [])
    if contingencies:
        doc.add_heading("Plano de Contingência", level=3)
        for cp in contingencies:
            trigger = cp.get("trigger", "")
            action = cp.get("action", "")
            responsible = cp.get("responsible", "")
            if trigger:
                p = doc.add_paragraph()
                p.add_run("Gatilho: ").font.bold = True
                p.add_run(trigger)
                p.add_run("  →  Ação: ").font.bold = True
                p.add_run(action)
                if responsible:
                    p.add_run("  (").font.bold = False
                    p.add_run(responsible)
                    p.add_run(")")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
